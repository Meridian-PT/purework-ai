#!/usr/bin/env python3
"""
Build Madison & Wall Benefits Executive Summary - One Page
Prepared by The People Group | July 2026
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

OUTPUT = "/home/aiciv/portal_uploads/to-portal/MadisonWall_Benefits_ExecSummary_July2026.docx"
NAVY = RGBColor(0x1A, 0x3A, 0x5C)
GRAY = RGBColor(0x66, 0x66, 0x66)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def set_cell_shading(cell, hex_color):
    """Apply background shading to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_margins(cell, top=20, bottom=20, left=40, right=40):
    """Set tight cell margins in twips."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:start w:w="{left}" w:type="dxa"/>'
        f'  <w:end w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)


def set_row_height(row, height_twips):
    """Set exact row height."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = parse_xml(
        f'<w:trHeight {nsdecls("w")} w:val="{height_twips}" w:hRule="atLeast"/>'
    )
    trPr.append(trHeight)


def style_header_row(row, font_size=Pt(8)):
    """Style a table header row with navy background and white bold text."""
    for cell in row.cells:
        set_cell_shading(cell, "1A3A5C")
        set_cell_margins(cell, top=15, bottom=15, left=40, right=40)
        for p in cell.paragraphs:
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = Pt(10)
            for run in p.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = font_size
                run.font.name = "Calibri"


def style_data_cell(cell, font_size=Pt(9)):
    """Style a data cell with tight spacing."""
    set_cell_margins(cell, top=12, bottom=12, left=40, right=40)
    for p in cell.paragraphs:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(10)
        for run in p.runs:
            run.font.size = font_size
            run.font.name = "Calibri"
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def add_table_borders(table, color="999999"):
    """Add thin borders to all cells."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def make_table(doc, headers, rows, col_widths=None):
    """Create a compact styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.name = "Calibri"
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.color.rgb = WHITE
    style_header_row(table.rows[0], font_size=Pt(8))

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = "Calibri"
            run.font.size = Pt(9)
            style_data_cell(cell, font_size=Pt(9))
        # Alternate row shading
        if r_idx % 2 == 1:
            for cell in row.cells:
                set_cell_shading(cell, "F2F6FA")

    add_table_borders(table)

    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    return table


def add_section_heading(doc, text):
    """Add a compact section heading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = Pt(12)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = NAVY
    return p


def add_body_text(doc, text, font_size=Pt(9)):
    """Add compact body text."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = Pt(11)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = font_size
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p


def add_horizontal_rule(doc):
    """Add a thin horizontal line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def build():
    doc = Document()

    # ---- Page Setup: Narrow margins (0.7" all around) ----
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    # Letter size
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)

    # ---- Default paragraph style ----
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(9)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.line_spacing = Pt(11)

    # ============================================================
    # HEADER
    # ============================================================
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(2)
    run = p_title.add_run("Madison & Wall Inc. - Employee Benefits Summary")
    run.font.name = "Calibri"
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = NAVY

    # Sub-header
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(2)
    run = p_sub.add_run("Prepared by The People Group | July 2026 | Effective August 1, 2026")
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY

    # Horizontal rule
    add_horizontal_rule(doc)

    # ============================================================
    # SECTION 1: How Canadian Healthcare Works
    # ============================================================
    add_section_heading(doc, "How Canadian Healthcare Works")
    add_body_text(
        doc,
        "OHIP (Ontario\u2019s universal healthcare) covers doctors, hospitals, ER, surgery, and diagnostics at no cost. "
        "Your employer plan fills the gaps: dental, prescriptions, vision, paramedical, and disability coverage."
    )

    # ============================================================
    # SECTION 2: Your Benefits at a Glance
    # ============================================================
    add_section_heading(doc, "Your Benefits at a Glance")

    benefits_headers = ["Tier", "What It Is", "M&W Annual Cost", "Employee Cost"]
    benefits_rows = [
        ["OHIP (Government)", "Doctors, hospitals, ER, surgery", "$0", "$0"],
        ["Group Benefits (Manulife)", "Rx, dental, vision, paramedical, life insurance", "$8,561/employee/yr", "$0"],
        ["Long-Term Disability", "66.7% income replacement if disabled", "$0", "$85.20/mo (tax-free if claimed)"],
        ["Critical Illness + Virtual Care", "$10K lump sum + 24/7 TELUS Health", "Included above", "$0"],
        ["Health Spending Account", "$5,000\u2013$10,000/yr for uncovered expenses", "$5,000\u2013$10,000/employee/yr", "$0"],
        ["Executive Medical (Medcan)", "Comprehensive annual physical", "$3,600/person/yr", "$0"],
        ["Concierge GP (Optional)", "Private GP relationship", "$0\u2013$2,195/family/yr", "$0"],
        ["International Health Insurance", "$5M lifetime, US facility access", "Quote from Amr", "$0"],
    ]
    make_table(doc, benefits_headers, benefits_rows, col_widths=[1.6, 2.3, 1.7, 1.5])

    # ============================================================
    # SECTION 3: Why LTD Is Employee-Paid
    # ============================================================
    add_section_heading(doc, "Why LTD Is Employee-Paid")
    add_body_text(
        doc,
        "If M&W pays the LTD premium, any benefit received is taxable income to the employee. "
        "If the employee pays, the benefit is tax-free. At $85.20/month, this is a small cost for significant tax protection. "
        "This is standard practice in Canada."
    )

    # ============================================================
    # SECTION 4: Why You Need an HSA
    # ============================================================
    add_section_heading(doc, "Why You Need an HSA")
    add_body_text(
        doc,
        "Your Manulife dental plan excludes major dental and orthodontics (requires 3+ employees). "
        "A Health Spending Account through getmyhsa.com fills this gap. "
        "HSA dollars are 100% tax-deductible to M&W and tax-free to you. "
        "CRA guideline: cannot exceed 25% of gross salary."
    )

    # ============================================================
    # SECTION 5: Total Estimated Annual Cost to M&W
    # ============================================================
    add_section_heading(doc, "Total Estimated Annual Cost to M&W")

    cost_headers = ["Component", "Per Employee", "Total (2 Employees)"]
    cost_rows = [
        ["Group Benefits (excl. LTD)", "$8,561", "$17,122"],
        ["HSA ($7,500 midpoint)", "$7,500", "$15,000"],
        ["Medcan Annual Assessment", "$3,600", "$7,200"],
        ["TOTAL", "$19,661", "$39,322"],
    ]
    cost_table = make_table(doc, cost_headers, cost_rows, col_widths=[3.0, 2.0, 2.1])

    # Bold the TOTAL row
    total_row = cost_table.rows[-1]
    for cell in total_row.cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = NAVY

    # ============================================================
    # SECTION 6: Our Recommendations
    # ============================================================
    add_section_heading(doc, "Our Recommendations")

    rec_headers = ["Item", "Recommendation"]
    rec_rows = [
        ["Group Benefits (Manulife)", "Proceed. M&W pays 100% except LTD. Strong plan, no changes needed."],
        ["Long-Term Disability", "Employee-paid. Ensures tax-free benefit if ever claimed. Standard in Canada."],
        ["Critical Illness + Virtual Care", "Include. $57/month combined. Low cost, high signal to employees."],
        ["Health Spending Account", "Yes, $5,000-$10,000/person via getmyhsa.com. Fills major dental/ortho gap."],
        ["Executive Medical (Medcan)", "Yes, $3,600/person. Market expectation at executive level. Preventive value."],
        ["Concierge GP (Tier 4)", "Skip. OHIP + Medcan annual covers you. Concierge is above market."],
        ["International Health Insurance", "Worth it if you want US facility access. Take the $10K or $20K deductible."],
    ]
    rec_table = make_table(doc, rec_headers, rec_rows, col_widths=[2.2, 4.9])

    # ============================================================
    # FOOTER
    # ============================================================
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_footer.paragraph_format.space_before = Pt(10)
    p_footer.paragraph_format.space_after = Pt(0)
    run = p_footer.add_run("The People Group | mike.daser@thepeoplegroup.ca | Confidential")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = GRAY

    # Save
    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")
    print(f"Size: {os.path.getsize(OUTPUT)} bytes")


if __name__ == "__main__":
    build()
