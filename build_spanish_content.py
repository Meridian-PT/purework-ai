#!/usr/bin/env python3
"""
Build 8 Spanish-language HR content documents for Pure Work AI.

Documents:
  1. Policy_Harassment_Spanish.docx
  2. Policy_HealthSafety_Spanish.docx
  3. Policy_DrugAlcohol_Spanish.docx
  4. Policy_Attendance_Spanish.docx
  5. Onboarding_Spanish.docx
  6. Training_SafetyOrientation_Spanish.docx
  7. SOP_LOTO_Spanish.docx
  8. Policy_PPE_Spanish.docx

Formatting specs (matching existing PureWorkAI files):
- Normal: Calibri 11pt, color #333333
- Heading 1: Calibri 18pt, bold, color #1A3A5C
- Heading 2: Calibri 14pt, bold, color #1A3A5C
- Heading 3: Calibri 12pt, bold, color #1A3A5C
- Title line: Calibri 22pt, bold, color #1A3A5C, centered
- Jurisdiction tag: Calibri 11pt, italic, color #666666, centered
- Table cells: Calibri 10pt
- Table header rows: shading #1A3A5C with white bold text
- Footer: "Preparado por Pure Work AI | Confidencial" on every page
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml
import os

# ---------------------------------------------------------------------------
# Directories
# ---------------------------------------------------------------------------
BASE = "/home/aiciv/purework-ai/public/content"
POLICY_DIR = os.path.join(BASE, "policies")
ONBOARD_DIR = os.path.join(BASE, "onboarding")
TRAIN_DIR = os.path.join(BASE, "training")
SOP_DIR = os.path.join(BASE, "sop-templates")

# ---------------------------------------------------------------------------
# Colors
# ---------------------------------------------------------------------------
NAVY = RGBColor(0x1A, 0x3A, 0x5C)
ACCENT = RGBColor(0x2C, 0x7B, 0xE5)
DARK_TEXT = RGBColor(0x33, 0x33, 0x33)
GRAY_TAG = RGBColor(0x66, 0x66, 0x66)
LIGHT_GRAY = RGBColor(0x99, 0x99, 0x99)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

FOOTER_TEXT = "Preparado por Pure Work AI | Confidencial"


# ===================================================================
# Helper functions -- matching existing PureWorkAI formatting exactly
# ===================================================================

def create_doc():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Emu(139700)  # 11pt
    style.font.color.rgb = DARK_TEXT

    h1 = doc.styles['Heading 1']
    h1.font.name = 'Calibri'
    h1.font.size = Emu(228600)  # 18pt
    h1.font.bold = True
    h1.font.color.rgb = NAVY
    h1.paragraph_format.space_before = Emu(304800)
    h1.paragraph_format.space_after = Emu(0)

    h2 = doc.styles['Heading 2']
    h2.font.name = 'Calibri'
    h2.font.size = Emu(177800)  # 14pt
    h2.font.bold = True
    h2.font.color.rgb = NAVY
    h2.paragraph_format.space_before = Emu(127000)
    h2.paragraph_format.space_after = Emu(0)

    h3 = doc.styles['Heading 3']
    h3.font.name = 'Calibri'
    h3.font.size = Emu(152400)  # 12pt
    h3.font.bold = True
    h3.font.color.rgb = NAVY
    h3.paragraph_format.space_before = Emu(127000)
    h3.paragraph_format.space_after = Emu(0)

    style_lb = doc.styles['List Bullet']
    style_lb.font.name = 'Calibri'
    style_lb.font.size = Emu(139700)
    style_lb.font.color.rgb = DARK_TEXT

    return doc


def add_title(doc, company_line="[NOMBRE DE LA EMPRESA]"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(company_line)
    run.font.name = 'Calibri'
    run.font.size = Emu(279400)  # 22pt
    run.font.bold = True
    run.font.color.rgb = NAVY


def add_jurisdiction_tag(doc, tag_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(tag_text)
    run.font.name = 'Calibri'
    run.font.size = Emu(139700)
    run.italic = True
    run.font.color.rgb = GRAY_TAG


def add_metadata_table(doc, policy_number):
    table = doc.add_table(rows=5, cols=2)
    data = [
        ("N\u00famero de Pol\u00edtica:", policy_number),
        ("Versi\u00f3n:", "1.0"),
        ("Fecha de Vigencia:", "[FECHA DE VIGENCIA]"),
        ("Fecha de Revisi\u00f3n:", "[FECHA DE REVISI\u00d3N]"),
        ("Responsable:", "Recursos Humanos"),
    ]
    for i, (label, value) in enumerate(data):
        for j, text in enumerate([label, value]):
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.name = 'Calibri'
            run.font.size = Emu(127000)  # 10pt


def H1(doc, text):
    doc.add_heading(text, level=1)

def H2(doc, text):
    doc.add_heading(text, level=2)

def H3(doc, text):
    doc.add_heading(text, level=3)

def P(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def P_bold(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    return p

def B(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def style_header_row(row, color_hex="1A3A5C"):
    for cell in row.cells:
        set_cell_shading(cell, color_hex)
        for pr in cell.paragraphs:
            for run in pr.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.name = "Calibri"

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '</w:tblBorders>')
    tblPr.append(borders)

def make_table(doc, headers, rows_data, bold_first_col=False):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        hdr.cells[i].text = h
    style_header_row(hdr)
    for row_data in rows_data:
        row = table.add_row()
        for i, text in enumerate(row_data):
            cell = row.cells[i]
            cell.text = ""
            pr = cell.paragraphs[0]
            run = pr.add_run(text)
            run.font.size = Pt(10)
            run.font.name = "Calibri"
            if bold_first_col and i == 0:
                run.font.bold = True
    doc.add_paragraph()
    return table

def add_footer_section(doc):
    """Add footer text to every section of the document."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(FOOTER_TEXT)
        run.font.size = Emu(101600)  # 8pt
        run.italic = True
        run.font.color.rgb = LIGHT_GRAY
        run.font.name = "Calibri"

def add_approval_table(doc):
    H2(doc, "Aprobaci\u00f3n")
    table = doc.add_table(rows=4, cols=3)
    headers = ["Cargo", "Nombre", "Firma / Fecha"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.font.name = 'Calibri'
        run.font.size = Emu(127000)
        run.bold = True
    roles = ["Presidente / Director General", "VP, Recursos Humanos", "Asesor Legal"]
    for r, role in enumerate(roles, 1):
        table.cell(r, 0).text = role
        table.cell(r, 1).text = "________________________"
        table.cell(r, 2).text = "________________________"

def add_acknowledgment(doc):
    H2(doc, "Acuse de Recibo del Empleado")
    P(doc, 'Yo, ________________________ (nombre en letra de molde), reconozco que he recibido, le\u00eddo y comprendido esta pol\u00edtica. Acepto cumplir con sus t\u00e9rminos y condiciones como requisito de mi empleo con [NOMBRE DE LA EMPRESA].')
    doc.add_paragraph()
    table = doc.add_table(rows=3, cols=2)
    data = [
        ("Firma del Empleado: ________________________", "Fecha: ________________________"),
        ("Nombre del Empleado: ________________________", "N\u00famero de Empleado: ________________________"),
        ("Departamento: ________________________", "Puesto: ________________________"),
    ]
    for i, (left, right) in enumerate(data):
        table.cell(i, 0).text = left
        table.cell(i, 1).text = right

def page_break(doc):
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(WD_BREAK.PAGE)


# ===================================================================
# 1. HARASSMENT POLICY -- Pol\u00edtica contra el Acoso Laboral
# ===================================================================
def build_harassment():
    doc = create_doc()
    add_title(doc)
    add_jurisdiction_tag(doc, "Idioma: Espa\u00f1ol | Aplicable: Estados Unidos y marcos universales")
    H1(doc, "Pol\u00edtica contra el Acoso Laboral")
    doc.add_paragraph()
    add_metadata_table(doc, "[POL-ES-001]")

    H2(doc, "1. Prop\u00f3sito")
    P(doc, "[NOMBRE DE LA EMPRESA] se compromete a mantener un ambiente de trabajo libre de acoso, intimidaci\u00f3n y discriminaci\u00f3n. Esta pol\u00edtica establece las normas de conducta, los mecanismos de denuncia y las consecuencias por violaciones, en cumplimiento con el T\u00edtulo VII de la Ley de Derechos Civiles de 1964, las directrices de la Comisi\u00f3n para la Igualdad de Oportunidades en el Empleo (EEOC) y los marcos universales de derechos humanos en el lugar de trabajo.")

    H2(doc, "2. Alcance")
    P(doc, "Esta pol\u00edtica aplica a todos los empleados, contratistas, pasantes, voluntarios, proveedores y visitantes de [NOMBRE DE LA EMPRESA] en todas las ubicaciones, incluyendo oficinas, sitios de trabajo remotos, viajes de negocios, eventos patrocinados por la empresa y comunicaciones electr\u00f3nicas relacionadas con el trabajo.")

    H2(doc, "3. Definiciones")
    H3(doc, "3.1 Acoso Laboral")
    P(doc, "El acoso laboral es cualquier conducta no deseada basada en raza, color, religi\u00f3n, sexo (incluyendo embarazo, identidad de g\u00e9nero y orientaci\u00f3n sexual), origen nacional, edad (40 a\u00f1os o m\u00e1s), discapacidad, informaci\u00f3n gen\u00e9tica u otra caracter\u00edstica protegida por la ley, que:")
    B(doc, "Tiene el prop\u00f3sito o efecto de interferir de manera irrazonable con el desempe\u00f1o laboral de una persona;")
    B(doc, "Crea un ambiente de trabajo intimidante, hostil u ofensivo;")
    B(doc, "Se utiliza como base para decisiones de empleo (contrataci\u00f3n, ascenso, terminaci\u00f3n).")

    H3(doc, "3.2 Acoso Sexual")
    P(doc, "El acoso sexual incluye avances sexuales no deseados, solicitudes de favores sexuales y cualquier otra conducta verbal, f\u00edsica o visual de naturaleza sexual que:")
    B(doc, "Se condiciona expl\u00edcita o impl\u00edcitamente al empleo de la persona (quid pro quo);")
    B(doc, "Interfiere de manera irrazonable con el desempe\u00f1o laboral;")
    B(doc, "Crea un ambiente de trabajo intimidante, hostil o degradante.")
    P(doc, "Ejemplos incluyen, sin limitarse a: comentarios sexuales o de doble sentido, contacto f\u00edsico no deseado, exhibici\u00f3n de material sexualmente expl\u00edcito, correos electr\u00f3nicos o mensajes de contenido sexual, miradas lascivas persistentes y presiones para obtener citas o favores sexuales.")

    H3(doc, "3.3 Tabla de Conductas Prohibidas")
    make_table(doc,
        ["Tipo de Conducta", "Ejemplos", "Gravedad"],
        [
            ["Acoso verbal", "Insultos, comentarios despectivos sobre raza, g\u00e9nero, religi\u00f3n o apariencia f\u00edsica; bromas ofensivas; apodos degradantes", "Moderada a grave"],
            ["Acoso f\u00edsico", "Contacto f\u00edsico no deseado, empujones, bloquear el paso, golpes, roces inapropiados", "Grave"],
            ["Acoso visual", "Exhibici\u00f3n de im\u00e1genes ofensivas, caricaturas, gestos obscenos, miradas intimidantes", "Moderada a grave"],
            ["Acoso sexual \u2014 Quid pro quo", "Condicionar ascensos, aumentos o permanencia en el empleo a favores sexuales", "Muy grave"],
            ["Acoso sexual \u2014 Ambiente hostil", "Comentarios sexuales repetidos, chistes de contenido sexual, distribuci\u00f3n de material pornogr\u00e1fico", "Grave"],
            ["Ciberacoso", "Mensajes amenazantes o degradantes por correo electr\u00f3nico, redes sociales o aplicaciones de mensajer\u00eda", "Moderada a grave"],
            ["Intimidaci\u00f3n (bullying)", "Humillaci\u00f3n p\u00fablica, exclusi\u00f3n deliberada, sabotaje del trabajo, amenazas", "Grave"],
            ["Represalia", "Cualquier acci\u00f3n adversa contra una persona por presentar una queja o participar en una investigaci\u00f3n", "Muy grave"],
        ],
        bold_first_col=True)

    H2(doc, "4. Canales de Denuncia")
    P(doc, "Cualquier persona que sea testigo o v\u00edctima de acoso debe reportarlo de inmediato a trav\u00e9s de cualquiera de los siguientes canales:")
    B(doc, "Supervisor inmediato o cualquier miembro de la gerencia;")
    B(doc, "Departamento de Recursos Humanos;")
    B(doc, "L\u00ednea de denuncia confidencial (si est\u00e1 disponible): [N\u00daMERO DE TEL\u00c9FONO];")
    B(doc, "Correo electr\u00f3nico designado: [CORREO DE DENUNCIAS];")
    B(doc, "Formulario de denuncia en l\u00ednea (si est\u00e1 disponible).")
    P(doc, "Las denuncias pueden presentarse de forma verbal o por escrito. Se aceptan denuncias an\u00f3nimas, aunque esto puede limitar la capacidad de la empresa para investigar a fondo.")

    H2(doc, "5. Procedimiento de Investigaci\u00f3n")
    P(doc, "Todas las denuncias ser\u00e1n tratadas con seriedad, confidencialidad y diligencia. El procedimiento de investigaci\u00f3n incluye:")
    B(doc, "Acuse de recibo de la denuncia dentro de las 48 horas h\u00e1biles siguientes a su recepci\u00f3n.", bold_prefix="Paso 1: ")
    B(doc, "Asignaci\u00f3n de un investigador imparcial (interno o externo) que no tenga conflicto de inter\u00e9s con las partes involucradas.", bold_prefix="Paso 2: ")
    B(doc, "Entrevistas con el denunciante, el presunto acosador y los testigos relevantes.", bold_prefix="Paso 3: ")
    B(doc, "Revisi\u00f3n de evidencia documental, electr\u00f3nica y f\u00edsica pertinente.", bold_prefix="Paso 4: ")
    B(doc, "Elaboraci\u00f3n de un informe de hallazgos y recomendaciones.", bold_prefix="Paso 5: ")
    B(doc, "Comunicaci\u00f3n de los resultados a las partes involucradas, respetando la confidencialidad.", bold_prefix="Paso 6: ")
    P(doc, "La investigaci\u00f3n se completar\u00e1 en un plazo razonable, generalmente no mayor a 30 d\u00edas calendario desde la recepci\u00f3n de la denuncia. Se implementar\u00e1n medidas provisionales cuando sea necesario para proteger al denunciante durante el proceso.")

    H2(doc, "6. Acciones Correctivas")
    P(doc, "Si la investigaci\u00f3n concluye que se ha producido acoso, [NOMBRE DE LA EMPRESA] tomar\u00e1 las acciones correctivas apropiadas, que pueden incluir:")
    B(doc, "Amonestaci\u00f3n verbal o escrita;")
    B(doc, "Capacitaci\u00f3n obligatoria sobre prevenci\u00f3n de acoso;")
    B(doc, "Reasignaci\u00f3n de funciones o traslado;")
    B(doc, "Suspensi\u00f3n sin goce de sueldo;")
    B(doc, "Terminaci\u00f3n del empleo;")
    B(doc, "Referencia a las autoridades competentes si la conducta constituye un delito.")
    P(doc, "La severidad de la acci\u00f3n correctiva depender\u00e1 de la gravedad de la conducta, la frecuencia del comportamiento, el historial disciplinario del infractor y el impacto sobre la v\u00edctima y el ambiente de trabajo.")

    H2(doc, "7. Prohibici\u00f3n de Represalias")
    P(doc, "[NOMBRE DE LA EMPRESA] proh\u00edbe estrictamente cualquier forma de represalia contra cualquier persona que, de buena fe, presente una denuncia de acoso, participe en una investigaci\u00f3n o se oponga a conductas de acoso. Las represalias incluyen, pero no se limitan a: terminaci\u00f3n, descenso de puesto, reducci\u00f3n de salario, cambio desfavorable de horario, exclusi\u00f3n de proyectos, hostigamiento y cualquier otra acci\u00f3n adversa.")
    P(doc, "Cualquier acto de represalia ser\u00e1 tratado como una violaci\u00f3n independiente de esta pol\u00edtica y estar\u00e1 sujeto a medidas disciplinarias, incluyendo la terminaci\u00f3n del empleo.")

    H2(doc, "8. Obligaciones de Capacitaci\u00f3n")
    P(doc, "[NOMBRE DE LA EMPRESA] proporcionar\u00e1 capacitaci\u00f3n sobre prevenci\u00f3n de acoso de la siguiente manera:")
    B(doc, "Todos los empleados nuevos recibir\u00e1n capacitaci\u00f3n durante el proceso de incorporaci\u00f3n.", bold_prefix="Incorporaci\u00f3n: ")
    B(doc, "Todos los empleados completar\u00e1n un curso de actualizaci\u00f3n anual sobre prevenci\u00f3n de acoso.", bold_prefix="Anual: ")
    B(doc, "Los supervisores y gerentes recibir\u00e1n capacitaci\u00f3n adicional sobre su responsabilidad de prevenir, identificar y responder a situaciones de acoso.", bold_prefix="Gerentes: ")
    P(doc, "La capacitaci\u00f3n cubrir\u00e1 los tipos de acoso, ejemplos de conducta prohibida, el proceso de denuncia, las responsabilidades de los supervisores y las consecuencias por incumplimiento.")

    H2(doc, "9. Marco Legal de Referencia")
    B(doc, "T\u00edtulo VII de la Ley de Derechos Civiles de 1964 (Estados Unidos)")
    B(doc, "Directrices de la EEOC sobre acoso sexual")
    B(doc, "Leyes estatales y locales aplicables contra el acoso y la discriminaci\u00f3n")
    B(doc, "Convenio 190 de la OIT sobre la Violencia y el Acoso en el Mundo del Trabajo")
    B(doc, "Declaraci\u00f3n Universal de los Derechos Humanos, Art\u00edculos 1 y 23")

    add_acknowledgment(doc)
    add_approval_table(doc)
    add_footer_section(doc)

    path = os.path.join(POLICY_DIR, "Policy_Harassment_Spanish.docx")
    doc.save(path)
    print(f"  [OK] {path}")
    return path


# ===================================================================
# 2. HEALTH & SAFETY -- Pol\u00edtica de Salud y Seguridad en el Trabajo
# ===================================================================
def build_health_safety():
    doc = create_doc()
    add_title(doc)
    add_jurisdiction_tag(doc, "Idioma: Espa\u00f1ol | Referencia: OSHA / Marcos universales")
    H1(doc, "Pol\u00edtica de Salud y Seguridad en el Trabajo")
    doc.add_paragraph()
    add_metadata_table(doc, "[POL-ES-002]")

    H2(doc, "1. Prop\u00f3sito")
    P(doc, "Esta pol\u00edtica establece el compromiso de [NOMBRE DE LA EMPRESA] con la protecci\u00f3n de la salud, la seguridad y el bienestar de todos los empleados, contratistas y visitantes en el lugar de trabajo. La empresa se compromete a cumplir con las normas de la Administraci\u00f3n de Seguridad y Salud Ocupacional (OSHA), las leyes estatales y locales aplicables y las mejores pr\u00e1cticas internacionales de seguridad laboral.")

    H2(doc, "2. Alcance")
    P(doc, "Esta pol\u00edtica aplica a todos los empleados, contratistas, subcontratistas, trabajadores temporales y visitantes en todas las instalaciones y sitios de trabajo operados o controlados por [NOMBRE DE LA EMPRESA].")

    H2(doc, "3. Responsabilidades del Empleador")
    B(doc, "Proporcionar un lugar de trabajo libre de peligros reconocidos que causen o puedan causar muerte o da\u00f1o f\u00edsico grave;")
    B(doc, "Cumplir con todas las normas, regulaciones y \u00f3rdenes emitidas por OSHA y las autoridades locales aplicables;")
    B(doc, "Establecer y mantener un programa de seguridad y salud ocupacional;")
    B(doc, "Proporcionar el equipo de protecci\u00f3n personal (EPP) necesario sin costo para el empleado;")
    B(doc, "Capacitar a los empleados sobre los peligros presentes en el lugar de trabajo y los procedimientos de seguridad;")
    B(doc, "Investigar todos los incidentes y cuasi-accidentes para prevenir recurrencias;")
    B(doc, "Mantener registros de lesiones y enfermedades ocupacionales conforme a los requisitos de OSHA (Formularios 300, 300A, 301);")
    B(doc, "Publicar el cartel de OSHA \u201cEs la Ley\u201d en un lugar visible del centro de trabajo.")

    H2(doc, "4. Responsabilidades del Empleado")
    B(doc, "Cumplir con todas las normas, regulaciones y procedimientos de seguridad establecidos por la empresa;")
    B(doc, "Utilizar correctamente el equipo de protecci\u00f3n personal proporcionado;")
    B(doc, "Reportar inmediatamente cualquier condici\u00f3n insegura, peligro, lesi\u00f3n o enfermedad laboral a su supervisor o al departamento de seguridad;")
    B(doc, "Participar en las capacitaciones de seguridad requeridas;")
    B(doc, "No operar equipos o maquinaria para los cuales no hayan sido capacitados y autorizados;")
    B(doc, "No desactivar ni manipular los dispositivos de seguridad de equipos o maquinaria;")
    B(doc, "Cooperar con las investigaciones de incidentes.")

    H2(doc, "5. Reporte de Peligros e Incidentes")
    P(doc, "Todos los empleados tienen la obligaci\u00f3n de reportar de inmediato:")
    B(doc, "Condiciones de trabajo inseguras (derrames, equipos da\u00f1ados, iluminaci\u00f3n inadecuada, etc.);")
    B(doc, "Actos inseguros observados;")
    B(doc, "Lesiones o enfermedades relacionadas con el trabajo, sin importar la gravedad;")
    B(doc, "Cuasi-accidentes (incidentes que pudieron haber resultado en lesi\u00f3n o da\u00f1o).")
    P(doc, "Los reportes deben realizarse al supervisor inmediato y/o al departamento de seguridad mediante el formulario de reporte de incidentes de la empresa. Ning\u00fan empleado ser\u00e1 sancionado por reportar de buena fe una condici\u00f3n o acto inseguro.")

    H2(doc, "6. Investigaci\u00f3n de Incidentes")
    P(doc, "Todos los accidentes, lesiones y cuasi-accidentes ser\u00e1n investigados para determinar las causas ra\u00edz e implementar medidas correctivas. La investigaci\u00f3n incluir\u00e1:")
    B(doc, "Aseguramiento inmediato del \u00e1rea del incidente;")
    B(doc, "Entrevistas con testigos y personas involucradas;")
    B(doc, "Documentaci\u00f3n fotogr\u00e1fica y recolecci\u00f3n de evidencia;")
    B(doc, "An\u00e1lisis de causas ra\u00edz (utilizando metodolog\u00edas como los 5 Porqu\u00e9s o el Diagrama de Ishikawa);")
    B(doc, "Plan de acciones correctivas y preventivas con responsables y fechas l\u00edmite;")
    B(doc, "Seguimiento para verificar la efectividad de las medidas implementadas.")

    H2(doc, "7. Requisitos de Equipo de Protecci\u00f3n Personal (EPP)")
    P(doc, "[NOMBRE DE LA EMPRESA] proporcionar\u00e1 EPP apropiado basado en la evaluaci\u00f3n de peligros del lugar de trabajo. Los empleados deben usar el EPP designado en todo momento cuando se encuentren en \u00e1reas donde sea requerido. El EPP incluye, pero no se limita a: cascos, gafas de seguridad, protectores auditivos, guantes, calzado de seguridad, arneses de seguridad y protecci\u00f3n respiratoria.")

    H2(doc, "8. Tabla de Peligros Comunes y Controles")
    make_table(doc,
        ["Peligro", "Ejemplos", "Controles Recomendados"],
        [
            ["Ca\u00eddas", "Superficies resbaladizas, trabajo en alturas, escaleras defectuosas", "Barandales, arn\u00e9s, calzado antideslizante, se\u00f1alizaci\u00f3n"],
            ["Golpes por objetos", "Herramientas que caen, materiales apilados, maquinaria m\u00f3vil", "Casco, zonas de exclusi\u00f3n, almacenamiento seguro"],
            ["Riesgos el\u00e9ctricos", "Cableado expuesto, equipos defectuosos, contacto con l\u00edneas el\u00e9ctricas", "Bloqueo/etiquetado (LOTO), EPP diel\u00e9ctrico, inspecci\u00f3n peri\u00f3dica"],
            ["Exposici\u00f3n qu\u00edmica", "Solventes, \u00e1cidos, gases, polvos", "Ventilaci\u00f3n, EPP, Hojas de Datos de Seguridad (SDS), capacitaci\u00f3n"],
            ["Ruido excesivo", "Maquinaria industrial, herramientas neum\u00e1ticas", "Protecci\u00f3n auditiva, controles de ingenier\u00eda, rotaci\u00f3n de puestos"],
            ["Riesgos ergon\u00f3micos", "Levantamiento repetitivo, posturas forzadas, uso prolongado de pantalla", "Capacitaci\u00f3n ergon\u00f3mica, ajuste de estaciones, pausas activas"],
            ["Temperaturas extremas", "Trabajo al aire libre, c\u00e1maras fr\u00edas, ambientes calurosos", "Hidrataci\u00f3n, descansos, ropa protectora, monitoreo de s\u00edntomas"],
            ["Riesgos biol\u00f3gicos", "Sangre, moho, bacterias, virus", "EPP, protocolo de higiene, vacunaci\u00f3n, eliminaci\u00f3n segura de residuos"],
        ],
        bold_first_col=True)

    H2(doc, "9. Procedimientos de Emergencia")
    P(doc, "Todos los empleados deben familiarizarse con los procedimientos de emergencia de su lugar de trabajo, incluyendo:")
    B(doc, "Ubicaci\u00f3n de salidas de emergencia y puntos de reuni\u00f3n;")
    B(doc, "Procedimiento de evacuaci\u00f3n y rutas de escape;")
    B(doc, "Ubicaci\u00f3n y uso de extintores de incendios;")
    B(doc, "Ubicaci\u00f3n de botiquines de primeros auxilios y equipos de emergencia;")
    B(doc, "N\u00fameros de emergencia (911, brigada de emergencia interna, n\u00famero de la planta);")
    B(doc, "Procedimiento para reportar emergencias.")
    P(doc, "Se realizar\u00e1n simulacros de evacuaci\u00f3n al menos dos veces al a\u00f1o. Todos los empleados deben participar en los simulacros.")

    H2(doc, "10. Primeros Auxilios")
    P(doc, "[NOMBRE DE LA EMPRESA] mantendr\u00e1 botiquines de primeros auxilios en ubicaciones accesibles y debidamente se\u00f1alizadas. Se designar\u00e1n empleados capacitados en primeros auxilios y RCP en cada turno de trabajo. La lista de socorristas y la ubicaci\u00f3n de los botiquines se publicar\u00e1n en los tableros de seguridad.")

    H2(doc, "11. Derecho a Rechazar Trabajo Inseguro")
    P(doc, "Todo empleado tiene el derecho de rechazar una tarea que considere razonablemente que representa un peligro inminente para su vida o salud, o la de sus compa\u00f1eros de trabajo. El empleado debe notificar de inmediato a su supervisor sobre la situaci\u00f3n. Ning\u00fan empleado ser\u00e1 sancionado por ejercer de buena fe este derecho. El supervisor investigar\u00e1 la situaci\u00f3n antes de requerir la reanudaci\u00f3n de la tarea.")

    add_acknowledgment(doc)
    add_approval_table(doc)
    add_footer_section(doc)

    path = os.path.join(POLICY_DIR, "Policy_HealthSafety_Spanish.docx")
    doc.save(path)
    print(f"  [OK] {path}")
    return path


# ===================================================================
# 3. DRUG & ALCOHOL -- Pol\u00edtica de Drogas y Alcohol
# ===================================================================
def build_drug_alcohol():
    doc = create_doc()
    add_title(doc)
    add_jurisdiction_tag(doc, "Idioma: Espa\u00f1ol | Referencia: DOT / DFWA / Marcos universales")
    H1(doc, "Pol\u00edtica de Drogas y Alcohol")
    doc.add_paragraph()
    add_metadata_table(doc, "[POL-ES-003]")

    H2(doc, "1. Prop\u00f3sito")
    P(doc, "[NOMBRE DE LA EMPRESA] se compromete a mantener un lugar de trabajo libre de drogas y alcohol para proteger la salud y seguridad de todos los empleados, clientes y el p\u00fablico. Esta pol\u00edtica cumple con la Ley de Lugar de Trabajo Libre de Drogas (Drug-Free Workplace Act) y las regulaciones del Departamento de Transporte (DOT) cuando sean aplicables.")

    H2(doc, "2. Alcance")
    P(doc, "Esta pol\u00edtica aplica a todos los empleados, contratistas y trabajadores temporales de [NOMBRE DE LA EMPRESA], en todas las ubicaciones y durante el horario de trabajo, incluyendo viajes de negocios y eventos patrocinados por la empresa.")

    H2(doc, "3. Sustancias Prohibidas")
    P(doc, "Se proh\u00edbe estrictamente:")
    B(doc, "El uso, posesi\u00f3n, distribuci\u00f3n, venta o fabricaci\u00f3n de drogas ilegales o sustancias controladas en las instalaciones de la empresa o durante el horario laboral;")
    B(doc, "Presentarse a trabajar bajo la influencia de drogas ilegales, sustancias controladas no prescritas o alcohol;")
    B(doc, "El uso de medicamentos con receta m\u00e9dica o de venta libre que afecten la capacidad del empleado para desempe\u00f1ar sus funciones de manera segura, sin haber notificado previamente a Recursos Humanos;")
    B(doc, "El consumo de bebidas alcoh\u00f3licas durante el horario de trabajo o en las instalaciones de la empresa, salvo en eventos expresamente autorizados por la gerencia;")
    B(doc, "El uso indebido de sustancias legales (como inhalantes o sustancias de venta libre) con fines de intoxicaci\u00f3n.")

    H2(doc, "4. Procedimientos de Pruebas")
    H3(doc, "4.1 Prueba Previa al Empleo")
    P(doc, "Todos los candidatos seleccionados para un puesto pueden estar sujetos a una prueba de detecci\u00f3n de drogas como condici\u00f3n para la contrataci\u00f3n. Un resultado positivo o la negativa a someterse a la prueba resultar\u00e1 en la rescisi\u00f3n de la oferta de empleo.")

    H3(doc, "4.2 Prueba por Sospecha Razonable")
    P(doc, "Se podr\u00e1 solicitar una prueba de detecci\u00f3n de drogas o alcohol cuando un supervisor capacitado observe comportamientos, apariencia, lenguaje corporal o desempe\u00f1o que indiquen razonablemente el uso de sustancias prohibidas. La sospecha razonable debe estar documentada por al menos un supervisor capacitado.")

    H3(doc, "4.3 Prueba Post-Accidente")
    P(doc, "Se realizar\u00e1n pruebas de detecci\u00f3n de drogas y alcohol a cualquier empleado involucrado en un accidente de trabajo que resulte en lesi\u00f3n que requiera atenci\u00f3n m\u00e9dica, da\u00f1o significativo a la propiedad o una violaci\u00f3n de seguridad. La prueba debe realizarse lo antes posible despu\u00e9s del incidente, idealmente dentro de las primeras 2 horas para alcohol y 32 horas para drogas.")

    H3(doc, "4.4 Pruebas Aleatorias")
    P(doc, "Los empleados en posiciones sensibles de seguridad o sujetos a regulaciones del DOT podr\u00e1n ser seleccionados aleatoriamente para pruebas de detecci\u00f3n de drogas y alcohol, de acuerdo con las regulaciones federales aplicables.")

    H3(doc, "4.5 Prueba de Regreso al Servicio")
    P(doc, "Cualquier empleado que regrese al trabajo despu\u00e9s de completar un programa de tratamiento por abuso de sustancias deber\u00e1 someterse a una prueba de detecci\u00f3n negativa antes de reincorporarse y estar\u00e1 sujeto a pruebas de seguimiento durante un per\u00edodo m\u00ednimo de 12 meses.")

    H2(doc, "5. Consecuencias por Violaciones")
    make_table(doc,
        ["Infracci\u00f3n", "Primera Ocurrencia", "Segunda Ocurrencia", "Tercera Ocurrencia"],
        [
            ["Resultado positivo en prueba", "Suspensi\u00f3n; referencia obligatoria al Programa de Asistencia al Empleado (EAP); prueba de regreso al servicio", "Terminaci\u00f3n del empleo", "N/A"],
            ["Posesi\u00f3n de sustancias prohibidas", "Suspensi\u00f3n sin goce de sueldo; referencia al EAP; investigaci\u00f3n interna", "Terminaci\u00f3n del empleo", "N/A"],
            ["Distribuci\u00f3n o venta de drogas", "Terminaci\u00f3n inmediata del empleo; referencia a las autoridades", "N/A", "N/A"],
            ["Negativa a someterse a prueba", "Tratada como resultado positivo; suspensi\u00f3n inmediata", "Terminaci\u00f3n del empleo", "N/A"],
            ["Presentarse bajo la influencia", "Retiro inmediato del lugar de trabajo; suspensi\u00f3n; referencia al EAP", "Terminaci\u00f3n del empleo", "N/A"],
            ["Adulteraci\u00f3n de muestra", "Terminaci\u00f3n inmediata del empleo", "N/A", "N/A"],
        ],
        bold_first_col=True)

    H2(doc, "6. Programa de Asistencia al Empleado (EAP)")
    P(doc, "[NOMBRE DE LA EMPRESA] ofrece un Programa de Asistencia al Empleado (EAP) confidencial que brinda orientaci\u00f3n, evaluaci\u00f3n y referencia a tratamiento para empleados con problemas de abuso de sustancias. Se alienta a los empleados a buscar ayuda voluntariamente antes de que su desempe\u00f1o laboral se vea afectado.")
    B(doc, "El uso voluntario del EAP no ser\u00e1 motivo de medidas disciplinarias;")
    B(doc, "La solicitud de ayuda despu\u00e9s de una violaci\u00f3n de esta pol\u00edtica no exime al empleado de las consecuencias disciplinarias;")
    B(doc, "Toda la informaci\u00f3n relacionada con el EAP ser\u00e1 tratada de manera estrictamente confidencial.")

    H2(doc, "7. Cumplimiento con Regulaciones del DOT")
    P(doc, "Los empleados que ocupen posiciones sensibles de seguridad reguladas por el Departamento de Transporte (DOT) estar\u00e1n sujetos a requisitos adicionales de pruebas y cumplimiento conforme a 49 CFR Parte 40 y las regulaciones espec\u00edficas de la agencia modal aplicable (FMCSA, FTA, FAA, FRA, PHMSA). Estos requisitos incluyen pruebas previas al empleo, aleatorias, post-accidente, de sospecha razonable, de regreso al servicio y de seguimiento.")

    add_acknowledgment(doc)
    add_approval_table(doc)
    add_footer_section(doc)

    path = os.path.join(POLICY_DIR, "Policy_DrugAlcohol_Spanish.docx")
    doc.save(path)
    print(f"  [OK] {path}")
    return path


# ===================================================================
# 4. ATTENDANCE -- Pol\u00edtica de Asistencia y Puntualidad
# ===================================================================
def build_attendance():
    doc = create_doc()
    add_title(doc)
    add_jurisdiction_tag(doc, "Idioma: Espa\u00f1ol | Aplicaci\u00f3n universal")
    H1(doc, "Pol\u00edtica de Asistencia y Puntualidad")
    doc.add_paragraph()
    add_metadata_table(doc, "[POL-ES-004]")

    H2(doc, "1. Prop\u00f3sito")
    P(doc, "La asistencia regular y puntual es esencial para el funcionamiento eficiente de [NOMBRE DE LA EMPRESA] y para brindar un servicio de calidad a nuestros clientes. Esta pol\u00edtica establece las expectativas de la empresa en relaci\u00f3n con la asistencia y la puntualidad, el procedimiento para reportar ausencias y las consecuencias por incumplimiento.")

    H2(doc, "2. Alcance")
    P(doc, "Esta pol\u00edtica aplica a todos los empleados de [NOMBRE DE LA EMPRESA], independientemente de su clasificaci\u00f3n (tiempo completo, tiempo parcial, temporal). Los supervisores son responsables de administrar esta pol\u00edtica de manera justa y consistente.")

    H2(doc, "3. Expectativas de Asistencia")
    P(doc, "Se espera que todos los empleados:")
    B(doc, "Se presenten a trabajar seg\u00fan el horario establecido y est\u00e9n listos para iniciar sus funciones a la hora de inicio de su turno;")
    B(doc, "Permanezcan en su puesto de trabajo durante todo su turno, excepto durante los per\u00edodos de descanso y almuerzo autorizados;")
    B(doc, "Notifiquen a su supervisor con la mayor anticipaci\u00f3n posible cuando no puedan presentarse a trabajar o cuando vayan a llegar tarde;")
    B(doc, "Registren su asistencia de acuerdo con el sistema establecido (reloj checador, registro electr\u00f3nico u otro sistema designado).")

    H2(doc, "4. Procedimiento para Reportar Ausencias")
    P(doc, "Cuando un empleado no pueda presentarse a trabajar, debe:")
    B(doc, "Notificar a su supervisor directo al menos [1 HORA] antes del inicio de su turno, a menos que circunstancias de emergencia lo impidan;")
    B(doc, "Comunicarse personalmente (por tel\u00e9fono o mediante el sistema designado); los mensajes de texto o correos electr\u00f3nicos solo son aceptables si el supervisor ha autorizado previamente estos medios;")
    B(doc, "Indicar el motivo de la ausencia y la duraci\u00f3n esperada;")
    B(doc, "En caso de ausencias de m\u00e1s de un d\u00eda, comunicarse diariamente con su supervisor, a menos que se haya acordado lo contrario;")
    B(doc, "Proporcionar documentaci\u00f3n de respaldo cuando sea requerida (nota m\u00e9dica, citaci\u00f3n judicial, etc.).")

    H2(doc, "5. Ausencias Justificadas vs. Injustificadas")
    H3(doc, "5.1 Ausencias Justificadas")
    P(doc, "Se consideran ausencias justificadas aquellas que est\u00e1n amparadas por la ley o la pol\u00edtica de la empresa, incluyendo:")
    B(doc, "Enfermedad o lesi\u00f3n del empleado (con documentaci\u00f3n m\u00e9dica si se requiere);")
    B(doc, "Licencia m\u00e9dica y familiar (FMLA, cuando aplique);")
    B(doc, "D\u00edas de duelo por fallecimiento de un familiar directo;")
    B(doc, "Servicio de jurado o comparecencia ante un tribunal;")
    B(doc, "Servicio militar activo o entrenamiento;")
    B(doc, "D\u00edas festivos oficiales de la empresa;")
    B(doc, "Vacaciones previamente aprobadas;")
    B(doc, "Emergencias debidamente documentadas.")

    H3(doc, "5.2 Ausencias Injustificadas")
    P(doc, "Se consideran ausencias injustificadas aquellas que no cuentan con la autorizaci\u00f3n previa del supervisor, no est\u00e1n amparadas por la ley o la pol\u00edtica de la empresa, o para las cuales no se proporcion\u00f3 la notificaci\u00f3n o documentaci\u00f3n requerida.")

    H2(doc, "6. Tardanzas")
    P(doc, "Se considera tardanza cuando un empleado llega a su puesto de trabajo despu\u00e9s de la hora de inicio programada. Las tardanzas frecuentes afectan la productividad del equipo y ser\u00e1n abordadas de acuerdo con el proceso de disciplina progresiva.")
    B(doc, "Se registrar\u00e1 como tardanza cualquier llegada despu\u00e9s de la hora de inicio del turno;")
    B(doc, "Tres (3) tardanzas en un per\u00edodo de 30 d\u00edas ser\u00e1n equivalentes a una (1) ausencia injustificada;")
    B(doc, "Salir antes de la hora de finalizaci\u00f3n del turno sin autorizaci\u00f3n tambi\u00e9n se considerar\u00e1 como tardanza.")

    H2(doc, "7. Ausencia sin Aviso (No-Call / No-Show)")
    P(doc, "La ausencia sin aviso ocurre cuando un empleado no se presenta a trabajar ni se comunica con su supervisor. Esta es una violaci\u00f3n grave de esta pol\u00edtica.")
    B(doc, "Una ausencia sin aviso generar\u00e1 una acci\u00f3n disciplinaria inmediata;")
    B(doc, "Tres (3) d\u00edas consecutivos de ausencia sin aviso se considerar\u00e1n como abandono voluntario del empleo y resultar\u00e1n en la terminaci\u00f3n autom\u00e1tica de la relaci\u00f3n laboral.")

    H2(doc, "8. Requisitos de Documentaci\u00f3n")
    B(doc, "Se podr\u00e1 solicitar una nota m\u00e9dica para ausencias por enfermedad de tres (3) o m\u00e1s d\u00edas consecutivos;")
    B(doc, "Se podr\u00e1 requerir documentaci\u00f3n para cualquier ausencia que siga un patr\u00f3n sospechoso (antes/despu\u00e9s de fines de semana o d\u00edas festivos);")
    B(doc, "La empresa se reserva el derecho de solicitar documentaci\u00f3n en cualquier momento.")

    H2(doc, "9. Disciplina Progresiva por Violaciones de Asistencia")
    make_table(doc,
        ["Etapa", "Acci\u00f3n Disciplinaria", "Criterio"],
        [
            ["1", "Conversaci\u00f3n verbal documentada", "2 ausencias injustificadas o equivalente en tardanzas en un per\u00edodo de 90 d\u00edas"],
            ["2", "Amonestaci\u00f3n escrita", "3 ausencias injustificadas o equivalente en tardanzas en un per\u00edodo de 90 d\u00edas"],
            ["3", "Amonestaci\u00f3n escrita final / Suspensi\u00f3n", "4 ausencias injustificadas o equivalente en tardanzas en un per\u00edodo de 90 d\u00edas"],
            ["4", "Terminaci\u00f3n del empleo", "5 o m\u00e1s ausencias injustificadas o equivalente en tardanzas en un per\u00edodo de 90 d\u00edas, o cualquier ausencia sin aviso (no-call/no-show) reiterada"],
        ],
        bold_first_col=True)
    P(doc, "Nota: La empresa se reserva el derecho de omitir etapas de disciplina progresiva en casos de violaciones graves, incluyendo ausencias sin aviso.")

    add_acknowledgment(doc)
    add_approval_table(doc)
    add_footer_section(doc)

    path = os.path.join(POLICY_DIR, "Policy_Attendance_Spanish.docx")
    doc.save(path)
    print(f"  [OK] {path}")
    return path


# ===================================================================
# 5. ONBOARDING -- Gu\u00eda de Incorporaci\u00f3n para Nuevos Empleados
# ===================================================================
def build_onboarding():
    doc = create_doc()
    add_title(doc)
    add_jurisdiction_tag(doc, "Idioma: Espa\u00f1ol | Gu\u00eda de incorporaci\u00f3n universal")
    H1(doc, "Gu\u00eda de Incorporaci\u00f3n para Nuevos Empleados")
    doc.add_paragraph()
    add_metadata_table(doc, "[ONB-ES-001]")

    H2(doc, "1. Prop\u00f3sito")
    P(doc, "Esta gu\u00eda establece el proceso de incorporaci\u00f3n (onboarding) para nuevos empleados de [NOMBRE DE LA EMPRESA]. Un programa de incorporaci\u00f3n estructurado asegura que los nuevos integrantes del equipo se sientan bienvenidos, comprendan la cultura y los valores de la empresa, y cuenten con las herramientas y el conocimiento necesarios para tener \u00e9xito en su puesto.")

    H2(doc, "2. D\u00eda 1 \u2014 Bienvenida e Integraci\u00f3n Inicial")
    H3(doc, "2.1 Bienvenida")
    B(doc, "Recepci\u00f3n por el supervisor directo o representante de Recursos Humanos;")
    B(doc, "Entrega de paquete de bienvenida (identificaci\u00f3n, materiales de la empresa, gu\u00eda del empleado);")
    B(doc, "Recorrido por las instalaciones: oficinas, \u00e1reas comunes, ba\u00f1os, comedor, salidas de emergencia, estacionamiento;")
    B(doc, "Presentaci\u00f3n al equipo de trabajo inmediato;")
    B(doc, "Asignaci\u00f3n de un compa\u00f1ero mentor (buddy) para las primeras semanas.")

    H3(doc, "2.2 Formularios de Recursos Humanos")
    B(doc, "Formulario W-4 (Certificado de Retenci\u00f3n de Impuestos Federales);", bold_prefix="W-4: ")
    B(doc, "Formulario I-9 (Verificaci\u00f3n de Elegibilidad de Empleo) \u2014 debe completarse el primer d\u00eda de trabajo; los documentos de identidad deben presentarse dentro de los primeros 3 d\u00edas h\u00e1biles;", bold_prefix="I-9: ")
    B(doc, "Formulario de dep\u00f3sito directo con informaci\u00f3n bancaria;", bold_prefix="Dep\u00f3sito Directo: ")
    B(doc, "Formulario de contacto de emergencia;")
    B(doc, "Inscripci\u00f3n en beneficios (seguro m\u00e9dico, dental, visi\u00f3n, plan de retiro);")
    B(doc, "Firma de acuse de recibo del manual del empleado;")
    B(doc, "Acuerdo de confidencialidad y propiedad intelectual (si aplica).")

    H3(doc, "2.3 Tecnolog\u00eda y Accesos")
    B(doc, "Configuraci\u00f3n de cuenta de correo electr\u00f3nico corporativo;")
    B(doc, "Acceso a sistemas y plataformas de trabajo (intranet, herramientas de colaboraci\u00f3n);")
    B(doc, "Entrega de equipo de c\u00f3mputo y perif\u00e9ricos;")
    B(doc, "Entrega de tarjeta de acceso / credencial de identificaci\u00f3n.")

    H2(doc, "3. Semana 1 \u2014 Orientaci\u00f3n y Capacitaci\u00f3n")
    B(doc, "Orientaci\u00f3n de seguridad laboral (salidas de emergencia, protocolos, EPP si aplica);")
    B(doc, "Capacitaci\u00f3n sobre pol\u00edticas clave: acoso, drogas y alcohol, asistencia, c\u00f3digo de conducta;")
    B(doc, "Introducci\u00f3n a la misi\u00f3n, visi\u00f3n y valores de la empresa;")
    B(doc, "Revisi\u00f3n de la descripci\u00f3n del puesto y expectativas de desempe\u00f1o;")
    B(doc, "Sesi\u00f3n con el supervisor para establecer objetivos iniciales;")
    B(doc, "Presentaciones con departamentos clave con los que interactuar\u00e1 el nuevo empleado;")
    B(doc, "Capacitaci\u00f3n en herramientas y sistemas espec\u00edficos del puesto.")

    H2(doc, "4. Hitos de 30 / 60 / 90 D\u00edas")
    make_table(doc,
        ["Per\u00edodo", "Actividades Clave", "Entregable / Evaluaci\u00f3n"],
        [
            ["30 d\u00edas", "Completar todas las capacitaciones obligatorias; Dominar las tareas b\u00e1sicas del puesto; Establecer relaciones con el equipo; Reuni\u00f3n de retroalimentaci\u00f3n con supervisor", "Formulario de evaluaci\u00f3n de 30 d\u00edas completado por el supervisor"],
            ["60 d\u00edas", "Demostrar competencia en funciones principales; Participar activamente en reuniones de equipo; Identificar \u00e1reas de mejora y oportunidades; Segunda reuni\u00f3n de retroalimentaci\u00f3n", "Formulario de evaluaci\u00f3n de 60 d\u00edas; Plan de desarrollo si es necesario"],
            ["90 d\u00edas", "Desempe\u00f1ar funciones de manera aut\u00f3noma; Contribuir a proyectos del equipo; Completar la evaluaci\u00f3n del per\u00edodo de prueba; Establecer objetivos de desempe\u00f1o para el pr\u00f3ximo trimestre", "Evaluaci\u00f3n final del per\u00edodo de prueba; Confirmaci\u00f3n de empleo o plan de mejora"],
        ],
        bold_first_col=True)

    H2(doc, "5. Lista de Verificaci\u00f3n de Incorporaci\u00f3n")
    make_table(doc,
        ["Tarea", "Responsable", "Fecha L\u00edmite", "Completado"],
        [
            ["Preparar estaci\u00f3n de trabajo y equipo", "TI / Supervisor", "Antes del D\u00eda 1", "\u2610"],
            ["Enviar correo de bienvenida al nuevo empleado", "Recursos Humanos", "Antes del D\u00eda 1", "\u2610"],
            ["Completar formularios W-4, I-9, dep\u00f3sito directo", "Nuevo Empleado / RRHH", "D\u00eda 1", "\u2610"],
            ["Recorrido por las instalaciones", "Supervisor / Mentor", "D\u00eda 1", "\u2610"],
            ["Presentaci\u00f3n al equipo", "Supervisor", "D\u00eda 1", "\u2610"],
            ["Orientaci\u00f3n de seguridad", "Seguridad / Supervisor", "Semana 1", "\u2610"],
            ["Capacitaci\u00f3n en pol\u00edticas clave", "Recursos Humanos", "Semana 1", "\u2610"],
            ["Revisi\u00f3n de descripci\u00f3n de puesto y objetivos", "Supervisor", "Semana 1", "\u2610"],
            ["Inscripci\u00f3n en beneficios", "Nuevo Empleado / RRHH", "Primeros 30 d\u00edas", "\u2610"],
            ["Evaluaci\u00f3n de 30 d\u00edas", "Supervisor", "D\u00eda 30", "\u2610"],
            ["Evaluaci\u00f3n de 60 d\u00edas", "Supervisor", "D\u00eda 60", "\u2610"],
            ["Evaluaci\u00f3n de 90 d\u00edas / Confirmaci\u00f3n", "Supervisor / RRHH", "D\u00eda 90", "\u2610"],
        ],
        bold_first_col=True)

    add_footer_section(doc)

    path = os.path.join(ONBOARD_DIR, "Onboarding_Spanish.docx")
    doc.save(path)
    print(f"  [OK] {path}")
    return path


# ===================================================================
# 6. SAFETY ORIENTATION TRAINING -- Capacitaci\u00f3n de Orientaci\u00f3n de Seguridad
# ===================================================================
def build_safety_orientation():
    doc = create_doc()
    add_title(doc)
    add_jurisdiction_tag(doc, "Idioma: Espa\u00f1ol | Material de capacitaci\u00f3n de seguridad")
    H1(doc, "Capacitaci\u00f3n de Orientaci\u00f3n de Seguridad")
    doc.add_paragraph()
    add_metadata_table(doc, "[CAP-ES-001]")

    H2(doc, "1. Prop\u00f3sito")
    P(doc, "Esta capacitaci\u00f3n de orientaci\u00f3n de seguridad tiene como objetivo proporcionar a todos los empleados nuevos y existentes el conocimiento fundamental sobre los protocolos de seguridad, los procedimientos de emergencia y las responsabilidades individuales para mantener un lugar de trabajo seguro en [NOMBRE DE LA EMPRESA].")

    H2(doc, "2. Salidas de Emergencia y Puntos de Reuni\u00f3n")
    B(doc, "Identifique las salidas de emergencia m\u00e1s cercanas a su \u00e1rea de trabajo. Cada planta cuenta con un m\u00ednimo de dos salidas de emergencia se\u00f1alizadas;")
    B(doc, "Las rutas de evacuaci\u00f3n est\u00e1n marcadas con se\u00f1ales verdes iluminadas. Familiar\u00edcese con ellas durante su primer d\u00eda;")
    B(doc, "Los puntos de reuni\u00f3n est\u00e1n ubicados en [UBICACI\u00d3N DEL PUNTO DE REUNI\u00d3N]. Despu\u00e9s de evacuar, dir\u00edjase al punto de reuni\u00f3n asignado y espere instrucciones;")
    B(doc, "Nunca utilice los elevadores durante una emergencia;")
    B(doc, "Nunca regrese al edificio hasta que las autoridades competentes lo autoricen.")

    H2(doc, "3. Extintores de Incendios")
    P(doc, "Los extintores de incendios est\u00e1n ubicados en puntos estrat\u00e9gicos de las instalaciones, se\u00f1alizados con letreros rojos.")
    P_bold(doc, "T\u00e9cnica P.A.S.S.:")
    B(doc, "Tire (Pull) del pasador de seguridad;", bold_prefix="P \u2014 ")
    B(doc, "Apunte (Aim) la boquilla hacia la base del fuego;", bold_prefix="A \u2014 ")
    B(doc, "Presione (Squeeze) la palanca de descarga;", bold_prefix="S \u2014 ")
    B(doc, "Barra (Sweep) de lado a lado hasta extinguir el fuego.", bold_prefix="S \u2014 ")
    P(doc, "Solo intente extinguir un fuego si es peque\u00f1o, tiene una ruta de escape despejada y ha sido capacitado. En caso de duda, evacue inmediatamente.")

    H2(doc, "4. Ubicaciones de Primeros Auxilios")
    B(doc, "Los botiquines de primeros auxilios est\u00e1n ubicados en: [LISTAR UBICACIONES];")
    B(doc, "Los desfibriladores externos autom\u00e1ticos (DEA) est\u00e1n ubicados en: [LISTAR UBICACIONES];")
    B(doc, "El personal capacitado en primeros auxilios y RCP incluye: [LISTAR NOMBRES];")
    B(doc, "En caso de emergencia m\u00e9dica, llame al 911 inmediatamente y luego notifique a su supervisor.")

    H2(doc, "5. Requisitos de Equipo de Protecci\u00f3n Personal (EPP)")
    P(doc, "Seg\u00fan su \u00e1rea de trabajo, se le puede requerir el uso de EPP espec\u00edfico:")
    make_table(doc,
        ["\u00c1rea / Actividad", "EPP Requerido"],
        [
            ["\u00c1rea de producci\u00f3n / planta", "Casco, gafas de seguridad, calzado de seguridad, protecci\u00f3n auditiva"],
            ["Trabajo con qu\u00edmicos", "Guantes qu\u00edmicos, gafas protectoras, respirador, delantal"],
            ["Trabajo en alturas (m\u00e1s de 1.8 m)", "Arn\u00e9s de cuerpo completo, l\u00ednea de vida, casco"],
            ["Trabajo con maquinaria", "Gafas de seguridad, guantes apropiados, protecci\u00f3n auditiva"],
            ["Trabajo de soldadura", "Careta de soldador, guantes de cuero, delantal ign\u00edfugo"],
            ["Oficina / \u00e1reas administrativas", "Generalmente no se requiere EPP; seguir se\u00f1alizaci\u00f3n al visitar planta"],
        ],
        bold_first_col=True)
    P(doc, "El EPP debe ser inspeccionado antes de cada uso. Reporte cualquier EPP da\u00f1ado o defectuoso a su supervisor para su reemplazo inmediato.")

    H2(doc, "6. Procedimiento de Reporte de Incidentes")
    P(doc, "Todo incidente, lesi\u00f3n, enfermedad laboral o cuasi-accidente debe reportarse de inmediato:")
    B(doc, "Notifique a su supervisor inmediato verbalmente lo antes posible;", bold_prefix="Paso 1: ")
    B(doc, "Complete el formulario de reporte de incidentes dentro de las primeras 24 horas;", bold_prefix="Paso 2: ")
    B(doc, "Coopere con la investigaci\u00f3n del incidente proporcionando informaci\u00f3n veraz y completa;", bold_prefix="Paso 3: ")
    B(doc, "Siga las instrucciones m\u00e9dicas y mantenga informado a Recursos Humanos sobre su recuperaci\u00f3n.", bold_prefix="Paso 4: ")
    P(doc, "Ning\u00fan empleado ser\u00e1 sancionado por reportar de buena fe un incidente o condici\u00f3n insegura.")

    H2(doc, "7. Identificaci\u00f3n de Peligros")
    P(doc, "Todos los empleados son responsables de identificar y reportar peligros potenciales en su \u00e1rea de trabajo. Busque activamente:")
    B(doc, "Derrames de l\u00edquidos o superficies resbaladizas;")
    B(doc, "Cables el\u00e9ctricos expuestos o da\u00f1ados;")
    B(doc, "Obstrucciones en pasillos, salidas o rutas de evacuaci\u00f3n;")
    B(doc, "Equipos o maquinaria da\u00f1ados o con protecciones faltantes;")
    B(doc, "Condiciones de iluminaci\u00f3n o ventilaci\u00f3n inadecuadas;")
    B(doc, "Almacenamiento inseguro de materiales (apilamiento inestable, productos qu\u00edmicos sin etiquetar).")

    H2(doc, "8. Derecho a Rechazar Trabajo Inseguro")
    P(doc, "Usted tiene el derecho de rechazar una tarea si cree razonablemente que representa un peligro inminente para su vida o salud. Para ejercer este derecho:")
    B(doc, "Informe a su supervisor de inmediato sobre la condici\u00f3n peligrosa;")
    B(doc, "Expl\u00edquele por qu\u00e9 considera que la tarea es insegura;")
    B(doc, "Permanezca en un lugar seguro mientras se investiga la situaci\u00f3n;")
    B(doc, "No se le podr\u00e1 sancionar por ejercer este derecho de buena fe.")

    page_break(doc)
    H1(doc, "Examen de Comprensi\u00f3n y Acuse de Recibo")
    P(doc, "Complete las siguientes preguntas para confirmar su comprensi\u00f3n del material de capacitaci\u00f3n de seguridad.")
    doc.add_paragraph()

    P_bold(doc, "1. \u00bfD\u00f3nde se encuentra la salida de emergencia m\u00e1s cercana a su \u00e1rea de trabajo?")
    P(doc, "Respuesta: ________________________________________________________________________")
    doc.add_paragraph()
    P_bold(doc, "2. \u00bfCu\u00e1l es el punto de reuni\u00f3n designado para su \u00e1rea?")
    P(doc, "Respuesta: ________________________________________________________________________")
    doc.add_paragraph()
    P_bold(doc, "3. Describa los cuatro pasos de la t\u00e9cnica P.A.S.S. para usar un extintor:")
    P(doc, "Respuesta: ________________________________________________________________________")
    doc.add_paragraph()
    P_bold(doc, "4. \u00bfQu\u00e9 EPP se requiere en su \u00e1rea de trabajo?")
    P(doc, "Respuesta: ________________________________________________________________________")
    doc.add_paragraph()
    P_bold(doc, "5. \u00bfA qui\u00e9n debe reportar un incidente o condici\u00f3n insegura?")
    P(doc, "Respuesta: ________________________________________________________________________")
    doc.add_paragraph()
    P_bold(doc, "6. \u00bfTiene usted derecho a rechazar una tarea que considere insegura? \u00bfQu\u00e9 pasos debe seguir?")
    P(doc, "Respuesta: ________________________________________________________________________")

    doc.add_paragraph()
    doc.add_paragraph()
    H2(doc, "Acuse de Recibo de Capacitaci\u00f3n")
    P(doc, "Yo, ________________________ (nombre en letra de molde), certifico que he asistido y completado la capacitaci\u00f3n de orientaci\u00f3n de seguridad de [NOMBRE DE LA EMPRESA]. Entiendo los procedimientos de seguridad, los requisitos de EPP, los protocolos de emergencia y mi responsabilidad de mantener un lugar de trabajo seguro.")
    doc.add_paragraph()
    table = doc.add_table(rows=4, cols=2)
    data = [
        ("Firma del Empleado: ________________________", "Fecha: ________________________"),
        ("Nombre del Empleado: ________________________", "N\u00famero de Empleado: ________________________"),
        ("Nombre del Capacitador: ________________________", "Firma del Capacitador: ________________________"),
        ("Departamento: ________________________", "Puesto: ________________________"),
    ]
    for i, (left, right) in enumerate(data):
        table.cell(i, 0).text = left
        table.cell(i, 1).text = right

    add_footer_section(doc)

    path = os.path.join(TRAIN_DIR, "Training_SafetyOrientation_Spanish.docx")
    doc.save(path)
    print(f"  [OK] {path}")
    return path


# ===================================================================
# 7. LOTO SOP -- Procedimiento de Bloqueo y Etiquetado
# ===================================================================
def build_loto():
    doc = create_doc()
    add_title(doc)
    add_jurisdiction_tag(doc, "Idioma: Espa\u00f1ol | Referencia: OSHA 29 CFR 1910.147")
    H1(doc, "Procedimiento de Bloqueo y Etiquetado (LOTO)")
    doc.add_paragraph()
    add_metadata_table(doc, "[SOP-ES-001]")

    H2(doc, "1. Prop\u00f3sito")
    P(doc, "Este Procedimiento Operativo Est\u00e1ndar (SOP) establece los requisitos m\u00ednimos para el control de energ\u00eda peligrosa (bloqueo/etiquetado o LOTO) durante el servicio y mantenimiento de m\u00e1quinas y equipos en los que la activaci\u00f3n inesperada o la liberaci\u00f3n de energ\u00eda almacenada podr\u00eda causar lesiones a los trabajadores. Este procedimiento cumple con la norma OSHA 29 CFR 1910.147 \u2014 Control de Energ\u00eda Peligrosa.")

    H2(doc, "2. Alcance")
    P(doc, "Este procedimiento aplica a todos los empleados, contratistas y proveedores de servicios de [NOMBRE DE LA EMPRESA] que realicen actividades de servicio, mantenimiento, reparaci\u00f3n, instalaci\u00f3n, ajuste, inspecci\u00f3n o limpieza de m\u00e1quinas y equipos en los que exista el riesgo de activaci\u00f3n inesperada o liberaci\u00f3n de energ\u00eda almacenada.")

    H2(doc, "3. Definiciones")
    B(doc, "Dispositivo de cierre utilizado para mantener un interruptor de aislamiento de energ\u00eda en posici\u00f3n segura, impidiendo la activaci\u00f3n del equipo.", bold_prefix="Candado (Lockout): ")
    B(doc, "Dispositivo de advertencia que se coloca en el punto de aislamiento de energ\u00eda, indicando que el equipo no debe operarse.", bold_prefix="Etiqueta (Tagout): ")
    B(doc, "Persona capacitada y autorizada que aplica dispositivos de bloqueo o etiquetado a los dispositivos de aislamiento de energ\u00eda.", bold_prefix="Empleado Autorizado: ")
    B(doc, "Persona cuyo trabajo requiere operar o utilizar una m\u00e1quina o equipo en el que se realiza servicio bajo bloqueo/etiquetado.", bold_prefix="Empleado Afectado: ")
    B(doc, "Cualquier forma de energ\u00eda incluyendo: el\u00e9ctrica, mec\u00e1nica, hidr\u00e1ulica, neum\u00e1tica, qu\u00edmica, t\u00e9rmica, gravitacional u otras formas de energ\u00eda.", bold_prefix="Energ\u00eda Peligrosa: ")

    H2(doc, "4. Responsabilidades")
    H3(doc, "4.1 Gerencia / Supervisi\u00f3n")
    B(doc, "Asegurar que todos los empleados afectados y autorizados reciban capacitaci\u00f3n inicial y de actualizaci\u00f3n anual;")
    B(doc, "Proporcionar candados, etiquetas y dispositivos de bloqueo adecuados;")
    B(doc, "Realizar inspecciones peri\u00f3dicas anuales del procedimiento LOTO;")
    B(doc, "Hacer cumplir este procedimiento y tomar acciones disciplinarias por incumplimiento.")

    H3(doc, "4.2 Empleados Autorizados")
    B(doc, "Seguir cada paso de este procedimiento al realizar bloqueo/etiquetado;")
    B(doc, "Utilizar \u00fanicamente sus propios candados y etiquetas personales asignados;")
    B(doc, "No retirar el candado o etiqueta de otro trabajador.")

    H3(doc, "4.3 Empleados Afectados")
    B(doc, "No intentar arrancar, operar o utilizar equipos que est\u00e9n bloqueados/etiquetados;")
    B(doc, "Notificar al supervisor si observan un equipo bloqueado sin supervisi\u00f3n.")

    H2(doc, "5. Procedimiento Paso a Paso")
    H3(doc, "5.1 Notificaci\u00f3n")
    P(doc, "Antes de iniciar el procedimiento de bloqueo/etiquetado, el empleado autorizado debe notificar a todos los empleados afectados y operadores del equipo que se va a realizar un servicio y que el equipo ser\u00e1 bloqueado.")

    H3(doc, "5.2 Identificaci\u00f3n de Fuentes de Energ\u00eda")
    P(doc, "Identifique todas las fuentes de energ\u00eda asociadas con el equipo utilizando la documentaci\u00f3n del equipo, las hojas de procedimiento espec\u00edfico de la m\u00e1quina y la inspecci\u00f3n visual.")

    H3(doc, "5.3 Apagado del Equipo")
    P(doc, "Apague el equipo utilizando el procedimiento normal de parada. Aseg\u00farese de que todas las partes m\u00f3viles se hayan detenido completamente antes de proceder.")

    H3(doc, "5.4 Aislamiento de Fuentes de Energ\u00eda")
    P(doc, "Opere los dispositivos de aislamiento de energ\u00eda (interruptores, v\u00e1lvulas, desconectores) para aislar completamente el equipo de todas las fuentes de energ\u00eda.")

    H3(doc, "5.5 Aplicaci\u00f3n de Dispositivos de Bloqueo/Etiquetado")
    P(doc, "Coloque su candado personal y etiqueta en cada dispositivo de aislamiento de energ\u00eda. La etiqueta debe indicar: nombre del empleado autorizado, fecha, raz\u00f3n del bloqueo y fecha estimada de finalizaci\u00f3n. Si m\u00faltiples empleados trabajan en el mismo equipo, cada uno debe colocar su propio candado.")

    H3(doc, "5.6 Liberaci\u00f3n de Energ\u00eda Almacenada")
    P(doc, "Despu\u00e9s del aislamiento, libere o controle toda la energ\u00eda almacenada o residual:")
    B(doc, "Descargue capacitores;")
    B(doc, "Libere la presi\u00f3n hidr\u00e1ulica y neum\u00e1tica;")
    B(doc, "Bloquee o asegure piezas que puedan moverse por gravedad;")
    B(doc, "Permita que las superficies calientes se enfr\u00eden;")
    B(doc, "Desconecte bater\u00edas o fuentes de energ\u00eda de respaldo.")

    H3(doc, "5.7 Verificaci\u00f3n del Aislamiento")
    P(doc, "Antes de iniciar el trabajo, el empleado autorizado debe verificar que el equipo est\u00e1 completamente aislado y des-energizado:")
    B(doc, "Intente arrancar el equipo utilizando los controles normales de operaci\u00f3n (el equipo NO debe arrancar);")
    B(doc, "Utilice equipos de medici\u00f3n (mult\u00edmetro, man\u00f3metro) para confirmar la ausencia de energ\u00eda;")
    B(doc, "Regrese los controles de operaci\u00f3n a la posici\u00f3n neutral/apagado despu\u00e9s de la verificaci\u00f3n.")

    H2(doc, "6. Tabla de Fuentes de Energ\u00eda y M\u00e9todos de Aislamiento")
    make_table(doc,
        ["Tipo de Energ\u00eda", "Fuente Com\u00fan", "Dispositivo de Aislamiento", "Verificaci\u00f3n"],
        [
            ["El\u00e9ctrica", "Paneles el\u00e9ctricos, transformadores, motores", "Interruptor/desconector de circuito", "Mult\u00edmetro / probador de voltaje"],
            ["Mec\u00e1nica", "Bandas, engranajes, volantes de inercia, resortes", "Bloqueo f\u00edsico, calzas, pines", "Inspecci\u00f3n visual de partes m\u00f3viles"],
            ["Hidr\u00e1ulica", "Bombas, cilindros, acumuladores", "V\u00e1lvulas de corte, purga de presi\u00f3n", "Man\u00f3metro (lectura de 0 PSI)"],
            ["Neum\u00e1tica", "Compresores, l\u00edneas de aire, cilindros", "V\u00e1lvulas de corte, purga de l\u00edneas", "Man\u00f3metro (lectura de 0 PSI)"],
            ["T\u00e9rmica", "Hornos, calderas, vapor, superficies calientes", "V\u00e1lvulas de corte, ventilaci\u00f3n", "Term\u00f3metro / medici\u00f3n de temperatura"],
            ["Qu\u00edmica", "Tuber\u00edas de proceso, tanques de almacenamiento", "V\u00e1lvulas de corte, purga, bridas ciegas", "Pruebas atmosf\u00e9ricas / detector de gas"],
            ["Gravitacional", "Componentes elevados, puertas de contenedores", "Calzas, bloques, soportes mec\u00e1nicos", "Inspecci\u00f3n visual de estabilidad"],
        ],
        bold_first_col=True)

    H2(doc, "7. Restauraci\u00f3n del Equipo al Servicio")
    P(doc, "Una vez completado el trabajo de servicio o mantenimiento:")
    B(doc, "Inspeccione el \u00e1rea de trabajo para asegurar que todas las herramientas y materiales han sido retirados;", bold_prefix="Paso 1: ")
    B(doc, "Verifique que todas las protecciones y dispositivos de seguridad del equipo est\u00e9n instalados correctamente;", bold_prefix="Paso 2: ")
    B(doc, "Confirme que todos los empleados est\u00e9n alejados del equipo y en una posici\u00f3n segura;", bold_prefix="Paso 3: ")
    B(doc, "Notifique a los empleados afectados que los dispositivos de bloqueo/etiquetado ser\u00e1n retirados;", bold_prefix="Paso 4: ")
    B(doc, "El empleado autorizado que coloc\u00f3 el candado lo retira personalmente;", bold_prefix="Paso 5: ")
    B(doc, "Reconecte las fuentes de energ\u00eda y arranque el equipo siguiendo el procedimiento normal.", bold_prefix="Paso 6: ")

    H2(doc, "8. Procedimiento de Cambio de Turno")
    P(doc, "Cuando el trabajo de mantenimiento se extiende m\u00e1s all\u00e1 de un turno, se debe garantizar la continuidad del bloqueo/etiquetado:")
    B(doc, "El empleado del turno entrante coloca su candado personal en el dispositivo de aislamiento ANTES de que el empleado del turno saliente retire su candado;")
    B(doc, "Se realiza una transferencia verbal de informaci\u00f3n sobre el estado del trabajo y los peligros presentes;")
    B(doc, "Se documenta la transferencia en el registro de bloqueo/etiquetado.")

    H2(doc, "9. Violaciones y Consecuencias")
    P(doc, "Las violaciones de este procedimiento ponen en peligro vidas y ser\u00e1n tratadas con la m\u00e1xima seriedad:")
    B(doc, "Retirar el candado o etiqueta de otro trabajador sin autorizaci\u00f3n \u2014 Terminaci\u00f3n inmediata;")
    B(doc, "No seguir el procedimiento LOTO antes del servicio \u2014 Suspensi\u00f3n sin goce de sueldo; recapacitaci\u00f3n obligatoria;")
    B(doc, "Operar equipo bloqueado/etiquetado \u2014 Suspensi\u00f3n sin goce de sueldo; investigaci\u00f3n;")
    B(doc, "No verificar el aislamiento antes de iniciar el trabajo \u2014 Amonestaci\u00f3n escrita; recapacitaci\u00f3n.")

    add_acknowledgment(doc)
    add_approval_table(doc)
    add_footer_section(doc)

    path = os.path.join(SOP_DIR, "SOP_LOTO_Spanish.docx")
    doc.save(path)
    print(f"  [OK] {path}")
    return path


# ===================================================================
# 8. PPE POLICY -- Pol\u00edtica de Equipo de Protecci\u00f3n Personal (EPP)
# ===================================================================
def build_ppe():
    doc = create_doc()
    add_title(doc)
    add_jurisdiction_tag(doc, "Idioma: Espa\u00f1ol | Referencia: OSHA 29 CFR 1910.132-138")
    H1(doc, "Pol\u00edtica de Equipo de Protecci\u00f3n Personal (EPP)")
    doc.add_paragraph()
    add_metadata_table(doc, "[POL-ES-005]")

    H2(doc, "1. Prop\u00f3sito")
    P(doc, "Esta pol\u00edtica establece los requisitos de [NOMBRE DE LA EMPRESA] para la selecci\u00f3n, uso, mantenimiento y reemplazo del Equipo de Protecci\u00f3n Personal (EPP) con el fin de proteger a los empleados contra peligros en el lugar de trabajo que puedan causar lesiones o enfermedades. Esta pol\u00edtica cumple con las normas de OSHA 29 CFR 1910.132 a 1910.138 y regulaciones estatales aplicables.")

    H2(doc, "2. Alcance")
    P(doc, "Esta pol\u00edtica aplica a todos los empleados, contratistas, subcontratistas, trabajadores temporales y visitantes en todas las instalaciones y sitios de trabajo de [NOMBRE DE LA EMPRESA] donde existan peligros que requieran el uso de EPP.")

    H2(doc, "3. Obligaci\u00f3n del Empleador")
    B(doc, "Realizar una evaluaci\u00f3n de peligros (Hazard Assessment) en cada \u00e1rea de trabajo para identificar los riesgos presentes y determinar el EPP apropiado;")
    B(doc, "Proporcionar el EPP necesario sin costo alguno para el empleado;")
    B(doc, "Asegurar que el EPP proporcionado sea del tama\u00f1o y ajuste correcto para cada empleado;")
    B(doc, "Capacitar a los empleados sobre el uso correcto, las limitaciones, el cuidado y el mantenimiento del EPP;")
    B(doc, "Reemplazar el EPP da\u00f1ado, defectuoso o al final de su vida \u00fatil sin costo para el empleado;")
    B(doc, "Mantener registros de la evaluaci\u00f3n de peligros, las certificaciones de EPP y la capacitaci\u00f3n proporcionada.")

    H2(doc, "4. Obligaci\u00f3n del Empleado")
    B(doc, "Usar correctamente el EPP designado para su \u00e1rea de trabajo y actividades en todo momento;")
    B(doc, "Inspeccionar el EPP antes de cada uso para detectar da\u00f1os o defectos;")
    B(doc, "Reportar inmediatamente al supervisor cualquier EPP da\u00f1ado, defectuoso o que no ajuste correctamente;")
    B(doc, "Cuidar, limpiar y almacenar el EPP de acuerdo con las instrucciones del fabricante;")
    B(doc, "Participar en todas las capacitaciones de EPP requeridas;")
    B(doc, "No modificar, alterar ni desactivar ning\u00fan componente del EPP.")

    H2(doc, "5. Tabla de EPP por Tipo de Peligro")
    make_table(doc,
        ["Tipo de Peligro", "EPP Requerido", "Norma de Referencia"],
        [
            ["Protecci\u00f3n de la cabeza", "Casco de seguridad (Tipo I o II seg\u00fan el peligro); casco contra impactos para riesgos menores", "OSHA 29 CFR 1910.135; ANSI/ISEA Z89.1"],
            ["Protecci\u00f3n ocular y facial", "Gafas de seguridad, goggles, caretas faciales, lentes con filtro para soldadura, lentes con protecci\u00f3n UV", "OSHA 29 CFR 1910.133; ANSI/ISEA Z87.1"],
            ["Protecci\u00f3n auditiva", "Tapones auditivos, orejeras (cuando la exposici\u00f3n al ruido excede 85 dBA TWA)", "OSHA 29 CFR 1910.95; ANSI S3.19 / S12.6"],
            ["Protecci\u00f3n de manos", "Guantes de cuero, guantes qu\u00edmicos (nitrilo, neopreno, PVC), guantes resistentes a cortes, guantes diel\u00e9ctricos", "OSHA 29 CFR 1910.138; ANSI/ISEA 105"],
            ["Protecci\u00f3n de pies", "Calzado de seguridad con puntera de acero o composite, calzado diel\u00e9ctrico, botas resistentes a qu\u00edmicos", "OSHA 29 CFR 1910.136; ASTM F2413"],
            ["Protecci\u00f3n respiratoria", "Respiradores de part\u00edculas (N95, P100), respiradores de media cara, respiradores de cara completa, equipos de suministro de aire (SCBA)", "OSHA 29 CFR 1910.134; NIOSH 42 CFR 84"],
            ["Protecci\u00f3n contra ca\u00eddas", "Arn\u00e9s de cuerpo completo, l\u00ednea de vida, puntos de anclaje, sistemas de detenci\u00f3n de ca\u00eddas (cuando se trabaja a m\u00e1s de 1.8 metros de altura)", "OSHA 29 CFR 1910.140; ANSI Z359.1"],
            ["Protecci\u00f3n del cuerpo", "Overoles, delantales, trajes qu\u00edmicos, chalecos de alta visibilidad, ropa ign\u00edfuga (FR)", "OSHA 29 CFR 1910.132; NFPA 2112 (FR)"],
        ],
        bold_first_col=True)

    H2(doc, "6. Requisitos de Inspecci\u00f3n")
    P(doc, "Todo el EPP debe ser inspeccionado regularmente para garantizar su funcionalidad y protecci\u00f3n:")
    B(doc, "El empleado debe inspeccionar visualmente su EPP antes de cada uso;", bold_prefix="Antes de cada uso: ")
    B(doc, "Inspecciones formales documentadas por el supervisor o el departamento de seguridad;", bold_prefix="Mensual: ")
    B(doc, "Evaluaci\u00f3n completa de todo el inventario de EPP por el coordinador de seguridad;", bold_prefix="Anual: ")
    B(doc, "Criterios de retiro: grietas, deformaciones, p\u00e9rdida de elasticidad, desgaste excesivo, da\u00f1o por qu\u00edmicos, da\u00f1o por impacto, caducidad del fabricante.")

    H2(doc, "7. Procedimiento de Reemplazo")
    B(doc, "El empleado notifica a su supervisor sobre la necesidad de reemplazo;")
    B(doc, "El supervisor autoriza el reemplazo y emite una solicitud al almac\u00e9n o proveedor;")
    B(doc, "El EPP da\u00f1ado o caducado se retira de servicio y se desecha de manera apropiada;")
    B(doc, "El nuevo EPP se entrega al empleado y se registra en el historial de EPP del empleado;")
    B(doc, "Todo reemplazo se realiza sin costo para el empleado.")

    H2(doc, "8. Requisitos de Capacitaci\u00f3n")
    P(doc, "[NOMBRE DE LA EMPRESA] proporcionar\u00e1 capacitaci\u00f3n sobre EPP que incluya:")
    B(doc, "Cu\u00e1ndo es necesario el uso de EPP;")
    B(doc, "Qu\u00e9 tipo de EPP es necesario para cada tarea y peligro;")
    B(doc, "C\u00f3mo ponerse, ajustar, usar y quitarse correctamente el EPP;")
    B(doc, "Limitaciones del EPP;")
    B(doc, "Cuidado, mantenimiento, vida \u00fatil y almacenamiento adecuado del EPP;")
    B(doc, "Procedimiento para reportar EPP da\u00f1ado o defectuoso.")
    P(doc, "La capacitaci\u00f3n se proporcionar\u00e1 al momento de la contrataci\u00f3n, cuando se asigne nuevo EPP, cuando cambien las condiciones de trabajo y como actualizaci\u00f3n anual.")

    H2(doc, "9. Consecuencias por Incumplimiento")
    P(doc, "El uso del EPP no es opcional cuando ha sido designado como requerido para un \u00e1rea o tarea. Las consecuencias por no usar el EPP requerido incluyen:")
    B(doc, "Advertencia verbal y correcci\u00f3n inmediata;", bold_prefix="Primera infracci\u00f3n: ")
    B(doc, "Amonestaci\u00f3n escrita y recapacitaci\u00f3n obligatoria;", bold_prefix="Segunda infracci\u00f3n: ")
    B(doc, "Suspensi\u00f3n sin goce de sueldo;", bold_prefix="Tercera infracci\u00f3n: ")
    B(doc, "Terminaci\u00f3n del empleo.", bold_prefix="Cuarta infracci\u00f3n: ")
    P(doc, "En situaciones de peligro inminente, la primera violaci\u00f3n puede resultar en medidas disciplinarias m\u00e1s severas.")

    add_acknowledgment(doc)
    add_approval_table(doc)
    add_footer_section(doc)

    path = os.path.join(POLICY_DIR, "Policy_PPE_Spanish.docx")
    doc.save(path)
    print(f"  [OK] {path}")
    return path


# ===================================================================
# MAIN -- Generate all 8 documents
# ===================================================================
def main():
    # Ensure directories exist
    for d in [POLICY_DIR, ONBOARD_DIR, TRAIN_DIR, SOP_DIR]:
        os.makedirs(d, exist_ok=True)

    print("Building Spanish-language HR content for Pure Work AI...\n")

    paths = []
    paths.append(build_harassment())
    paths.append(build_health_safety())
    paths.append(build_drug_alcohol())
    paths.append(build_attendance())
    paths.append(build_onboarding())
    paths.append(build_safety_orientation())
    paths.append(build_loto())
    paths.append(build_ppe())

    print(f"\nDone. {len(paths)} documents generated successfully.")

    # Verification: check file sizes
    print("\nVerification:")
    all_ok = True
    for p in paths:
        if os.path.exists(p):
            size = os.path.getsize(p)
            status = "OK" if size > 1000 else "WARNING: small file"
            print(f"  {os.path.basename(p):50s} {size:>8,} bytes  [{status}]")
            if size < 1000:
                all_ok = False
        else:
            print(f"  {p:50s}  MISSING!")
            all_ok = False

    if all_ok:
        print("\nAll files generated and verified.")
    else:
        print("\nSome files may have issues. Check warnings above.")


if __name__ == "__main__":
    main()
