#!/usr/bin/env python3
"""
Build 9 South Africa HR content documents for Pure Work AI.

Documents:
  1. Policy_BCEA_SouthAfrica.docx          (Basic Conditions of Employment Act)
  2. Policy_LRA_SouthAfrica.docx           (Labour Relations Act)
  3. Policy_EEA_SouthAfrica.docx           (Employment Equity Act)
  4. Policy_OHSA_SouthAfrica.docx          (Occupational Health & Safety Act)
  5. Policy_COIDA_SouthAfrica.docx         (Compensation for Occupational Injuries)
  6. Policy_UIF_PAYE_SDL_SouthAfrica.docx  (Statutory Deductions)
  7. Onboarding_SouthAfrica.docx           (SA-specific onboarding playbook)
  8. Training_SALabourLaw.docx             (Labour law training for managers)
  9. SOP_SouthAfrica_Pack.docx             (SA compliance SOP pack)

Formatting specs (matching existing PureWorkAI files):
- Normal: Calibri 11pt, color #333333
- Heading 1: Calibri 18pt, bold, color #1A3A5C
- Heading 2: Calibri 14pt, bold, color #1A3A5C
- Heading 3: Calibri 12pt, bold, color #1A3A5C
- Title line: Calibri 22pt, bold, color #1A3A5C, centered
- Jurisdiction tag: Calibri 11pt, italic, color #666666, centered
- Table cells: Calibri 10pt
- Table header rows: shading #1A3A5C with white bold text
- Footer line: Calibri 8pt, italic, color #999999, centered
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml
import os

# ---------------------------------------------------------------------------
# Directories
# ---------------------------------------------------------------------------
BASE = "/home/aiciv/purework-ai/public/content"
POLICY_DIR = os.path.join(BASE, "policies")
ONBOARD_DIR = os.path.join(BASE, "onboarding")
TRAIN_DIR = os.path.join(BASE, "training")
SOP_DIR = os.path.join(BASE, "sop-templates")

# ---------------------------------------------------------------------------
# Colors
# ---------------------------------------------------------------------------
NAVY = RGBColor(0x1A, 0x3A, 0x5C)
ACCENT = RGBColor(0x2C, 0x7B, 0xE5)
DARK_TEXT = RGBColor(0x33, 0x33, 0x33)
GRAY_TAG = RGBColor(0x66, 0x66, 0x66)
LIGHT_GRAY = RGBColor(0x99, 0x99, 0x99)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


# ===================================================================
# Helper functions -- matching existing PureWorkAI formatting exactly
# ===================================================================

def create_doc():
    """Create a new document with styles matching existing PureWorkAI policies."""
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Emu(139700)  # 11pt
    style.font.color.rgb = DARK_TEXT

    h1 = doc.styles['Heading 1']
    h1.font.name = 'Calibri'
    h1.font.size = Emu(228600)  # 18pt
    h1.font.bold = True
    h1.font.color.rgb = NAVY
    h1.paragraph_format.space_before = Emu(304800)
    h1.paragraph_format.space_after = Emu(0)

    h2 = doc.styles['Heading 2']
    h2.font.name = 'Calibri'
    h2.font.size = Emu(177800)  # 14pt
    h2.font.bold = True
    h2.font.color.rgb = NAVY
    h2.paragraph_format.space_before = Emu(127000)
    h2.paragraph_format.space_after = Emu(0)

    h3 = doc.styles['Heading 3']
    h3.font.name = 'Calibri'
    h3.font.size = Emu(152400)  # 12pt
    h3.font.bold = True
    h3.font.color.rgb = NAVY
    h3.paragraph_format.space_before = Emu(127000)
    h3.paragraph_format.space_after = Emu(0)

    style_lb = doc.styles['List Bullet']
    style_lb.font.name = 'Calibri'
    style_lb.font.size = Emu(139700)
    style_lb.font.color.rgb = DARK_TEXT

    return doc


def add_title(doc, company_line="[COMPANY NAME]"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(company_line)
    run.font.name = 'Calibri'
    run.font.size = Emu(279400)  # 22pt
    run.font.bold = True
    run.font.color.rgb = NAVY


def add_jurisdiction_tag(doc, tag_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(tag_text)
    run.font.name = 'Calibri'
    run.font.size = Emu(139700)
    run.italic = True
    run.font.color.rgb = GRAY_TAG


def add_metadata_table(doc, policy_number):
    table = doc.add_table(rows=5, cols=2)
    data = [
        ("Policy Number:", policy_number),
        ("Version:", "1.0"),
        ("Effective Date:", "[EFFECTIVE DATE]"),
        ("Review Date:", "[REVIEW DATE]"),
        ("Owner:", "Human Resources"),
    ]
    for i, (label, value) in enumerate(data):
        for j, text in enumerate([label, value]):
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.name = 'Calibri'
            run.font.size = Emu(127000)  # 10pt


def H1(doc, text):
    doc.add_heading(text, level=1)


def H2(doc, text):
    doc.add_heading(text, level=2)


def H3(doc, text):
    doc.add_heading(text, level=3)


def P(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def P_bold(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    return p


def B(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def style_header_row(row, color_hex="1A3A5C"):
    for cell in row.cells:
        set_cell_shading(cell, color_hex)
        for pr in cell.paragraphs:
            for run in pr.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.name = "Calibri"


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '</w:tblBorders>')
    tblPr.append(borders)


def make_table(doc, headers, rows_data, bold_first_col=False):
    """Create a table with navy header row and optional bold first column."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        hdr.cells[i].text = h
    style_header_row(hdr)
    for row_data in rows_data:
        row = table.add_row()
        for i, text in enumerate(row_data):
            cell = row.cells[i]
            cell.text = ""
            pr = cell.paragraphs[0]
            run = pr.add_run(text)
            run.font.size = Pt(10)
            run.font.name = "Calibri"
            if bold_first_col and i == 0:
                run.font.bold = True
    doc.add_paragraph()
    return table


def add_approval_table(doc):
    H2(doc, "Approval")
    table = doc.add_table(rows=4, cols=3)
    headers = ["Role", "Name", "Signature / Date"]
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.font.name = 'Calibri'
        run.font.size = Emu(127000)
        run.bold = True
    roles = ["President / CEO", "VP, Human Resources", "Legal Counsel"]
    for r, role in enumerate(roles, 1):
        table.cell(r, 0).text = role
        table.cell(r, 1).text = "________________________"
        table.cell(r, 2).text = "________________________"


def add_acknowledgment(doc):
    H2(doc, "Employee Acknowledgment")
    P(doc, 'I, ________________________ (print name), acknowledge that I have received, read, and understood this policy. I agree to comply with its terms and conditions as a condition of my employment with [COMPANY NAME].')
    doc.add_paragraph()
    table = doc.add_table(rows=3, cols=2)
    data = [
        ("Employee Signature: ________________________", "Date: ________________________"),
        ("Employee Name (Print): ________________________", "Employee ID: ________________________"),
        ("Department: ________________________", "Position: ________________________"),
    ]
    for i, (left, right) in enumerate(data):
        table.cell(i, 0).text = left
        table.cell(i, 1).text = right


def add_footer_line(doc, policy_name):
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Prepared by [COMPANY NAME] | {policy_name} | [DATE]")
    run.font.size = Emu(101600)  # 8pt
    run.italic = True
    run.font.color.rgb = LIGHT_GRAY


def page_break(doc):
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(WD_BREAK.PAGE)


# ===================================================================
# 1. BCEA - Basic Conditions of Employment Act
# ===================================================================
def build_bcea():
    doc = create_doc()
    add_title(doc)
    add_jurisdiction_tag(doc, "[SOUTH AFRICA] -- Basic Conditions of Employment Act 75 of 1997")
    H1(doc, "Basic Conditions of Employment Act (BCEA) Policy")
    doc.add_paragraph()
    add_metadata_table(doc, "[POLICY-SA-001]")

    H2(doc, "1. Purpose")
    P(doc, "This policy outlines the minimum conditions of employment as prescribed by the Basic Conditions of Employment Act 75 of 1997 (BCEA), as amended. [COMPANY NAME] commits to meeting or exceeding the statutory minimum standards set out in the BCEA for all employees within its South African operations. This policy applies alongside the employee's contract of employment, applicable Sectoral Determinations, and any Bargaining Council agreements.")

    H2(doc, "2. Scope")
    P(doc, "This policy applies to all employees of [COMPANY NAME] employed in the Republic of South Africa, except where specific sections of the BCEA do not apply to senior managerial employees or employees earning above the BCEA earnings threshold (as gazetted by the Minister of Employment and Labour from time to time).")

    H3(doc, "2.1 BCEA Earnings Threshold")
    P(doc, "Employees earning above the BCEA earnings threshold (currently R254,371.67 per annum, as updated periodically by Government Gazette) are excluded from certain provisions of the BCEA, including those relating to ordinary hours of work, overtime, rest periods, and meal intervals. These employees remain covered by provisions relating to leave, particulars of employment, and termination of employment.")

    H2(doc, "3. Ordinary Hours of Work")
    P(doc, "The BCEA prescribes the following maximum ordinary working hours:")
    make_table(doc,
        ["Work Schedule", "Maximum Daily Hours", "Maximum Weekly Hours"],
        [
            ["5-day work week", "9 hours per day", "45 hours per week"],
            ["6-day work week", "8 hours per day", "45 hours per week"],
        ])
    P(doc, "Employees may not be required or permitted to work more than 45 ordinary hours per week or more than the daily maximum without a written agreement to work overtime.")

    H3(doc, "3.1 Meal Intervals")
    P(doc, "An employee is entitled to a meal interval of at least one continuous hour after five hours of work. An agreement may reduce the meal interval to 30 minutes. The meal interval is unpaid unless the employee is required to remain available for work or to perform duties during the interval.")

    H3(doc, "3.2 Rest Periods")
    P(doc, "An employee is entitled to a daily rest period of at least 12 consecutive hours between ending work on one day and commencing work on the next day, and a weekly rest period of at least 36 consecutive hours, which must include a Sunday unless otherwise agreed.")

    H2(doc, "4. Overtime")
    P(doc, "Overtime may only be worked by agreement between the employer and the employee. The BCEA prescribes the following conditions for overtime:")
    B(doc, "Maximum overtime: 10 hours per week (or as extended by a collective agreement, not exceeding 15 hours per week);")
    B(doc, "An employee may not work more than 12 hours in any day (including ordinary hours and overtime);")
    B(doc, "Overtime compensation: 1.5 times the employee's normal hourly wage rate;")
    B(doc, "Work on a Sunday (for employees who do not ordinarily work on a Sunday): 2 times the employee's normal hourly wage rate;")
    B(doc, "Work on a public holiday: 2 times the employee's normal hourly wage rate;")
    B(doc, "By agreement, overtime may be compensated by paid time off equivalent to the overtime worked, granted within one month of the overtime being performed.")

    H2(doc, "5. Annual Leave")
    P(doc, "Every employee is entitled to annual leave as follows:")
    B(doc, "21 consecutive calendar days of annual leave per annual leave cycle (equivalent to 15 working days for a 5-day work week);")
    B(doc, "The annual leave cycle is 12 months of employment from the date of commencement or the completion of the previous leave cycle;")
    B(doc, "Annual leave must be granted by the employer within 6 months after the end of the annual leave cycle in which it accrued;")
    B(doc, "Annual leave may not be replaced by payment in lieu, except on termination of employment;")
    B(doc, "An employee is entitled to be paid the employee's normal remuneration for the period of annual leave before the commencement of the leave.")

    H2(doc, "6. Sick Leave")
    P(doc, "Sick leave is provided in a 36-month cycle commencing on the first day of employment:")
    B(doc, "An employee is entitled to an amount of paid sick leave equal to the number of days the employee would normally work during a period of six weeks (e.g., 30 days for a 5-day week worker, 36 days for a 6-day week worker) in each 36-month cycle;")
    B(doc, "During the first six months of employment, an employee is entitled to one day of paid sick leave for every 26 days worked;")
    B(doc, "A medical certificate is required if the employee is absent for more than two consecutive days, or on more than two occasions during an eight-week period;")
    B(doc, "The medical certificate must be issued by a medical practitioner, traditional health practitioner (registered under the Traditional Health Practitioners Act 22 of 2007), or any other person certified to diagnose and treat patients under the laws of the country.")

    H2(doc, "7. Family Responsibility Leave")
    P(doc, "An employee who has worked for at least four months and who works at least four days per week is entitled to three days of paid family responsibility leave per annual leave cycle. This leave may be taken:")
    B(doc, "When the employee's child is born;")
    B(doc, "When the employee's child is sick;")
    B(doc, "In the event of the death of the employee's spouse or life partner, parent, adoptive parent, grandparent, child, adopted child, grandchild, or sibling.")
    P(doc, "The employer may require reasonable proof of the event (e.g., birth certificate, medical certificate, death certificate).")

    H2(doc, "8. Maternity Leave")
    P(doc, "A pregnant employee is entitled to at least four consecutive months of maternity leave. The following conditions apply:")
    B(doc, "An employee may commence maternity leave at any time from four weeks before the expected date of birth, or earlier if a medical practitioner or midwife certifies it is necessary;")
    B(doc, "An employee may not work for six weeks after the birth of her child unless a medical practitioner or midwife certifies that she is fit to do so;")
    B(doc, "Maternity leave is unpaid under the BCEA. However, employees who have contributed to the Unemployment Insurance Fund (UIF) may claim maternity benefits from the UIF;")
    B(doc, "The employer may not require or permit a pregnant employee to perform work that is hazardous to her health or the health of her child;")
    B(doc, "No employee may be dismissed for reasons related to her pregnancy, and any such dismissal is automatically unfair under the Labour Relations Act 66 of 1995.")

    H2(doc, "9. Parental Leave")
    P(doc, "In terms of the Labour Laws Amendment Act 10 of 2018, which amended the BCEA:")
    B(doc, "An employee who is a parent of a child born on or after 1 January 2020 is entitled to at least 10 consecutive days of parental leave;")
    B(doc, "Parental leave commences on the day the child is born or the date of the adoption order;")
    B(doc, "Parental leave is unpaid under the BCEA but qualifying employees may claim from the UIF;")
    B(doc, "The 2025 expansion (Labour Laws Amendment Act): Shared parental leave allows parents to share up to four months of UIF-funded parental leave between them, improving flexibility and promoting gender equality in caregiving;")
    B(doc, "The employee must provide written notice to the employer at least one month before the expected date of commencement of parental leave.")

    H2(doc, "10. Notice Periods")
    P(doc, "Either party may terminate the contract of employment by giving written notice as follows:")
    make_table(doc,
        ["Period of Employment", "Minimum Notice Period"],
        [
            ["Less than 6 months", "1 week"],
            ["6 months to less than 1 year", "2 weeks"],
            ["1 year or more", "4 weeks"],
        ])
    P(doc, "Notice must be in writing, except when given by an illiterate employee. The notice period may not run concurrently with any period of leave, except sick leave.")

    H2(doc, "11. Severance Pay")
    P(doc, "An employee dismissed for reasons based on the employer's operational requirements (as defined in Section 189 of the Labour Relations Act) is entitled to severance pay of at least one week's remuneration for each completed year of continuous service with the employer. An employee who unreasonably refuses an offer of alternative employment by the employer or any other employer is not entitled to severance pay.")

    H2(doc, "12. Prohibited Conduct")
    P(doc, "The following conduct is prohibited under the BCEA:")
    B(doc, "Employment of a child under the age of 15 years (Section 43);")
    B(doc, "Employment of a child in work that is inappropriate for the child's age or that places the child at risk (Section 44);")
    B(doc, "Forced labour (Section 48);")
    B(doc, "Requiring an employee to waive their rights under the BCEA (Section 4).")

    H2(doc, "13. Record-Keeping")
    P(doc, "The employer must keep written records of the prescribed particulars of each employee for a period of three years after the termination of employment. Records must include:")
    B(doc, "Employee's name and occupation;")
    B(doc, "Time worked and remuneration paid;")
    B(doc, "Date of birth (if under 18);")
    B(doc, "Leave records;")
    B(doc, "Any other prescribed information.")

    H2(doc, "14. Policy Review")
    P(doc, "This policy will be reviewed annually, or more frequently as required by amendments to the BCEA, Sectoral Determinations, or changes in the earnings threshold.")

    add_approval_table(doc)
    add_acknowledgment(doc)
    add_footer_line(doc, "BCEA Policy -- South Africa")

    path = os.path.join(POLICY_DIR, "Policy_BCEA_SouthAfrica.docx")
    doc.save(path)
    return path


# ===================================================================
# 2. LRA - Labour Relations Act
# ===================================================================
def build_lra():
    doc = create_doc()
    add_title(doc)
    add_jurisdiction_tag(doc, "[SOUTH AFRICA] -- Labour Relations Act 66 of 1995")
    H1(doc, "Labour Relations Act (LRA) Policy")
    doc.add_paragraph()
    add_metadata_table(doc, "[POLICY-SA-002]")

    H2(doc, "1. Purpose")
    P(doc, "This policy sets out [COMPANY NAME]'s approach to employment relations in accordance with the Labour Relations Act 66 of 1995 (LRA), as amended. The LRA is the primary legislation governing the relationship between employers, employees, and trade unions in South Africa. This policy ensures that all employment decisions, including discipline, dismissal, and dispute resolution, are conducted fairly and lawfully.")

    H2(doc, "2. Scope")
    P(doc, "This policy applies to all employees of [COMPANY NAME] in South Africa, regardless of their level, function, or duration of employment. Independent contractors are not covered by the LRA unless they are deemed to be employees as defined in Section 200A (the presumption of employment).")

    H2(doc, "3. No At-Will Employment")
    P(doc, "South Africa does not recognize at-will employment. Under the LRA, every employee has the right not to be unfairly dismissed. This means:")
    B(doc, "The employer must have a fair reason for dismissal, recognized in law (substantive fairness);")
    B(doc, "The employer must follow a fair procedure before dismissing an employee (procedural fairness);")
    B(doc, "The burden of proof rests on the employer to demonstrate both substantive and procedural fairness.")

    H2(doc, "4. Types of Dismissal")
    P(doc, "The LRA recognizes three grounds for fair dismissal:")

    H3(doc, "4.1 Misconduct")
    P(doc, "Dismissal for misconduct arises when the employee has contravened a workplace rule or standard of conduct. To be fair, the misconduct must be serious, the rule must have been known (or should have been known) to the employee, and the rule must be consistently applied. Progressive discipline should normally be applied before dismissal for misconduct, except in cases of gross misconduct where summary dismissal may be warranted.")

    H3(doc, "4.2 Incapacity")
    P(doc, "Incapacity includes both poor work performance and ill health or injury. The approach differs for each:")
    B(doc, "Poor work performance: The employer must counsel the employee, set clear performance standards, provide training or support, allow a reasonable period for improvement, and consider alternative positions before dismissal;", bold_prefix="Poor Performance: ")
    B(doc, "Ill health or injury: The employer must investigate the extent of the incapacity, the cause, the likely duration, the possibility of securing temporary replacement, and consider alternative work or adapted duties before dismissal.", bold_prefix="Ill Health/Injury: ")

    H3(doc, "4.3 Operational Requirements")
    P(doc, "Dismissal for operational requirements (retrenchment) is based on the employer's economic, technological, structural, or similar needs. The employer must follow the consultation process prescribed by Section 189 (or Section 189A for large-scale retrenchments affecting 50 or more employees). The employer must:")
    B(doc, "Issue a written notice to the consulting parties (employees/union) disclosing the reasons for contemplated dismissals;")
    B(doc, "Consult in good faith with the aim of reaching consensus on avoiding, minimizing, or mitigating the effects of the dismissals;")
    B(doc, "Apply fair and objective selection criteria (e.g., LIFO -- Last In, First Out -- unless otherwise agreed);")
    B(doc, "Pay severance pay of at least one week's remuneration per completed year of service.")

    H2(doc, "5. Disciplinary Hearing Requirements")
    P(doc, "Before dismissing an employee for misconduct, the employer must conduct a fair hearing. The hearing must meet the following requirements:")
    B(doc, "The employee must be given reasonable notice of the charges against them (in writing);")
    B(doc, "The charges must be specific enough for the employee to prepare a defence;")
    B(doc, "The employee must be informed of the right to be represented by a fellow employee or trade union representative;")
    B(doc, "The employee must be given a reasonable opportunity to respond to the charges, present evidence, call witnesses, and cross-examine the employer's witnesses;")
    B(doc, "The chairperson of the hearing must be unbiased and should not be the complainant or direct supervisor (where possible);")
    B(doc, "The decision must be communicated to the employee in writing, with reasons and the right to appeal or refer the matter to the CCMA.")

    H2(doc, "6. CCMA Dispute Resolution")
    P(doc, "The Commission for Conciliation, Mediation and Arbitration (CCMA) is the primary body for resolving employment disputes in South Africa. The process operates as follows:")

    H3(doc, "6.1 Referral")
    P(doc, "An employee who believes they have been unfairly dismissed or subjected to an unfair labour practice must refer the dispute to the CCMA within 30 calendar days of the date of dismissal (or the date of the unfair conduct). Late referrals may be condoned in exceptional circumstances.")

    H3(doc, "6.2 Conciliation")
    P(doc, "The CCMA will first attempt to resolve the dispute through conciliation (a facilitated negotiation). Conciliation must be completed within 30 days of referral. If the dispute is resolved, the terms are recorded in a settlement agreement that is binding and enforceable as an arbitration award.")

    H3(doc, "6.3 Arbitration")
    P(doc, "If conciliation fails, the matter proceeds to arbitration (for unfair dismissal disputes) or the Labour Court (for automatically unfair dismissals and certain other matters). The CCMA arbitrator will hear evidence, make findings of fact, and issue a binding award. The employer bears the burden of proving the dismissal was fair.")

    H2(doc, "7. Protected Strikes and Lockouts")
    P(doc, "The LRA protects the right of employees to strike and the right of employers to lock out, provided the procedural requirements are met:")
    B(doc, "The dispute must first be referred to the CCMA for conciliation;")
    B(doc, "A certificate of non-resolution must be issued, or 30 days must have elapsed since the referral;")
    B(doc, "48 hours' written notice of the strike or lockout must be given;")
    B(doc, "No employee may be dismissed or disciplined for participating in a protected strike;")
    B(doc, "Unprotected strikes (those not following the correct procedure) may result in interdicts, damages claims, or dismissal.")

    H2(doc, "8. Unfair Dismissal and Unfair Labour Practices")
    P(doc, "An employee may challenge a dismissal as unfair if the employer cannot demonstrate both a fair reason and a fair procedure. Remedies for unfair dismissal include:")
    B(doc, "Reinstatement (the employee is restored to the position they held before dismissal);")
    B(doc, "Re-employment (the employee is employed in a comparable position);")
    B(doc, "Compensation (up to 12 months' remuneration, or 24 months for automatically unfair dismissals).")
    P(doc, "Unfair labour practices include unfair conduct relating to promotion, demotion, probation, training, provision of benefits, suspension, and failure to reinstate or re-employ a former employee.")

    H2(doc, "9. Automatically Unfair Dismissals")
    P(doc, "Section 187 of the LRA identifies dismissals that are automatically unfair, regardless of whether a fair procedure was followed. These include dismissals based on:")
    B(doc, "The employee's pregnancy, intended pregnancy, or any reason related to pregnancy;")
    B(doc, "The employee's participation or intended participation in lawful trade union activities or protected strikes;")
    B(doc, "The employee's refusal to agree to a demand in respect of any matter of mutual interest between employer and employee;")
    B(doc, "The employee's exercise of any right conferred by the LRA or other employment legislation;")
    B(doc, "Discrimination on any arbitrary ground, including race, gender, sex, ethnic or social origin, colour, sexual orientation, age, disability, religion, conscience, belief, political opinion, culture, language, marital status, or family responsibility;")
    B(doc, "A transfer or merger of the business under Section 197 of the LRA;")
    B(doc, "A protected disclosure (whistleblowing) under the Protected Disclosures Act 26 of 2000.")
    P(doc, "Compensation for automatically unfair dismissal may be up to 24 months' remuneration.")

    H2(doc, "10. New Code of Good Practice on Dismissal (September 2025 Update)")
    P(doc, "The updated Code of Good Practice on Dismissal (Government Gazette, September 2025) introduced the following clarifications and updates:")
    B(doc, "Reinforcement of the requirement for employers to apply progressive discipline, reserving summary dismissal for cases of gross misconduct;")
    B(doc, "Expanded guidance on the investigation process, emphasizing the need for thorough and unbiased investigation before initiating formal charges;")
    B(doc, "Updated guidelines on the use of suspension (including precautionary suspension), requiring employers to review the necessity of continued suspension at regular intervals;")
    B(doc, "Clarification that dismissal for a first offence is only appropriate where the misconduct is so serious that it makes a continued employment relationship intolerable;")
    B(doc, "Enhanced guidance on the role of the chairperson in disciplinary hearings, including the expectation of independence and impartiality;")
    B(doc, "Emphasis on the duty of the employer to consider the employee's personal circumstances, length of service, and disciplinary record as mitigating factors before deciding on dismissal.")

    H2(doc, "11. Policy Review")
    P(doc, "This policy will be reviewed annually, or as required by amendments to the LRA, the Code of Good Practice, or decisions of the CCMA, Labour Court, or Labour Appeal Court.")

    add_approval_table(doc)
    add_acknowledgment(doc)
    add_footer_line(doc, "LRA Policy -- South Africa")

    path = os.path.join(POLICY_DIR, "Policy_LRA_SouthAfrica.docx")
    doc.save(path)
    return path


# ===================================================================
# 3. EEA - Employment Equity Act
# ===================================================================
def build_eea():
    doc = create_doc()
    add_title(doc)
    add_jurisdiction_tag(doc, "[SOUTH AFRICA] -- Employment Equity Act 55 of 1998")
    H1(doc, "Employment Equity Act (EEA) Policy")
    doc.add_paragraph()
    add_metadata_table(doc, "[POLICY-SA-003]")

    H2(doc, "1. Purpose")
    P(doc, "This policy sets out [COMPANY NAME]'s commitment to employment equity in accordance with the Employment Equity Act 55 of 1998 (EEA), as amended by the Employment Equity Amendment Act 4 of 2022 (effective 1 January 2025). The EEA aims to achieve equity in the workplace by promoting equal opportunity and fair treatment through the elimination of unfair discrimination, and by implementing affirmative action measures to redress the disadvantages in employment experienced by designated groups.")

    H2(doc, "2. Scope")
    P(doc, "This policy applies to all employees of [COMPANY NAME] in South Africa. Specific obligations under Chapters II and III of the EEA apply to designated employers.")

    H3(doc, "2.1 Designated Employers")
    P(doc, "A designated employer is an employer that:")
    B(doc, "Employs 50 or more employees; or")
    B(doc, "Has a total annual turnover that is equal to or above the applicable threshold prescribed in Schedule 4 of the EEA (varying by sector); or")
    B(doc, "Is a municipality, an organ of state, or is bound by a collective agreement to implement employment equity; or")
    B(doc, "Has been designated by the Minister of Employment and Labour.")

    H3(doc, "2.2 Designated Groups")
    P(doc, "Designated groups means Black people (including Africans, Coloureds, and Indians), women, and people with disabilities who are citizens of the Republic of South Africa by birth or descent, or who became citizens by naturalisation before the commencement of the Constitution of the Republic of South Africa, 1993.")

    H2(doc, "3. Prohibition of Unfair Discrimination")
    P(doc, "No person may unfairly discriminate, directly or indirectly, against an employee in any employment policy or practice on one or more grounds, including race, gender, sex, pregnancy, marital status, family responsibility, ethnic or social origin, colour, sexual orientation, age, disability, religion, HIV status, conscience, belief, political opinion, culture, language, birth, or on any other arbitrary ground.")
    P(doc, "Discrimination on a listed ground is presumed to be unfair unless the employer can demonstrate that the discrimination is fair.")

    H2(doc, "4. Employment Equity Plans")
    P(doc, "A designated employer must prepare and implement an employment equity plan (EE Plan). The EE Plan must include:")
    B(doc, "The objectives to be achieved for each year of the plan;")
    B(doc, "The affirmative action measures to be implemented;")
    B(doc, "The numerical goals to achieve equitable representation of suitably qualified people from designated groups in each occupational level;")
    B(doc, "The timetable for each year of the plan;")
    B(doc, "The duration of the plan (between three and five years);")
    B(doc, "The internal monitoring and evaluation procedures;")
    B(doc, "The internal dispute resolution mechanisms.")

    H3(doc, "4.1 Five-Year Cycle: 2025-2030")
    P(doc, "The current five-year employment equity cycle runs from 2025 to 2030. Employers must ensure their EE Plans are aligned with the sector-specific numerical targets published by the Minister of Employment and Labour (see Section 7 below).")

    H2(doc, "5. Affirmative Action Measures")
    P(doc, "Affirmative action measures are measures designed to ensure that suitably qualified people from designated groups have equal employment opportunities and are equitably represented in all occupational levels of the workforce. These measures include:")
    B(doc, "Identifying and eliminating employment barriers that adversely affect designated groups;")
    B(doc, "Measures to further diversity in the workplace;")
    B(doc, "Making reasonable accommodation for people from designated groups to ensure they enjoy equal opportunities;")
    B(doc, "Retention and development strategies for designated groups;")
    B(doc, "Preferential treatment and numerical goals (but not quotas) in recruitment, promotion, and training.")

    H2(doc, "6. Income Differentials and Equal Pay")
    P(doc, "Section 27 of the EEA requires designated employers to report on their remuneration and benefits received in each occupational level. The employer must take steps to progressively reduce disproportionate income differentials. The Employment Conditions Commission may investigate and make recommendations on income differentials.")
    P(doc, "Equal pay for work of equal value is a fundamental principle of the EEA. The Regulations on Equal Pay (2014) provide that a difference in terms and conditions of employment between employees of the same employer performing the same or substantially the same work, or work of equal value, that is directly or indirectly based on any one or more of the listed grounds of unfair discrimination, is unfair.")

    H2(doc, "7. Sector-Specific Numerical Targets (2025 Regulations)")
    P(doc, "The Employment Equity Amendment Act 4 of 2022 introduced the power of the Minister to set sector-specific numerical targets for designated groups. These regulations, effective from 2025, require:")
    B(doc, "Designated employers to align their EE Plans with sector-specific targets published by the Minister for their economic sector;")
    B(doc, "Targets are set based on the national and regional demographics and are intended to reflect the national economically active population (EAP);")
    B(doc, "Employers must demonstrate good-faith efforts to achieve the targets, including evidence of active recruitment, training, and promotion of designated group members;")
    B(doc, "Non-compliance may result in the employer's compliance certificate being refused, which affects the employer's ability to contract with state entities.")

    H2(doc, "8. EEA Reporting Requirements")
    P(doc, "Designated employers must submit an annual Employment Equity Report (EEA2 / EEA4) to the Department of Employment and Labour:")
    B(doc, "Employers with fewer than 150 employees: Report by 15 January of each year (or the first working day thereafter);")
    B(doc, "Employers with 150 or more employees: Report by 1 October of each year (or the first working day thereafter);")
    B(doc, "Reports must be submitted electronically via the Department's online portal;")
    B(doc, "The report must include the employer's workforce profile, income differentials, and progress against the EE Plan targets;")
    B(doc, "A summary of the EE Report must be displayed at the workplace for employee access.")

    H2(doc, "9. Penalties for Non-Compliance")
    P(doc, "The EEA prescribes significant penalties for non-compliance, which were substantially increased by the 2022 Amendment Act:")
    make_table(doc,
        ["Contravention", "Maximum Penalty"],
        [
            ["First contravention", "R1,500,000 or 2% of turnover (whichever is greater)"],
            ["Second contravention", "R1,800,000 or 4% of turnover"],
            ["Third contravention", "R2,100,000 or 6% of turnover"],
            ["Fourth and subsequent", "R2,700,000 or 8% of turnover"],
        ], bold_first_col=True)
    P(doc, "In addition to financial penalties, the Director-General may apply to the Labour Court for an order compelling the employer to comply, and the employer may be barred from contracting with any organ of state.")

    H2(doc, "10. Compliance Certificate")
    P(doc, "A designated employer must obtain a compliance certificate from the Minister of Employment and Labour in order to contract with any organ of state. The certificate confirms that the employer is in compliance with the EEA. An employer may be refused a certificate if it has not submitted an EE Report, has not prepared or implemented an EE Plan, or has not made good-faith efforts to achieve its numerical targets.")

    H2(doc, "11. Policy Review")
    P(doc, "This policy will be reviewed annually, or as required by amendments to the EEA, the publication of new sector-specific targets, or changes in the employer's designated employer status.")

    add_approval_table(doc)
    add_acknowledgment(doc)
    add_footer_line(doc, "EEA Policy -- South Africa")

    path = os.path.join(POLICY_DIR, "Policy_EEA_SouthAfrica.docx")
    doc.save(path)
    return path


# ===================================================================
# 4. OHSA - Occupational Health and Safety Act
# ===================================================================
def build_ohsa():
    doc = create_doc()
    add_title(doc)
    add_jurisdiction_tag(doc, "[SOUTH AFRICA] -- Occupational Health and Safety Act 85 of 1993")
    H1(doc, "Occupational Health and Safety Act (OHSA) Policy")
    doc.add_paragraph()
    add_metadata_table(doc, "[POLICY-SA-004]")

    H2(doc, "1. Purpose")
    P(doc, "This policy establishes [COMPANY NAME]'s commitment to providing and maintaining a safe and healthy working environment for all employees, contractors, and visitors in accordance with the Occupational Health and Safety Act 85 of 1993 (OHSA) and its regulations. The OHSA requires every employer to provide and maintain, as far as is reasonably practicable, a working environment that is safe and without risk to the health of employees.")

    H2(doc, "2. Scope")
    P(doc, "This policy applies to all workplaces, employees, contractors, subcontractors, and visitors of [COMPANY NAME] within the Republic of South Africa. It extends to all activities conducted on behalf of [COMPANY NAME], including work at client sites, project locations, and remote work environments where applicable.")

    H2(doc, "3. Employer Duties (Section 8)")
    P(doc, "In terms of Section 8 of the OHSA, [COMPANY NAME], as the employer, has a duty to:")
    B(doc, "Provide and maintain, as far as is reasonably practicable, a working environment that is safe and without risk to the health of employees;")
    B(doc, "Identify hazards and assess risks in the workplace and take steps to eliminate or mitigate them;")
    B(doc, "Establish and maintain systems of work that are safe and without risk to health;")
    B(doc, "Provide such information, instructions, training, and supervision as may be necessary to ensure the health and safety of employees;")
    B(doc, "Provide and maintain all necessary personal protective equipment (PPE) at no cost to employees;")
    B(doc, "Ensure that all equipment, machinery, and substances are used, handled, stored, and transported safely;")
    B(doc, "Not permit any employee to do any work or produce, process, use, handle, store, or transport any substance unless the necessary precautionary measures have been taken;")
    B(doc, "Enforce health and safety measures and take disciplinary action against employees who contravene safety rules.")

    H2(doc, "4. Employee Duties (Section 14)")
    P(doc, "In terms of Section 14 of the OHSA, every employee has a duty to:")
    B(doc, "Take reasonable care for the health and safety of themselves and of other persons who may be affected by their acts or omissions;")
    B(doc, "Cooperate with the employer in complying with the provisions of the OHSA;")
    B(doc, "Carry out any lawful order given and obey the health and safety rules and procedures of the employer;")
    B(doc, "Report any unsafe or unhealthy condition to the employer, the health and safety representative, or the health and safety committee as soon as possible;")
    B(doc, "Report any incident that may affect their health or that has caused an injury, to the employer as soon as possible;")
    B(doc, "Use and take proper care of any PPE or safety equipment provided.")

    H2(doc, "5. Health and Safety Representatives (Sections 17-18)")
    P(doc, "The OHSA requires the designation of health and safety representatives (HSRs) as follows:")
    B(doc, "Every employer with 20 or more employees must designate health and safety representatives in writing;")
    B(doc, "At least one HSR must be designated for every 50 employees (or as prescribed by the Minister);")
    B(doc, "HSRs must be full-time employees who are familiar with the conditions in the workplace;")
    B(doc, "HSRs have the right to inspect the workplace, investigate incidents, attend meetings of the health and safety committee, and make representations to the employer on health and safety matters;")
    B(doc, "The employer must not victimise, discipline, or dismiss an employee for performing functions as an HSR.")

    H2(doc, "6. Health and Safety Committees (Sections 19-20)")
    P(doc, "Where two or more HSRs have been designated, the employer must establish a health and safety committee for the workplace:")
    B(doc, "The committee must consist of HSRs and an equal number of employer-nominated members;")
    B(doc, "The committee must meet at least once every three months, or more frequently if required;")
    B(doc, "Functions include making recommendations to the employer on health and safety matters, investigating incidents, and reviewing the effectiveness of health and safety measures;")
    B(doc, "Minutes of committee meetings must be kept and made available to an inspector upon request.")

    H2(doc, "7. Incident Reporting and Investigation")
    P(doc, "The OHSA and the General Administrative Regulations require the reporting of certain incidents to the Department of Employment and Labour:")
    B(doc, "Any incident in which a person dies, becomes unconscious, loses a limb, or is otherwise seriously injured must be reported to a provincial inspector (Section 24);")
    B(doc, "The report must be made as soon as possible, but not later than the prescribed timeframe;")
    B(doc, "A full investigation must be conducted to determine the root cause and prevent recurrence;")
    B(doc, "Records of all incidents, investigations, and corrective actions must be maintained for at least three years;")
    B(doc, "The incident scene must not be disturbed until an inspector has authorized it, except to prevent further injury or damage.")

    H2(doc, "8. Construction Regulations")
    P(doc, "For construction work, the Construction Regulations (2014) under the OHSA impose additional duties:")
    B(doc, "A Health and Safety File must be prepared and maintained for each construction site;")
    B(doc, "A Health and Safety Plan must be submitted to the client before construction commences;")
    B(doc, "A competent person must be appointed as the Construction Health and Safety Agent;")
    B(doc, "Fall protection plans must be in place for work at heights exceeding 2 metres;")
    B(doc, "A risk assessment must be conducted before each phase of construction;")
    B(doc, "All contractors must be registered with the Compensation Fund or a licensed mutual association (see COIDA policy).")

    H2(doc, "9. Personal Protective Equipment (PPE)")
    P(doc, "In accordance with the General Safety Regulations and the PPE Regulations under the OHSA:")
    B(doc, "The employer must identify PPE requirements based on the risk assessment for each work activity;")
    B(doc, "PPE must be provided at no cost to the employee;")
    B(doc, "Employees must be trained in the correct use, care, and limitations of PPE;")
    B(doc, "PPE must be properly maintained, inspected, and replaced when damaged or worn;")
    B(doc, "Records of PPE issued and training provided must be maintained.")

    H2(doc, "10. Penalties")
    P(doc, "Non-compliance with the OHSA may result in severe penalties:")
    make_table(doc,
        ["Offence", "Penalty"],
        [
            ["Contravention of a provision of the OHSA", "Fine up to R100,000 and/or imprisonment up to 12 months"],
            ["Contravention resulting in death or serious injury", "Fine up to R1,000,000 and/or imprisonment up to 2 years"],
            ["Failure to comply with a prohibition notice", "Fine up to R200,000 and/or imprisonment up to 24 months"],
            ["Repeat offence within 12 months", "Double the prescribed penalty"],
        ], bold_first_col=True)
    P(doc, "In addition, the employer may face civil claims for damages from injured employees or their dependants, and the Department of Employment and Labour may issue compliance orders, prohibition notices, or improvement notices.")

    H2(doc, "11. Policy Review")
    P(doc, "This policy will be reviewed annually, or as required by amendments to the OHSA, the promulgation of new regulations, or significant changes in workplace conditions.")

    add_approval_table(doc)
    add_acknowledgment(doc)
    add_footer_line(doc, "OHSA Policy -- South Africa")

    path = os.path.join(POLICY_DIR, "Policy_OHSA_SouthAfrica.docx")
    doc.save(path)
    return path


# ===================================================================
# 5. COIDA - Compensation for Occupational Injuries and Diseases
# ===================================================================
def build_coida():
    doc = create_doc()
    add_title(doc)
    add_jurisdiction_tag(doc, "[SOUTH AFRICA] -- Compensation for Occupational Injuries and Diseases Act 130 of 1993")
    H1(doc, "Compensation for Occupational Injuries and Diseases Act (COIDA) Policy")
    doc.add_paragraph()
    add_metadata_table(doc, "[POLICY-SA-005]")

    H2(doc, "1. Purpose")
    P(doc, "This policy outlines [COMPANY NAME]'s obligations and procedures under the Compensation for Occupational Injuries and Diseases Act 130 of 1993 (COIDA), as amended. COIDA provides a no-fault compensation system for employees who are injured, contract diseases, or die as a result of workplace accidents or occupational diseases arising out of and in the course of their employment.")

    H2(doc, "2. Scope")
    P(doc, "This policy applies to all employees of [COMPANY NAME] in South Africa, including full-time, part-time, temporary, and casual employees. It does not apply to members of the South African National Defence Force, the South African Police Service, or domestic workers in private households (who are covered under separate arrangements).")

    H2(doc, "3. Employer Registration")
    P(doc, "Every employer in South Africa that employs one or more employees is required to register with the Compensation Fund (administered by the Department of Employment and Labour) or with a licensed mutual association (such as the Federated Employers' Mutual Assurance Company, the Rand Mutual Assurance Company, or others operating in specific sectors).")
    B(doc, "Registration must be completed within seven days of employing the first employee;")
    B(doc, "The employer must submit a return of earnings (W.As.8 form) annually, by 31 March, declaring the total earnings paid to all employees during the preceding year;")
    B(doc, "Failure to register is a criminal offence and exposes the employer to personal liability for compensation claims.")

    H2(doc, "4. Assessment and Premium Calculations")
    P(doc, "The Compensation Fund calculates the employer's annual assessment (premium) based on:")
    B(doc, "The industry classification and associated tariff rate for the employer's sector;")
    B(doc, "The total annual earnings of all employees (as declared in the return of earnings);")
    B(doc, "The employer's claims history (the merit/demerit rebate system rewards employers with low claims ratios and penalises those with high claims);")
    B(doc, "Assessments must be paid within 30 days of the date of the assessment notice. Late payment incurs a 10% penalty plus interest.")

    H2(doc, "5. Reporting Workplace Injuries")
    P(doc, "An employer must report a workplace injury or occupational disease to the Compensation Commissioner as follows:")
    B(doc, "Notification: The employer must report the accident within seven (7) days of the date on which the accident occurred or the date on which the employer became aware of it;", bold_prefix="Step 1 -- ")
    B(doc, "Form W.Cl.2: The employer must complete and submit the Employer's Report of an Accident (Form W.Cl.2) to the Compensation Commissioner or the relevant mutual association;", bold_prefix="Step 2 -- ")
    B(doc, "Medical Report: The first medical report (Form W.Cl.4) must be submitted by the attending medical practitioner within 14 days of the first consultation;", bold_prefix="Step 3 -- ")
    B(doc, "Progress Reports: Subsequent progress and final medical reports must be submitted as required by the Compensation Commissioner;", bold_prefix="Step 4 -- ")
    B(doc, "Death: If the employee dies as a result of the accident, the employer must report the death immediately.")
    P(doc, "Failure to report within the prescribed period is an offence. The employer may also be held liable for any additional costs incurred as a result of the late report.")

    H2(doc, "6. Claims Process")
    P(doc, "The claims process under COIDA is as follows:")
    B(doc, "The employee sustains an injury or is diagnosed with an occupational disease arising out of and in the course of employment;")
    B(doc, "The employee reports the injury or disease to the employer as soon as possible;")
    B(doc, "The employer completes and submits Form W.Cl.2 to the Compensation Commissioner within seven days;")
    B(doc, "The employee's medical practitioner submits the medical report (Form W.Cl.4);")
    B(doc, "The Compensation Commissioner assesses the claim and determines the level of disability (if any);")
    B(doc, "The employee receives compensation in the form of medical expenses, temporary disability benefits, permanent disability benefits, or dependants' benefits (in the case of death).")
    P(doc, "The employee does not need to prove negligence on the part of the employer. In exchange for the no-fault compensation system, the employee is barred from suing the employer for damages arising from the workplace injury (the 'COIDA bar').")

    H2(doc, "7. Temporary and Permanent Disability Benefits")

    H3(doc, "7.1 Temporary Disability")
    P(doc, "An employee who is temporarily disabled as a result of a workplace injury is entitled to:")
    B(doc, "Compensation for the period of temporary disability, calculated as 75% of the employee's monthly earnings (subject to the maximum earnings ceiling prescribed by the Minister);")
    B(doc, "The employer must pay the employee for the first three months of disability and is then reimbursed by the Compensation Fund;")
    B(doc, "The employee is entitled to all reasonable and necessary medical treatment at the expense of the Compensation Fund.")

    H3(doc, "7.2 Permanent Disability")
    P(doc, "If the employee suffers a permanent disability, the level of disability is assessed as a percentage by the Compensation Commissioner based on the medical evidence:")
    B(doc, "Permanent disability of less than 30%: The employee receives a once-off lump sum payment;")
    B(doc, "Permanent disability of 30% or more: The employee receives a monthly pension for life (calculated as a percentage of the employee's earnings, subject to the maximum earnings ceiling);")
    B(doc, "100% permanent disability: The employee receives a full pension equivalent to 75% of their monthly earnings.")

    H3(doc, "7.3 Death Benefits")
    P(doc, "If an employee dies as a result of a workplace injury or occupational disease, the dependants of the employee are entitled to a monthly pension and/or a lump sum payment, as determined by the Compensation Commissioner.")

    H2(doc, "8. 2026 Amendments")
    P(doc, "The Compensation for Occupational Injuries and Diseases Amendment Act (2026) introduced the following key changes:")
    B(doc, "Updated compliance duties: Employers must now maintain digital records of all workplace injuries and submit reports electronically through the CompEasy online system;")
    B(doc, "Administrative penalties: The Commissioner may now impose administrative penalties (up to R500,000) for non-compliance with reporting obligations, without the need for criminal prosecution;")
    B(doc, "Expanded coverage: Clarification that employees working from home or remotely are covered by COIDA if the injury arises out of and in the course of their employment;")
    B(doc, "Faster adjudication: The establishment of a Compensation Court to expedite the resolution of disputed claims;")
    B(doc, "Enhanced medical provider management: The Commissioner may prescribe approved medical service providers and fee schedules to control costs.")

    H2(doc, "9. Employer Obligations Summary")
    make_table(doc,
        ["Obligation", "Deadline / Requirement"],
        [
            ["Register with Compensation Fund", "Within 7 days of first employee"],
            ["Submit Return of Earnings (W.As.8)", "By 31 March annually"],
            ["Pay annual assessment", "Within 30 days of notice"],
            ["Report workplace injury (W.Cl.2)", "Within 7 days of accident"],
            ["Report death", "Immediately"],
            ["Maintain injury records", "For at least 4 years"],
            ["Display COIDA poster", "At all times in the workplace"],
        ], bold_first_col=True)

    H2(doc, "10. Policy Review")
    P(doc, "This policy will be reviewed annually, or as required by amendments to COIDA, changes in the assessment tariff, or directives from the Compensation Commissioner.")

    add_approval_table(doc)
    add_acknowledgment(doc)
    add_footer_line(doc, "COIDA Policy -- South Africa")

    path = os.path.join(POLICY_DIR, "Policy_COIDA_SouthAfrica.docx")
    doc.save(path)
    return path


# ===================================================================
# 6. UIF / PAYE / SDL - Statutory Deductions
# ===================================================================
def build_uif_paye_sdl():
    doc = create_doc()
    add_title(doc)
    add_jurisdiction_tag(doc, "[SOUTH AFRICA] -- Statutory Deductions: UIF, PAYE, SDL, COIDA")
    H1(doc, "Statutory Deductions Policy: UIF, PAYE, SDL, and COIDA Contributions")
    doc.add_paragraph()
    add_metadata_table(doc, "[POLICY-SA-006]")

    H2(doc, "1. Purpose")
    P(doc, "This policy outlines [COMPANY NAME]'s obligations regarding statutory deductions and contributions required by South African law. Employers are required to deduct certain amounts from employee remuneration and to make additional employer contributions, and to remit these to the relevant government authorities within prescribed timeframes. Non-compliance may result in penalties, interest charges, and criminal prosecution.")

    H2(doc, "2. Scope")
    P(doc, "This policy applies to [COMPANY NAME] and all employees employed in the Republic of South Africa. The payroll department is responsible for the accurate calculation, deduction, and remittance of all statutory amounts.")

    H2(doc, "3. Unemployment Insurance Fund (UIF)")
    P(doc, "The Unemployment Insurance Act 63 of 2001 and the Unemployment Insurance Contributions Act 4 of 2002 require:")

    H3(doc, "3.1 Contributions")
    B(doc, "Employee contribution: 1% of the employee's remuneration (deducted from salary);")
    B(doc, "Employer contribution: 1% of the employee's remuneration (paid by the employer);")
    B(doc, "Total UIF contribution: 2% of remuneration (1% employee + 1% employer);")
    B(doc, "Remuneration is subject to the UIF earnings ceiling (currently R17,712 per month; maximum contribution per employee is R177.12 per month);")
    B(doc, "Employees who work less than 24 hours per month for an employer, and learners on learnership agreements, are excluded from UIF contributions.")

    H3(doc, "3.2 Filing and Payment")
    B(doc, "UIF contributions must be paid to SARS monthly, by the 7th day of the month following the month in which the remuneration was paid;")
    B(doc, "Monthly declarations are submitted via the EMP201 return;")
    B(doc, "The annual UIF reconciliation is submitted with the EMP501 annual reconciliation by 31 May each year.")

    H3(doc, "3.3 UIF Benefits")
    P(doc, "Employees who have contributed to the UIF may claim the following benefits: unemployment benefits (up to 238 days), illness benefits, maternity benefits, adoption benefits, parental leave benefits, and dependant's benefits (in the event of the contributor's death).")

    H2(doc, "4. Pay As You Earn (PAYE)")
    P(doc, "The Income Tax Act 58 of 1962 requires employers to deduct income tax (PAYE) from employees' remuneration and remit it to the South African Revenue Service (SARS).")

    H3(doc, "4.1 Obligations")
    B(doc, "The employer must register as an employer with SARS within 21 business days of becoming an employer;")
    B(doc, "PAYE must be deducted from all remuneration paid to employees, based on the applicable tax tables issued by SARS;")
    B(doc, "PAYE is calculated on a cumulative basis over the tax year (1 March to 28/29 February);")
    B(doc, "Tax deductions must take into account the employee's tax rebates (primary, secondary, tertiary), medical tax credits, and any approved deductions.")

    H3(doc, "4.2 Filing and Payment")
    B(doc, "Monthly PAYE must be paid to SARS by the 7th day of the month following the month in which the remuneration was paid (EMP201 return);")
    B(doc, "The bi-annual reconciliation (EMP501) must be submitted by September (interim) and by May (annual);")
    B(doc, "IRP5/IT3(a) certificates must be issued to each employee annually, reflecting total remuneration and tax deducted;")
    B(doc, "Late payment or non-payment results in a 10% penalty on the outstanding amount plus interest at the prescribed rate.")

    H2(doc, "5. Skills Development Levy (SDL)")
    P(doc, "The Skills Development Levies Act 9 of 1999 requires:")

    H3(doc, "5.1 Obligations")
    B(doc, "Every employer whose total annual payroll exceeds R500,000 must pay the SDL;")
    B(doc, "The SDL is calculated at 1% of the total leviable amount (total remuneration paid to all employees);")
    B(doc, "The SDL is an employer cost and may not be deducted from employee remuneration;")
    B(doc, "National and provincial government departments, public entities listed in Schedule 1 and Part A of Schedule 3 of the PFMA, municipalities, and certain exempt employers are not required to pay the SDL.")

    H3(doc, "5.2 Filing and Payment")
    B(doc, "The SDL must be paid to SARS monthly, together with UIF and PAYE, via the EMP201 return;")
    B(doc, "Payment is due by the 7th day of the month following the month in which the remuneration was paid;")
    B(doc, "Late payment results in a 10% penalty on the outstanding amount plus interest.")

    H3(doc, "5.3 SETA Grants")
    P(doc, "Employers who pay the SDL may claim grants from their relevant Sector Education and Training Authority (SETA):")
    B(doc, "Mandatory grant: 20% of the SDL paid, subject to the employer submitting a Workplace Skills Plan (WSP) and Annual Training Report (ATR) by 30 April each year;")
    B(doc, "Discretionary grant: Allocated by the SETA for specific training and skills development projects.")

    H2(doc, "6. COIDA Assessment Contributions")
    P(doc, "In addition to the requirements detailed in the COIDA Policy (Policy SA-005), the employer must:")
    B(doc, "Pay the annual COIDA assessment to the Compensation Fund (or licensed mutual association) within 30 days of the assessment notice;")
    B(doc, "Submit the annual Return of Earnings by 31 March each year;")
    B(doc, "The COIDA assessment is an employer cost and may not be deducted from employee remuneration.")

    H2(doc, "7. Summary of Statutory Deductions and Contributions")
    make_table(doc,
        ["Item", "Rate", "Who Pays", "Paid To", "Deadline"],
        [
            ["UIF (employee)", "1% of remuneration", "Employee (deducted)", "SARS", "7th of following month"],
            ["UIF (employer)", "1% of remuneration", "Employer", "SARS", "7th of following month"],
            ["PAYE", "Per tax tables", "Employee (deducted)", "SARS", "7th of following month"],
            ["SDL", "1% of payroll", "Employer", "SARS", "7th of following month"],
            ["COIDA assessment", "Per sector tariff", "Employer", "Compensation Fund", "Within 30 days of notice"],
        ], bold_first_col=True)

    H2(doc, "8. Penalties for Non-Compliance")
    make_table(doc,
        ["Offence", "Penalty"],
        [
            ["Late payment of PAYE/UIF/SDL", "10% penalty + interest at prescribed rate"],
            ["Failure to register as employer", "Criminal prosecution, fines"],
            ["Late filing of EMP201/EMP501", "Administrative penalties per SARS"],
            ["Failure to issue IRP5 certificates", "R250 per day per certificate"],
            ["Late COIDA Return of Earnings", "10% penalty + interest"],
            ["Failure to register with Compensation Fund", "Criminal offence, personal liability"],
        ], bold_first_col=True)

    H2(doc, "9. Policy Review")
    P(doc, "This policy will be reviewed annually, or upon changes to tax tables, contribution rates, earnings ceilings, or other legislative amendments.")

    add_approval_table(doc)
    add_acknowledgment(doc)
    add_footer_line(doc, "Statutory Deductions Policy -- South Africa")

    path = os.path.join(POLICY_DIR, "Policy_UIF_PAYE_SDL_SouthAfrica.docx")
    doc.save(path)
    return path


# ===================================================================
# 7. Onboarding - South Africa
# ===================================================================
def build_onboarding_sa():
    doc = create_doc()
    add_title(doc)
    add_jurisdiction_tag(doc, "[SOUTH AFRICA] -- New Employee Onboarding Playbook")
    H1(doc, "South Africa Onboarding Playbook")
    doc.add_paragraph()
    add_metadata_table(doc, "[ONBOARD-SA-001]")

    H2(doc, "1. Purpose")
    P(doc, "This playbook provides a structured, legally compliant onboarding process for new employees joining [COMPANY NAME] in South Africa. It ensures all statutory, regulatory, and company-specific requirements are met, and provides a positive first experience for the new employee. Compliance items are mapped to the relevant legislation (BCEA, LRA, OHSA, COIDA, UIF, EEA).")

    H2(doc, "2. Pre-Start (Before Day 1)")
    P(doc, "The following must be completed before the employee's first day of work:")

    H3(doc, "2.1 Documentation Collection")
    B(doc, "Certified copy of South African Identity Document (ID) or valid work permit/visa (for foreign nationals);")
    B(doc, "IRP5 tax reference number (if available from previous employer) or completion of the IT77 form to register with SARS;")
    B(doc, "Bank account details (for salary payment via EFT -- bank name, branch code, account number);")
    B(doc, "Tax number (or application for a tax number if the employee does not have one);")
    B(doc, "Proof of residential address (not older than 3 months);")
    B(doc, "Highest qualification certificates (certified copies);")
    B(doc, "Criminal record check (where applicable for the role);")
    B(doc, "Reference check completion and clearance.")

    H3(doc, "2.2 Employment Contract")
    B(doc, "The written contract of employment must be provided to the employee before or on the first day of work, as required by Section 29 of the BCEA;")
    B(doc, "The contract must contain the prescribed particulars: full name of employer and employee, place of work, job title and description, date of commencement, ordinary hours of work, remuneration, leave entitlements, notice period, and any applicable Bargaining Council agreement;")
    B(doc, "The employee must sign and return the contract. A copy must be retained by the employer and a copy provided to the employee.")

    H3(doc, "2.3 System Setup")
    B(doc, "Create employee record in HR/payroll system;")
    B(doc, "Register employee for UIF (via next EMP201 submission to SARS);")
    B(doc, "Set up access credentials, email, and IT equipment;")
    B(doc, "Assign workspace, tools, and PPE (if applicable).")

    H2(doc, "3. Day 1 -- Orientation and Induction")

    H3(doc, "3.1 OHSA Safety Induction (Mandatory Before Work Starts)")
    P_bold(doc, "No employee may commence any work before completing the OHSA safety induction.")
    B(doc, "General workplace safety rules and emergency procedures;")
    B(doc, "Location of fire extinguishers, first aid kits, and emergency exits;")
    B(doc, "Identification of health and safety representative(s) and committee;")
    B(doc, "Reporting procedures for accidents, incidents, and near-misses;")
    B(doc, "PPE requirements for the employee's role (if applicable);")
    B(doc, "Signed acknowledgment of safety induction must be retained on file.")

    H3(doc, "3.2 Administrative Onboarding")
    B(doc, "Welcome and introductions to team, manager, and HR;")
    B(doc, "Office tour (physical or virtual);")
    B(doc, "Review and acknowledgment of key policies: Code of Conduct, Disciplinary Code, Leave Policy, Health and Safety Policy;")
    B(doc, "IT setup and systems orientation;")
    B(doc, "Issue employee handbook (hard copy or electronic).")

    H2(doc, "4. Week 1 -- Statutory Registrations and Benefits Enrolment")

    H3(doc, "4.1 UIF Registration")
    P(doc, "Confirm that the employee has been included in the UIF contribution records for the next EMP201 submission. Ensure the employee's ID number and banking details are correctly captured.")

    H3(doc, "4.2 Provident / Pension Fund Enrolment")
    B(doc, "If [COMPANY NAME] offers a provident fund or pension fund, enrol the employee within the prescribed period (typically from the first day of employment or after a probationary period, as specified in the fund rules);")
    B(doc, "Provide the employee with a copy of the fund rules, the member guide, and the nomination of beneficiaries form;")
    B(doc, "Explain the employee and employer contribution rates, vesting periods, and withdrawal rules.")

    H3(doc, "4.3 Medical Aid Enrolment")
    B(doc, "If [COMPANY NAME] offers group medical aid, provide the employee with scheme options, costs, and enrolment forms;")
    B(doc, "Membership typically commences from the first day of employment or within 30 days (per scheme rules);")
    B(doc, "Explain the employer subsidy (if applicable) and the employee's co-payment obligations.")

    H3(doc, "4.4 EEA Self-Declaration Form")
    P(doc, "In accordance with the Employment Equity Act, the employee should complete a self-declaration form indicating their race, gender, and disability status (if any). This information is used for EE reporting purposes and is voluntary. The form must be kept confidential and used solely for employment equity planning and reporting.")

    H2(doc, "5. Probation Period")
    P(doc, "Probation at [COMPANY NAME] is typically three to six months, as specified in the employment contract. In accordance with the LRA and the Code of Good Practice on Dismissal:")
    B(doc, "Probation is an extension of the selection process, not a licence to dismiss without process;")
    B(doc, "The employer must set clear performance standards at the outset and communicate them to the employee;")
    B(doc, "The employer must provide appropriate training, instruction, evaluation, and guidance during probation;")
    B(doc, "Regular performance discussions should take place during probation;")
    B(doc, "If the employee's performance is not satisfactory, the employer must counsel the employee and allow a reasonable opportunity to improve;")
    B(doc, "Dismissal during probation requires substantive and procedural fairness (although the standard of fairness may be less rigid than for a permanent employee);")
    B(doc, "The employee has the right to refer an unfair dismissal dispute to the CCMA even during probation.")

    H2(doc, "6. 30 / 60 / 90-Day Check-ins")
    make_table(doc,
        ["Milestone", "Focus Areas", "Actions"],
        [
            ["30 Days", "Settling in, initial training progress, cultural fit",
             "Manager-employee check-in; review training completion; address concerns; verify payroll accuracy"],
            ["60 Days", "Performance progress, deepening of role knowledge",
             "Formal mid-probation review; documented performance feedback; identify development needs"],
            ["90 Days", "Probation assessment, confirmation or extension",
             "End-of-probation review; decision to confirm, extend, or terminate; documented outcome"],
        ], bold_first_col=True)

    H2(doc, "7. Onboarding Compliance Checklist")
    make_table(doc,
        ["Item", "Legislation", "Status"],
        [
            ["Written contract of employment provided", "BCEA s.29", "[  ]"],
            ["OHSA safety induction completed (Day 1)", "OHSA s.8", "[  ]"],
            ["Certified ID copy on file", "BCEA / Company policy", "[  ]"],
            ["IRP5 / tax number captured", "Income Tax Act", "[  ]"],
            ["Bank details captured for EFT salary", "Company policy", "[  ]"],
            ["UIF registration confirmed", "Unemployment Insurance Act", "[  ]"],
            ["Provident/pension fund enrolment", "Fund rules", "[  ]"],
            ["Medical aid enrolment", "Scheme rules", "[  ]"],
            ["EEA self-declaration form completed", "EEA", "[  ]"],
            ["PPE issued (if applicable)", "OHSA General Safety Regs", "[  ]"],
            ["Employee handbook acknowledged", "Company policy", "[  ]"],
            ["Disciplinary code acknowledged", "LRA / Company policy", "[  ]"],
            ["30-day check-in completed", "Best practice", "[  ]"],
            ["60-day check-in completed", "Best practice", "[  ]"],
            ["90-day probation review completed", "LRA Code of Good Practice", "[  ]"],
        ], bold_first_col=True)

    H2(doc, "8. Policy Review")
    P(doc, "This onboarding playbook will be reviewed annually, or as required by changes in legislation or company policy.")

    add_footer_line(doc, "Onboarding Playbook -- South Africa")

    path = os.path.join(ONBOARD_DIR, "Onboarding_SouthAfrica.docx")
    doc.save(path)
    return path


# ===================================================================
# 8. Training - SA Labour Law for Managers
# ===================================================================
def build_training_sa():
    doc = create_doc()
    add_title(doc)
    add_jurisdiction_tag(doc, "[SOUTH AFRICA] -- Labour Law Training for Managers")
    H1(doc, "South African Labour Law Training for Managers")
    doc.add_paragraph()
    add_metadata_table(doc, "[TRAIN-SA-001]")

    H2(doc, "1. Purpose")
    P(doc, "This training guide equips line managers and supervisors at [COMPANY NAME] with the essential knowledge of South African labour law needed to manage their teams lawfully and fairly. Non-compliance with labour legislation exposes both the company and individual managers to significant legal and financial risk, including CCMA awards, Labour Court orders, and reputational damage.")

    H2(doc, "2. Audience")
    P(doc, "This training is mandatory for all line managers, supervisors, team leaders, and any employee with authority to hire, discipline, manage performance, or terminate employment at [COMPANY NAME] South Africa.")

    H2(doc, "3. Module 1: BCEA Essentials for Line Managers")
    P(doc, "The Basic Conditions of Employment Act (BCEA) sets the minimum conditions of employment. Key points for managers:")

    H3(doc, "3.1 Working Hours")
    B(doc, "Maximum 45 ordinary hours per week (9 hours/day for 5-day week, 8 hours/day for 6-day week);")
    B(doc, "Overtime requires employee agreement and is limited to 10 hours per week;")
    B(doc, "Overtime rate: 1.5x normal rate (2x for Sundays and public holidays);")
    B(doc, "Employees earning above the BCEA threshold are excluded from hours-of-work provisions.")

    H3(doc, "3.2 Leave Management")
    B(doc, "Annual leave: 21 consecutive days (15 working days) per cycle;")
    B(doc, "Sick leave: 36 days per 36-month cycle (for 6-day workers; 30 days for 5-day workers);")
    B(doc, "Medical certificate required for absences exceeding 2 consecutive days;")
    B(doc, "Family responsibility leave: 3 days per year (birth, illness, or death of family member);")
    B(doc, "Maternity leave: 4 months (UIF claimable);")
    B(doc, "Parental leave: 10 days per parent.")

    H3(doc, "3.3 Common Manager Mistakes")
    B(doc, "Denying leave without a valid reason;")
    B(doc, "Requiring overtime without agreement;")
    B(doc, "Failing to record working hours accurately;")
    B(doc, "Paying incorrect overtime rates;")
    B(doc, "Allowing employees to work through meal intervals without agreement.")

    H2(doc, "4. Module 2: Disciplinary Hearing Procedure (Step-by-Step)")
    P(doc, "This is the most critical module. Failure to follow a fair procedure results in unfair dismissal findings at the CCMA.")

    H3(doc, "4.1 Step 1: Investigation")
    B(doc, "Investigate the alleged misconduct before initiating formal charges;")
    B(doc, "Gather evidence, interview witnesses, and document findings;")
    B(doc, "Determine whether the matter warrants informal counselling or a formal hearing;")
    B(doc, "Do not pre-judge the outcome. The investigation must be unbiased.")

    H3(doc, "4.2 Step 2: Charge Sheet and Notice of Hearing")
    B(doc, "Draft a charge sheet specifying the charges clearly and in detail;")
    B(doc, "Issue a written notice to the employee at least 48 hours before the hearing (minimum);")
    B(doc, "The notice must state: the date, time, and venue of the hearing; the charges; the employee's right to representation by a fellow employee or trade union representative; and the possible outcomes (including dismissal, if applicable).")

    H3(doc, "4.3 Step 3: Conducting the Hearing")
    B(doc, "Appoint an impartial chairperson (not the direct supervisor or the complainant);")
    B(doc, "The initiator (employer) presents the case, calls witnesses, and presents evidence;")
    B(doc, "The employee (or representative) may cross-examine, present a defence, call witnesses, and submit evidence;")
    B(doc, "The chairperson must allow the employee a full and fair opportunity to respond;")
    B(doc, "Record the proceedings (written minutes at minimum).")

    H3(doc, "4.4 Step 4: Finding and Sanction")
    B(doc, "The chairperson makes a finding of guilty or not guilty on each charge;")
    B(doc, "If guilty, the chairperson considers aggravating and mitigating factors (nature of the offence, employee's service record, remorse, consistency of treatment);")
    B(doc, "The chairperson determines the appropriate sanction (verbal warning, written warning, final written warning, or dismissal);")
    B(doc, "The decision must be communicated in writing, with reasons.")

    H3(doc, "4.5 Step 5: Appeal and Post-Hearing")
    B(doc, "Inform the employee of the right to an internal appeal (typically within 5 working days);")
    B(doc, "Inform the employee of the right to refer the matter to the CCMA within 30 calendar days;")
    B(doc, "File all hearing documentation in the employee's personnel file.")

    H2(doc, "5. Module 3: How to Handle Poor Performance (Incapacity Process)")
    P(doc, "Poor performance is treated as incapacity, not misconduct. The process is counselling-based, not punitive.")

    H3(doc, "5.1 Step 1: Identify and Communicate")
    B(doc, "Clearly identify the performance shortfall with specific, measurable examples;")
    B(doc, "Meet with the employee to discuss the gap between expected and actual performance;")
    B(doc, "Ensure the employee understands the required standard.")

    H3(doc, "5.2 Step 2: Support and Develop")
    B(doc, "Provide training, coaching, mentoring, or other support to help the employee improve;")
    B(doc, "Set a reasonable timeframe for improvement (typically 4-12 weeks depending on the role);")
    B(doc, "Document the performance improvement plan (PIP) with clear goals and deadlines.")

    H3(doc, "5.3 Step 3: Evaluate")
    B(doc, "At the end of the PIP period, assess whether the employee has met the required standards;")
    B(doc, "If improved, confirm in writing and continue to monitor;")
    B(doc, "If not improved, conduct a formal incapacity hearing.")

    H3(doc, "5.4 Step 4: Incapacity Hearing")
    B(doc, "The hearing must be fair (similar procedural requirements to a disciplinary hearing);")
    B(doc, "Consider whether alternative positions are available before dismissal;")
    B(doc, "Dismissal is a last resort, after all reasonable steps have been exhausted.")

    H2(doc, "6. Module 4: CCMA Referral Process and Timelines")
    P(doc, "Managers must understand the CCMA process, as they may be required to represent the company.")
    make_table(doc,
        ["Stage", "Timeline", "Key Points"],
        [
            ["Employee referral", "Within 30 days of dismissal", "Employee files at CCMA (or Bargaining Council)"],
            ["Conciliation", "30 days from referral", "Facilitated negotiation; aim for settlement"],
            ["Certificate of non-resolution", "If conciliation fails", "Issued by CCMA commissioner"],
            ["Arbitration", "Within 90 days of certificate", "Formal hearing; CCMA makes binding award"],
            ["Labour Court review", "Within 6 weeks of award", "Only on grounds of gross irregularity"],
        ], bold_first_col=True)

    H2(doc, "7. Module 5: Common Dismissal Pitfalls")
    P(doc, "The following are the most common reasons employers lose at the CCMA:")
    B(doc, "No hearing before dismissal (procedural unfairness);", bold_prefix="Pitfall 1: ")
    B(doc, "Inadequate notice of hearing or vague charges;", bold_prefix="Pitfall 2: ")
    B(doc, "Chairperson was biased or had a conflict of interest;", bold_prefix="Pitfall 3: ")
    B(doc, "Employee not given the opportunity to respond or call witnesses;", bold_prefix="Pitfall 4: ")
    B(doc, "Sanction was too harsh (dismissal for a first minor offence);", bold_prefix="Pitfall 5: ")
    B(doc, "Inconsistent application of discipline (similar offence, different sanction);", bold_prefix="Pitfall 6: ")
    B(doc, "Dismissal for poor performance without a PIP;", bold_prefix="Pitfall 7: ")
    B(doc, "Retrenchment without proper Section 189 consultation;", bold_prefix="Pitfall 8: ")
    B(doc, "Automatically unfair dismissal (pregnancy, union activity, whistleblowing);", bold_prefix="Pitfall 9: ")
    B(doc, "Failure to keep records (no hearing minutes, no warning letters on file).", bold_prefix="Pitfall 10: ")

    H2(doc, "8. Module 6: Record-Keeping Requirements")
    P(doc, "Proper record-keeping is essential for defending the company at the CCMA and Labour Court. Managers must maintain:")
    B(doc, "Signed employment contracts for all employees;")
    B(doc, "Leave records (applications, approvals, balances);")
    B(doc, "Time and attendance records;")
    B(doc, "Performance records (KPIs, reviews, PIPs);")
    B(doc, "Disciplinary records (charge sheets, hearing minutes, warning letters, outcome letters);")
    B(doc, "Training records (OHSA induction, skills training, EE development);")
    B(doc, "Payroll records (payslips, EMP201 submissions, IRP5 certificates);")
    B(doc, "OHSA incident reports and investigation records.")
    P(doc, "Records must be retained for at least three years after termination of employment (five years for tax records).")

    H2(doc, "9. Assessment")
    P(doc, "All managers are required to complete the South African Labour Law assessment within 30 days of completing this training. A minimum score of 80% is required. Managers who do not achieve the minimum score must retake the training and assessment.")

    add_footer_line(doc, "SA Labour Law Training for Managers")

    path = os.path.join(TRAIN_DIR, "Training_SALabourLaw.docx")
    doc.save(path)
    return path


# ===================================================================
# 9. SOP Pack - South Africa Compliance
# ===================================================================
def build_sop_sa():
    doc = create_doc()
    add_title(doc)
    add_jurisdiction_tag(doc, "[SOUTH AFRICA] -- Compliance Standard Operating Procedures")
    H1(doc, "South Africa Compliance SOP Pack")
    doc.add_paragraph()
    add_metadata_table(doc, "[SOP-SA-001]")

    P(doc, "This document contains five standard operating procedures (SOPs) for South African compliance. Each SOP provides step-by-step instructions for critical HR and legal compliance processes.")

    # ----- SOP 1: OHSA Incident Reporting -----
    page_break(doc)
    H1(doc, "SOP 1: OHSA Incident Reporting and Investigation")
    make_table(doc,
        ["Field", "Details"],
        [
            ["SOP Number", "SOP-SA-001-01"],
            ["Version", "1.0"],
            ["Department", "Health & Safety / Human Resources"],
            ["Effective Date", "[YYYY-MM-DD]"],
            ["Review Date", "[YYYY-MM-DD]"],
            ["Prepared By", "[Name, Title]"],
            ["Approved By", "[Name, Title]"],
        ], bold_first_col=True)

    H2(doc, "1.1 Purpose")
    P(doc, "To establish a standardised procedure for reporting and investigating workplace incidents (accidents, injuries, diseases, near-misses, and dangerous occurrences) in compliance with the Occupational Health and Safety Act 85 of 1993 (OHSA) and the General Administrative Regulations.")

    H2(doc, "1.2 Scope")
    P(doc, "This SOP applies to all incidents occurring on [COMPANY NAME] premises, at project/client sites, during work-related travel, and any location where work is being performed on behalf of [COMPANY NAME].")

    H2(doc, "1.3 Procedure")
    P_bold(doc, "Step 1: Immediate Response (0-30 minutes)")
    B(doc, "Render first aid to the injured person(s) and call emergency services if required;")
    B(doc, "Secure the scene to prevent further injuries;")
    B(doc, "Do NOT disturb the scene unless necessary to prevent further harm (OHSA Section 24(3));")
    B(doc, "Notify the line manager and the Health and Safety Representative immediately.")

    P_bold(doc, "Step 2: Initial Report (Within 24 hours)")
    B(doc, "Complete the internal Incident Report Form (capturing date, time, location, description, witnesses, injuries);")
    B(doc, "Take photographs of the scene, equipment, and any contributing factors;")
    B(doc, "Record witness statements.")

    P_bold(doc, "Step 3: Statutory Notification (Section 24 -- as applicable)")
    B(doc, "If the incident involves death, unconsciousness, loss of a limb, or serious injury: notify the Provincial Director of the Department of Employment and Labour immediately (telephonically) and follow up in writing within 24 hours;")
    B(doc, "Complete Form W.Cl.2 (COIDA Employer's Report of an Accident) and submit to the Compensation Commissioner within 7 days (see SOP 2).")

    P_bold(doc, "Step 4: Investigation (Within 48 hours)")
    B(doc, "Appoint an investigation team (line manager + H&S representative + additional members as needed);")
    B(doc, "Conduct root cause analysis (using 5 Whys, Fishbone/Ishikawa, or similar methodology);")
    B(doc, "Identify contributing factors (human, environmental, procedural, equipment);")
    B(doc, "Document findings in the Investigation Report.")

    P_bold(doc, "Step 5: Corrective Actions (Within 7 days)")
    B(doc, "Implement immediate corrective actions to prevent recurrence;")
    B(doc, "Assign responsibility and deadlines for longer-term corrective actions;")
    B(doc, "Update risk assessments and safe work procedures as required;")
    B(doc, "Communicate findings and corrective actions to all affected employees.")

    P_bold(doc, "Step 6: Record Retention")
    B(doc, "Retain all incident reports, investigation reports, and corrective action records for a minimum of 3 years;")
    B(doc, "Records must be available for inspection by a Department of Employment and Labour inspector.")

    # ----- SOP 2: COIDA Injury Claim Filing -----
    page_break(doc)
    H1(doc, "SOP 2: COIDA Injury Claim Filing")
    make_table(doc,
        ["Field", "Details"],
        [
            ["SOP Number", "SOP-SA-001-02"],
            ["Version", "1.0"],
            ["Department", "Human Resources / Finance"],
            ["Effective Date", "[YYYY-MM-DD]"],
            ["Review Date", "[YYYY-MM-DD]"],
        ], bold_first_col=True)

    H2(doc, "2.1 Purpose")
    P(doc, "To establish a standardised procedure for filing compensation claims under the Compensation for Occupational Injuries and Diseases Act 130 of 1993 (COIDA) when an employee sustains a workplace injury or is diagnosed with an occupational disease.")

    H2(doc, "2.2 Procedure")
    P_bold(doc, "Step 1: Employee Reports Injury (Day 0)")
    B(doc, "Employee reports the injury or disease to their line manager as soon as possible after the incident;")
    B(doc, "The line manager records the details and ensures the employee receives medical attention.")

    P_bold(doc, "Step 2: Employer Completes Form W.Cl.2 (Within 7 days)")
    B(doc, "HR completes the Employer's Report of an Accident (Form W.Cl.2);")
    B(doc, "The form must include: employee details, date/time/place of accident, nature of injury, description of how the accident occurred, witness details, and the employer's assessment number;")
    B(doc, "Submit the completed form to the Compensation Commissioner (or the relevant mutual association).")

    P_bold(doc, "Step 3: Medical Practitioner Submits First Medical Report (Within 14 days)")
    B(doc, "The attending medical practitioner completes Form W.Cl.4 (First Medical Report);")
    B(doc, "HR follows up with the medical practitioner to ensure timely submission;")
    B(doc, "Ensure the employee attends all follow-up appointments.")

    P_bold(doc, "Step 4: Employer Pays First 3 Months of Temporary Disability")
    B(doc, "If the employee is temporarily disabled, the employer must pay 75% of the employee's earnings for the first 3 months;")
    B(doc, "Submit a claim for reimbursement to the Compensation Fund;")
    B(doc, "The Compensation Fund takes over payment from month 4 onwards (if disability continues).")

    P_bold(doc, "Step 5: Follow-Up and Closure")
    B(doc, "Monitor the employee's recovery and return-to-work progress;")
    B(doc, "Submit progress reports (Form W.Cl.5) and the final medical report (Form W.Cl.26) as required;")
    B(doc, "If permanent disability is assessed, submit the application for permanent disability benefits;")
    B(doc, "Retain all documentation for at least 4 years.")

    H2(doc, "2.3 Key Deadlines Summary")
    make_table(doc,
        ["Action", "Deadline"],
        [
            ["Report injury to employer", "As soon as possible"],
            ["Employer submits W.Cl.2", "Within 7 days"],
            ["First Medical Report (W.Cl.4)", "Within 14 days of first consultation"],
            ["Employer pays temp disability (first 3 months)", "On normal pay dates"],
            ["Report death", "Immediately"],
        ], bold_first_col=True)

    # ----- SOP 3: Disciplinary Hearing Procedure -----
    page_break(doc)
    H1(doc, "SOP 3: Disciplinary Hearing Procedure (LRA Compliant)")
    make_table(doc,
        ["Field", "Details"],
        [
            ["SOP Number", "SOP-SA-001-03"],
            ["Version", "1.0"],
            ["Department", "Human Resources / Line Management"],
            ["Effective Date", "[YYYY-MM-DD]"],
            ["Review Date", "[YYYY-MM-DD]"],
        ], bold_first_col=True)

    H2(doc, "3.1 Purpose")
    P(doc, "To establish a standardised, LRA-compliant disciplinary hearing procedure that ensures both substantive and procedural fairness in all disciplinary matters at [COMPANY NAME].")

    H2(doc, "3.2 Progressive Discipline Framework")
    P(doc, "Before initiating a formal hearing, managers should apply progressive discipline:")
    make_table(doc,
        ["Offence Severity", "Recommended Action", "Validity Period"],
        [
            ["Minor first offence", "Verbal warning (documented)", "3 months"],
            ["Minor repeated offence", "First written warning", "6 months"],
            ["Serious offence / repeated minor", "Final written warning", "12 months"],
            ["Gross misconduct / repeated after final warning", "Disciplinary hearing (dismissal possible)", "N/A"],
        ], bold_first_col=True)

    H2(doc, "3.3 Formal Hearing Procedure")

    P_bold(doc, "Step 1: Investigation")
    B(doc, "Investigate the alleged misconduct thoroughly before charging the employee;")
    B(doc, "Gather evidence (documents, CCTV, witness statements);")
    B(doc, "Consider precautionary suspension (with full pay) only if the employee's presence may interfere with the investigation or endanger others.")

    P_bold(doc, "Step 2: Charge Sheet and Notice")
    B(doc, "Draft the charge sheet with specific, detailed charges;")
    B(doc, "Issue written notice of the hearing at least 48 hours in advance;")
    B(doc, "Notify the employee of: the date, time, and venue; the specific charges; the right to representation (fellow employee or union representative); the right to call witnesses and present evidence; and the possible sanctions.")

    P_bold(doc, "Step 3: Hearing")
    B(doc, "Appoint an impartial chairperson (not the complainant or the employee's direct supervisor);")
    B(doc, "The initiator presents the case and calls witnesses;")
    B(doc, "The employee (or representative) cross-examines, presents a defence, and calls witnesses;")
    B(doc, "Record the proceedings (minutes or audio recording);")
    B(doc, "The chairperson may adjourn to consider the evidence.")

    P_bold(doc, "Step 4: Finding")
    B(doc, "The chairperson makes a finding (guilty/not guilty) on each charge;")
    B(doc, "If guilty, consider mitigating and aggravating factors (length of service, disciplinary record, remorse, personal circumstances, severity of the offence, impact on the business);")
    B(doc, "Determine the appropriate sanction.")

    P_bold(doc, "Step 5: Communication and Appeal")
    B(doc, "Communicate the outcome in writing within 3 working days;")
    B(doc, "Inform the employee of the right to internal appeal (within 5 working days);")
    B(doc, "Inform the employee of the right to refer the matter to the CCMA (within 30 calendar days of dismissal);")
    B(doc, "File all documentation (charge sheet, notice, hearing minutes, outcome letter, appeal).")

    # ----- SOP 4: Retrenchment Consultation (Section 189/189A) -----
    page_break(doc)
    H1(doc, "SOP 4: Retrenchment Consultation (Section 189/189A)")
    make_table(doc,
        ["Field", "Details"],
        [
            ["SOP Number", "SOP-SA-001-04"],
            ["Version", "1.0"],
            ["Department", "Human Resources / Executive Management"],
            ["Effective Date", "[YYYY-MM-DD]"],
            ["Review Date", "[YYYY-MM-DD]"],
        ], bold_first_col=True)

    H2(doc, "4.1 Purpose")
    P(doc, "To establish a standardised procedure for conducting retrenchments (dismissals for operational requirements) in compliance with Section 189 and Section 189A of the Labour Relations Act 66 of 1995.")

    H2(doc, "4.2 Applicability")
    make_table(doc,
        ["Scenario", "Applicable Section"],
        [
            ["Employer contemplating retrenchment of fewer than 50 employees", "Section 189"],
            ["Employer with 50+ employees contemplating retrenchment of 50 or more employees, or the number specified in the Act", "Section 189A (large-scale retrenchment)"],
        ], bold_first_col=True)

    H2(doc, "4.3 Procedure (Section 189)")
    P_bold(doc, "Step 1: Issue Section 189(3) Notice")
    B(doc, "Issue a written notice to the consulting parties (recognised trade union, workplace forum, or affected employees) disclosing:")
    B(doc, "  (a) The reasons for the proposed dismissals;")
    B(doc, "  (b) The alternatives considered before resorting to retrenchment;")
    B(doc, "  (c) The number of employees likely to be affected;")
    B(doc, "  (d) The proposed selection criteria;")
    B(doc, "  (e) The proposed timing of dismissals;")
    B(doc, "  (f) The severance pay proposed;")
    B(doc, "  (g) Any assistance the employer proposes (e.g., job seeking, retraining).")

    P_bold(doc, "Step 2: Meaningful Consultation")
    B(doc, "Engage in meaningful joint consultation with the consulting parties;")
    B(doc, "Consultation must be in good faith and aimed at reaching consensus on:")
    B(doc, "  (a) Measures to avoid dismissals (hiring freeze, short-time, reduced overtime, voluntary retrenchment);")
    B(doc, "  (b) Measures to minimise the number of dismissals;")
    B(doc, "  (c) Measures to mitigate the effects (outplacement, retraining, extended benefits);")
    B(doc, "  (d) The method for selecting employees to be dismissed;")
    B(doc, "  (e) The timing of dismissals;")
    B(doc, "  (f) Severance pay.")
    P(doc, "The employer must consider and respond to representations made by the consulting parties.")

    P_bold(doc, "Step 3: Selection Criteria")
    B(doc, "Apply fair and objective selection criteria (LIFO is the default unless an alternative is agreed or justified);")
    B(doc, "Selection criteria must not directly or indirectly discriminate on any prohibited ground;")
    B(doc, "Document the selection process and the rationale for each affected employee.")

    P_bold(doc, "Step 4: Notice and Severance Pay")
    B(doc, "Provide the statutory notice period (1 week / 2 weeks / 4 weeks, depending on length of service);")
    B(doc, "Pay severance pay of at least one week's remuneration per completed year of continuous service;")
    B(doc, "Issue the employee with a certificate of service;")
    B(doc, "Provide a letter confirming the retrenchment and the employee's right to refer a dispute to the CCMA within 30 days.")

    P_bold(doc, "Step 5: Section 189A Additional Requirements (Large-Scale)")
    B(doc, "If Section 189A applies, the consultation period must be at least 60 days;")
    B(doc, "Either party may request a facilitator from the CCMA;")
    B(doc, "Employees may strike in protest against the retrenchment (after following the required procedures);")
    B(doc, "The Labour Court may adjudicate the substantive and procedural fairness of the retrenchment.")

    # ----- SOP 5: EE Plan Review and Reporting -----
    page_break(doc)
    H1(doc, "SOP 5: Employment Equity Plan Review and Reporting")
    make_table(doc,
        ["Field", "Details"],
        [
            ["SOP Number", "SOP-SA-001-05"],
            ["Version", "1.0"],
            ["Department", "Human Resources"],
            ["Effective Date", "[YYYY-MM-DD]"],
            ["Review Date", "[YYYY-MM-DD]"],
        ], bold_first_col=True)

    H2(doc, "5.1 Purpose")
    P(doc, "To establish a standardised procedure for developing, reviewing, and reporting on the [COMPANY NAME] Employment Equity Plan in compliance with the Employment Equity Act 55 of 1998 (EEA), as amended.")

    H2(doc, "5.2 Annual Cycle")
    make_table(doc,
        ["Month", "Activity", "Responsible"],
        [
            ["January", "Collect workforce profile data (race, gender, disability by occupational level)", "HR / Payroll"],
            ["February", "Conduct analysis: compare workforce profile to EAP and sector targets", "HR / EE Committee"],
            ["March", "Review barriers to employment equity (policies, practices, workplace culture)", "EE Committee"],
            ["April", "Update EE Plan: set/revise numerical goals, affirmative action measures", "HR / EE Committee"],
            ["May", "Conduct income differential analysis (Section 27)", "HR / Finance"],
            ["June", "Present updated EE Plan to EE Consultative Forum/Committee for input", "HR"],
            ["July", "Finalise EE Plan and submit to CEO/Managing Director for approval", "HR / CEO"],
            ["August-September", "Prepare EEA2/EEA4 report for electronic submission", "HR"],
            ["1 October", "Submit EE Report to Department of Employment and Labour (150+ employees)", "HR"],
            ["15 January", "Submit EE Report to Department of Employment and Labour (<150 employees)", "HR"],
            ["Ongoing", "Monitor progress against targets and implement affirmative action measures", "HR / Line Managers"],
        ], bold_first_col=True)

    H2(doc, "5.3 EE Committee/Consultative Forum")
    B(doc, "Establish an EE Committee with representatives from all occupational levels, designated groups, and trade unions;")
    B(doc, "The committee must meet at least quarterly to review progress;")
    B(doc, "Minutes of meetings must be kept and made available to inspectors;")
    B(doc, "The committee provides input on the EE Plan, reviews barriers, and monitors progress.")

    H2(doc, "5.4 Display and Communication")
    B(doc, "A summary of the most recent EE Report must be displayed at a prominent place in the workplace accessible to all employees;")
    B(doc, "Employees must be informed of the contents of the EE Plan and their role in achieving equity targets;")
    B(doc, "The EE Plan must be made available for inspection by any employee or inspector upon request.")

    H2(doc, "5.5 Compliance Certificate")
    B(doc, "Before contracting with any organ of state, confirm that [COMPANY NAME] has a valid compliance certificate;")
    B(doc, "If the certificate has been refused, take immediate corrective action and reapply;")
    B(doc, "Monitor the certificate's expiry date and renew proactively.")

    H2(doc, "5.6 Record Retention")
    B(doc, "Retain all EE Plans, EE Reports, committee minutes, workforce analyses, and correspondence with the Department for at least five years;")
    B(doc, "Records must be produced upon request by a Department inspector.")

    # ----- Revision History & Approval -----
    page_break(doc)
    H2(doc, "Revision History")
    make_table(doc,
        ["Version", "Date", "Author", "Changes"],
        [
            ["1.0", "[YYYY-MM-DD]", "[Name]", "Initial release -- 5 SOPs"],
        ])

    H2(doc, "Approval Signatures")
    make_table(doc,
        ["Role", "Name", "Signature", "Date"],
        [
            ["Prepared By", "", "", ""],
            ["Reviewed By (Legal/HR)", "", "", ""],
            ["Approved By (Management)", "", "", ""],
        ])

    add_footer_line(doc, "SA Compliance SOP Pack")

    path = os.path.join(SOP_DIR, "SOP_SouthAfrica_Pack.docx")
    doc.save(path)
    return path


# ===================================================================
# MAIN
# ===================================================================
def main():
    # Ensure all output directories exist
    for d in [POLICY_DIR, ONBOARD_DIR, TRAIN_DIR, SOP_DIR]:
        os.makedirs(d, exist_ok=True)

    builders = [
        ("BCEA Policy", build_bcea),
        ("LRA Policy", build_lra),
        ("EEA Policy", build_eea),
        ("OHSA Policy", build_ohsa),
        ("COIDA Policy", build_coida),
        ("UIF/PAYE/SDL Policy", build_uif_paye_sdl),
        ("Onboarding SA", build_onboarding_sa),
        ("Training SA Labour Law", build_training_sa),
        ("SOP SA Pack", build_sop_sa),
    ]

    results = []
    for name, builder in builders:
        try:
            path = builder()
            size = os.path.getsize(path)
            results.append((name, path, size, "OK"))
            print(f"  [OK] {name}: {path} ({size:,} bytes)")
        except Exception as e:
            results.append((name, "", 0, str(e)))
            print(f"  [FAIL] {name}: {e}")

    print(f"\n{'='*70}")
    print(f"  Total files: {len(results)}")
    ok = sum(1 for _, _, _, s in results if s == "OK")
    print(f"  Success: {ok} / {len(results)}")
    if ok < len(results):
        print("  FAILURES:")
        for name, _, _, status in results:
            if status != "OK":
                print(f"    - {name}: {status}")
    print(f"{'='*70}")

    return ok == len(results)


if __name__ == "__main__":
    success = main()
    exit(0 if success else 1)
