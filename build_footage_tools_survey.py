#!/usr/bin/env python3
"""
Build Footage Tools - Tools & Resources Follow-Up Survey Analysis (July 2026)
Professional HR report format with Calibri, navy headers, and shaded tables.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

NAVY = RGBColor(0x1A, 0x3A, 0x5C)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
GRAY = RGBColor(0x66, 0x66, 0x66)
OUTPUT_PATH = "/home/aiciv/portal_uploads/to-portal/FootageTools_ToolsResources_SurveyAnalysis_July2026.docx"


def set_cell_shading(cell, color_hex):
    """Set cell background shading."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def style_header_row(row, fill="1A3A5C"):
    """Style a table header row: navy background, white bold text."""
    for cell in row.cells:
        set_cell_shading(cell, fill)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.name = "Calibri"
                run.font.size = Pt(10)


def make_table(doc, headers, rows, col_widths=None):
    """Create a bordered table with navy header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    style_header_row(table.rows[0])

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            run.font.color.rgb = BLACK
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Apply column widths if specified
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Inches(width)

    return table


def add_body(doc, text):
    """Add body text paragraph (Calibri 11pt)."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(2)
    return p


def add_heading1(doc, text):
    """Add Heading 1 styled paragraph (Calibri 16pt bold navy)."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    # Add bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="4" w:color="1A3A5C"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p


def add_heading2(doc, text):
    """Add Heading 2 styled paragraph (Calibri 13pt bold navy)."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_footer(doc):
    """Add footer to all sections."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.text = ""
        run = p.add_run("Prepared by The People Group | Confidential")
        run.font.name = "Calibri"
        run.font.size = Pt(8)
        run.font.italic = True
        run.font.color.rgb = GRAY
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_document():
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = BLACK

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # =====================================================================
    # TITLE PAGE
    # =====================================================================
    # Add vertical spacing before title
    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

    # Title lines - centered
    title_lines = [
        ("Footage Tools", 28, True),
        ("Tools & Resources Follow-Up Survey", 20, True),
        ("Analysis and Recommendations", 16, True),
        ("July 2026", 14, False),
        ("10 Responses | June 29 - July 7, 2026", 12, False),
        ("Prepared by The People Group", 12, False),
    ]

    for text, size, bold in title_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.color.rgb = NAVY
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.space_before = Pt(4)

    # PAGE BREAK
    doc.add_page_break()

    # =====================================================================
    # SECTION 1: Executive Summary
    # =====================================================================
    add_heading1(doc, "1. Executive Summary")

    add_body(doc,
        "Footage Tools conducted a Tools & Resources Follow-Up Survey in late June 2026, "
        "receiving 10 responses across 6 functional areas. The survey assessed employee "
        "satisfaction with current tools and systems, ease of requesting new resources, "
        "and the impact of any gaps on daily work."
    )

    add_body(doc,
        "Key findings: 70% of respondents report gaps in their tools or resources. "
        "30% have stopped raising requests entirely, believing they will not be addressed. "
        "The most affected departments are CNC Operations, Engineering/Quality Control, "
        "and Sales. Quality issues, workarounds, and morale decline are the primary consequences."
    )

    # =====================================================================
    # SECTION 2: Response Demographics
    # =====================================================================
    add_heading1(doc, "2. Response Demographics")

    make_table(doc,
        ["Functional Area", "Responses"],
        [
            ["Engineering / Quality Control", "2"],
            ["CNC Operations", "3"],
            ["Finance / HR", "1"],
            ["Sales / Customer Support", "1"],
            ["Shipping / Receiving / Purchasing", "2"],
            ["Service & Assembly", "1"],
            ["Total", "10"],
        ],
        col_widths=[4.5, 1.5]
    )

    # =====================================================================
    # SECTION 3: Current Satisfaction
    # =====================================================================
    add_heading1(doc, "3. Current Satisfaction")

    make_table(doc,
        ["Satisfaction Level", "Count", "Percentage"],
        [
            ["I have everything I need and it works well", "1", "10%"],
            ["I have most of what I need, but some things are missing, outdated, or frustrating", "6", "60%"],
            ["I regularly work around gaps - missing tools, slow systems, or outdated processes", "3", "30%"],
        ],
        col_widths=[4.0, 0.8, 1.2]
    )

    add_body(doc,
        "Only 1 respondent (Finance/HR) reports having everything they need. "
        "The remaining 9 employees experience some level of gap, with 3 employees (30%) "
        "reporting they regularly work around missing tools, slow systems, or outdated processes. "
        "The departments most affected are CNC Operations, Engineering/QC, and Sales."
    )

    # =====================================================================
    # SECTION 4: Ease of Requesting Resources
    # =====================================================================
    add_heading1(doc, "4. Ease of Requesting Resources")

    make_table(doc,
        ["Response", "Count", "Percentage"],
        [
            ["Very easy - I raise it and it gets addressed quickly", "1", "10%"],
            ["Somewhat easy - it takes time but usually happens", "5", "50%"],
            ["Difficult - requests get delayed or lost", "1", "10%"],
            ["I have stopped raising it", "3", "30%"],
        ],
        col_widths=[4.0, 0.8, 1.2]
    )

    add_body(doc,
        "This is the most concerning finding. Three employees (30%) have stopped raising "
        "resource requests entirely. This indicates a broken feedback loop where employees "
        "no longer believe their needs will be addressed. The three who stopped raising issues "
        "are in CNC Operations (2) and Sales (1) - departments critical to production and revenue."
    )

    add_body(doc,
        "When employees stop asking, leadership loses visibility into operational problems. "
        "Issues compound silently until they surface as quality failures, turnover, or customer complaints."
    )

    # =====================================================================
    # SECTION 5: Impact of Gaps
    # =====================================================================
    add_heading1(doc, "5. Impact of Gaps")

    make_table(doc,
        ["Impact (Multi-Select)", "Count"],
        [
            ["Rework, errors, or quality issues", "3"],
            ["Having to use workarounds or manual processes", "3"],
            ["Frustration or lower morale", "3"],
            ["Slower output or missed deadlines", "1"],
            ["Difficulty collaborating with other departments", "1"],
            ["None of the above", "2"],
        ],
        col_widths=[4.5, 1.5]
    )

    add_body(doc,
        "Quality issues, workarounds, and morale are tied as the top three impacts. "
        "These are not comfort complaints. Rework costs real money. Workarounds introduce risk. "
        "Morale decline drives turnover. The fact that all three are equally prevalent suggests "
        "a systemic resource gap, not isolated incidents."
    )

    # =====================================================================
    # SECTION 6: Top Requests by Theme
    # =====================================================================
    add_heading1(doc, "6. Top Requests by Theme")

    make_table(doc,
        ["Theme", "Specific Requests", "Priority"],
        [
            [
                "People / Capacity",
                "Add staff to Engineering and QC teams; Supervisor Lead for CNC Operations",
                "High"
            ],
            [
                "CRM / Systems Integration",
                "CRM integrated with ERP system (Sales); Claude AI subscription (Finance)",
                "High"
            ],
            [
                "Training",
                "CNC programming training; DCR/ECN filing training; refresher sessions on new tools; proper operator training and support system",
                "High"
            ],
            [
                "Equipment",
                "Small CMM machine for QC; 15,000 hydraulic pump; HSS drills for aluminum; upgraded computer",
                "Medium"
            ],
            [
                "Organization",
                "Better tool organization and availability; tool tracking to prevent loss; cataloging and standardizing jigs for assembly/service; kitting traveller packages with complete documentation",
                "Medium"
            ],
            [
                "Access",
                "Warehouse camera access for Shipping investigations; M drive access; tooling room access for operators",
                "Low"
            ],
        ],
        col_widths=[1.5, 3.5, 1.0]
    )

    # =====================================================================
    # SECTION 7: Department Spotlight - Red Flags
    # =====================================================================
    add_heading1(doc, "7. Department Spotlight - Red Flags")

    add_heading2(doc, "CNC Operations (3 respondents)")
    add_body(doc,
        "All three CNC respondents report gaps. One states that 'operators are not being "
        "trained properly and effectively for the tasks that are being given to them' and that "
        "the support system is failing. Another has stopped raising issues entirely. A third "
        "requests a Supervisor Lead to provide structure and oversight. This department needs "
        "immediate attention: a structured training plan, consideration of a supervisory hire, "
        "and a visible response to the concerns raised."
    )

    add_heading2(doc, "Sales / Customer Support (1 respondent)")
    add_body(doc,
        "The Sales respondent wants CRM-ERP integration and has stopped raising the request. "
        "When a revenue-generating employee gives up asking for a tool that would improve "
        "their effectiveness, the cost of inaction is invisible but real. This person should "
        "be re-engaged directly: acknowledge the request, provide a timeline or honest explanation."
    )

    add_heading2(doc, "Engineering / Quality Control (2 respondents)")
    add_body(doc,
        "Both Engineering respondents say the real problem is headcount, not equipment. One "
        "states: 'Adding more people to the team would allow the team to be more proactive "
        "instead of reactive.' One also requests a small CMM machine to 'mitigate human "
        "errors.' When quality professionals tell you they need tools to prevent errors, the "
        "cost of ignoring them shows up in scrap, rework, and customer returns."
    )

    # =====================================================================
    # SECTION 8: Recommended Actions
    # =====================================================================
    add_heading1(doc, "8. Recommended Actions")

    make_table(doc,
        ["Priority", "Action", "Expected Outcome"],
        [
            [
                "1 - Immediate",
                "CNC Operations: Build a structured training plan for operators. Address the 'operators are not being trained properly' feedback directly. Consider supervisory hire.",
                "Reduced rework, improved morale, operators stop feeling unsupported"
            ],
            [
                "2 - Immediate",
                "Re-engage Sales on CRM-ERP integration. Acknowledge the request, scope the project, provide a timeline or honest explanation for delay.",
                "Revenue team re-engaged, broken feedback loop repaired"
            ],
            [
                "3 - This Month",
                "Engineering/QC headcount conversation. Evaluate whether adding 1 person to each team is feasible. If not, be transparent about timeline.",
                "Shift from reactive to proactive quality management"
            ],
            [
                "4 - This Month",
                "Implement tool organization and tracking system across CNC and Service & Assembly. Standardize jig cataloging. Introduce kitting traveller packages.",
                "Reduced downtime searching for tools, fewer lost/misplaced items"
            ],
            [
                "5 - This Week",
                "Quick wins: Grant warehouse camera access (Shipping), M drive access (Shipping), stock HSS drills for aluminum (CNC), upgrade computer (Shipping), grant tooling room access to operators (CNC).",
                "Immediate visible response showing the survey led to action"
            ],
            [
                "6 - This Quarter",
                "Evaluate CMM machine purchase for Quality Control. Get quotes, assess ROI against current rework costs.",
                "Reduced measurement errors, faster deviation resolution"
            ],
            [
                "7 - Ongoing",
                "Establish a quarterly resource request review. Ensure every request gets a response (yes, no, or timeline). Never let employees reach 'I stopped asking.'",
                "Sustained feedback loop, employees trust the process"
            ],
        ],
        col_widths=[1.2, 3.0, 1.8]
    )

    # =====================================================================
    # SECTION 9: Conclusion
    # =====================================================================
    add_heading1(doc, "9. Conclusion")

    add_body(doc,
        "The survey reveals a workforce that largely knows what it needs but has mixed "
        "confidence that requests will be addressed. The most critical finding is that 30% "
        "of respondents have stopped asking. This is not a tools problem. It is a trust "
        "problem. The fastest way to rebuild that trust is visible action on the quick wins "
        "(camera access, M drive, HSS drills, computer upgrade) and direct engagement with "
        "the three employees who have given up."
    )

    add_body(doc,
        "Footage Tools has a skilled and engaged team. They are telling you exactly what "
        "they need. The question is whether leadership responds."
    )

    # =====================================================================
    # FOOTER
    # =====================================================================
    add_footer(doc)

    # =====================================================================
    # SAVE
    # =====================================================================
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Document saved to: {OUTPUT_PATH}")
    print(f"File size: {os.path.getsize(OUTPUT_PATH):,} bytes")


if __name__ == "__main__":
    build_document()
