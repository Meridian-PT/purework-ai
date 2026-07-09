#!/usr/bin/env python3
"""
Build Master To-Do List for Mike Daser
July 8, 2026 (Updated)
Output: /home/aiciv/portal_uploads/to-portal/MasterToDoList_July8_2026.docx
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# Constants
NAVY = RGBColor(0x1A, 0x3A, 0x5C)
GRAY = RGBColor(0x66, 0x66, 0x66)
RED = RGBColor(0xCC, 0x00, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY_BG = "F2F2F2"
CHECKBOX = "\u2610"  # Unicode ballot box


def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_borders(cell, color="999999", sz="4"):
    """Set borders on a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'</w:tcBorders>'
    )
    tcBorders_existing = tcPr.find(qn('w:tcBorders'))
    if tcBorders_existing is not None:
        tcPr.remove(tcBorders_existing)
    tcPr.append(tcBorders)


def set_paragraph_font(para, font_name="Calibri", font_size=Pt(11), color=BLACK, bold=False, italic=False):
    """Set font properties on a paragraph's run(s)."""
    for run in para.runs:
        run.font.name = font_name
        run.font.size = font_size
        run.font.color.rgb = color
        run.font.bold = bold
        run.font.italic = italic


def add_horizontal_rule(doc):
    """Add a thin horizontal rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="CCCCCC"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p


def add_section_heading(doc, text):
    """Add a client section heading: Calibri 14pt bold navy with thin bottom border."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = NAVY
    # Thin bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="1A3A5C"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p


def add_sub_label(doc, text):
    """Add a sub-label like 'This Week' in Calibri 10pt bold red."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RED
    return p


def add_todo_item(doc, text):
    """Add a to-do item with checkbox prefix, Calibri 11pt."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(f"{CHECKBOX}  {text}")
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK
    return p


def add_footer(doc):
    """Add footer: 'The People Group | Confidential' Calibri 8pt italic gray centered."""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.clear()
    run = p.add_run("The People Group | Confidential")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = GRAY


def build_summary_box(doc, data, total):
    """Build the summary box at the bottom with light gray background."""
    # Summary title
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Summary")
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = NAVY

    # Create summary table
    num_rows = len(data) + 2  # header + data rows + total row
    table = doc.add_table(rows=num_rows, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(3.5)
        row.cells[1].width = Inches(1.5)

    # Header row
    hdr = table.rows[0]
    for i, text in enumerate(["Client", "Open Items"]):
        cell = hdr.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = WHITE
        set_cell_shading(cell, "1A3A5C")
        set_cell_borders(cell, "999999", "4")

    # Data rows
    for row_idx, (client, count) in enumerate(data):
        row = table.rows[row_idx + 1]
        for col_idx, text in enumerate([client, str(count)]):
            cell = row.cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.name = "Calibri"
            run.font.size = Pt(10)
            run.font.color.rgb = BLACK
            set_cell_shading(cell, LIGHT_GRAY_BG)
            set_cell_borders(cell, "999999", "4")

    # TOTAL row
    total_row = table.rows[-1]
    for col_idx, text in enumerate(["TOTAL", str(total)]):
        cell = total_row.cells[col_idx]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = NAVY
        set_cell_shading(cell, LIGHT_GRAY_BG)
        set_cell_borders(cell, "999999", "4")

    return table


def main():
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = BLACK

    # Set margins
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    # ---- TITLE ----
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(2)
    run = title_p.add_run("Master To-Do List")
    run.font.name = "Calibri"
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = NAVY

    # ---- SUBTITLE ----
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_before = Pt(0)
    sub_p.paragraph_format.space_after = Pt(4)
    run = sub_p.add_run("July 8, 2026 (Updated)")
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.color.rgb = GRAY

    # ---- HORIZONTAL RULE ----
    add_horizontal_rule(doc)

    # =========================================================
    # PROPERTY VISTA
    # =========================================================
    add_section_heading(doc, "PROPERTY VISTA")
    add_sub_label(doc, "This Week")
    items_pv = [
        "AE Screening \u2014 Screen Account Executive candidates",
        "Payroll \u2014 Set up / review",
        "Bloom+ \u2014 Follow up",
        "Organize event",
        "401(k) compliance \u2014 Review 2025 testing results and action (Brennan Moore / NFP)",
        "Caroline comp restructure \u2014 Move from discretionary to variable. Awaiting feedback from Scott (back July 7)",
        "KS comp + medical accommodation",
        "Book readings, ship equipment",
        "Looker Dashboard \u2014 Review/setup",
        "Pull full list for Jason \u2014 key people to meet",
    ]
    for item in items_pv:
        add_todo_item(doc, item)

    # =========================================================
    # FOOTAGE TOOLS
    # =========================================================
    add_section_heading(doc, "FOOTAGE TOOLS")
    add_sub_label(doc, "This Week")
    items_ft = [
        "CNC salary data \u2014 Research and compile",
        "CCLP membership sponsorship \u2014 Meet with Mani",
        "Boltman H&S \u2014 Follow up",
        "Melissa re: Top Grading",
        "Deep Dive Survey \u2014 Tools & Resources",
        "Recruit Now \u2014 CNC Lathe Machinist, CNC Mill Turn Machinist, CNC Tool Room Coordinator, Service & Assembly Tech, Facilities & Equipment Maintenance Tech",
        "Working Alone Plan \u2014 Build and implement",
        "Eric Lim Offboarding \u2014 RRSP, Exit Interview",
    ]
    for item in items_ft:
        add_todo_item(doc, item)

    # =========================================================
    # GREENER SOLUTIONS
    # =========================================================
    add_section_heading(doc, "GREENER SOLUTIONS")
    add_sub_label(doc, "This Week")
    items_gs = [
        "Humi clean-up \u2014 Update codes for lieu time",
        "Contract Creation \u2014 Contract docs clean up in Humi",
    ]
    for item in items_gs:
        add_todo_item(doc, item)

    # =========================================================
    # STRATCO (Strategy Collective)
    # =========================================================
    add_section_heading(doc, "STRATCO (Strategy Collective)")
    add_sub_label(doc, "This Week")
    items_sc = [
        "Call NY Workers Comp",
        "Gusto NY Taxes",
    ]
    for item in items_sc:
        add_todo_item(doc, item)

    # =========================================================
    # HOST GENIUS
    # =========================================================
    add_section_heading(doc, "HOST GENIUS")
    add_sub_label(doc, "This Week")
    items_hg = [
        "Sales Lead recruit \u2014 Screen candidates",
    ]
    for item in items_hg:
        add_todo_item(doc, item)

    # =========================================================
    # PURE TECHNOLOGY
    # =========================================================
    add_section_heading(doc, "PURE TECHNOLOGY")
    add_sub_label(doc, "This Week")
    items_pt = [
        "D&O insurance \u2014 Approval needed (Corgi)",
    ]
    for item in items_pt:
        add_todo_item(doc, item)

    # =========================================================
    # MADISON AND WALL
    # =========================================================
    add_section_heading(doc, "MADISON AND WALL")
    add_sub_label(doc, "This Week")
    items_mw = [
        "Finalize 30-day handbook + sprint plan (update handbook)",
        "Benefits work with Amr \u2014 Close out (exec summary sent, awaiting response)",
        "RRSP provider selection \u2014 Collage integration or standalone GRSP+DPSP",
        "Brian employment agreement \u2014 Meet with Heather",
        "IT policy \u2014 Draft and review",
        "Job descriptions \u2014 Build out",
        "Heather signed contract \u2014 Follow up, load into system",
        "Red flags on MedCan \u2014 Medical underwriting concerns",
        "Employment agreements for Brian \u2014 Finalize",
        "Heather onboarding \u2014 Target August 1",
        "Benefits \u2014 Finalize enrollment",
    ]
    for item in items_mw:
        add_todo_item(doc, item)

    # =========================================================
    # BOUNTY HUNTER PUREBRAIN
    # =========================================================
    add_section_heading(doc, "BOUNTY HUNTER PUREBRAIN")
    add_sub_label(doc, "Active")
    items_bh = [
        "Bounty Hunter World / Andy Ryan \u2014 Partnership call done, waiting for his specs, then build out onboarding/training/development integration. Follow up Wednesday.",
    ]
    for item in items_bh:
        add_todo_item(doc, item)

    # =========================================================
    # OTHER
    # =========================================================
    add_section_heading(doc, "OTHER")
    add_sub_label(doc, "This Week")
    items_other = [
        "Miko Castillo \u2014 Next steps?",
    ]
    for item in items_other:
        add_todo_item(doc, item)

    # =========================================================
    # SUMMARY BOX
    # =========================================================
    summary_data = [
        ("Property Vista", 10),
        ("Footage Tools", 8),
        ("Greener Solutions", 2),
        ("StratCo", 2),
        ("Host Genius", 1),
        ("Pure Technology", 1),
        ("Madison and Wall", 11),
        ("Bounty Hunter", 1),
        ("Other", 1),
    ]
    build_summary_box(doc, summary_data, 37)

    # ---- FOOTER ----
    add_footer(doc)

    # Save
    output_path = "/home/aiciv/portal_uploads/to-portal/MasterToDoList_July8_2026.docx"
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Saved: {output_path}")
    print(f"Size: {os.path.getsize(output_path)} bytes")


if __name__ == "__main__":
    main()
