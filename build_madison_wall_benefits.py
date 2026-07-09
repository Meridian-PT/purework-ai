#!/usr/bin/env python3
"""
Build Madison & Wall Benefits Advisory document for Mike Daser.

Output: /home/aiciv/portal_uploads/to-portal/MadisonWall_Benefits_Advisory_July2026.docx

Formatting:
- Calibri throughout
- Body: 11pt
- Section headers: 14pt bold, NAVY
- Tables: bordered, header rows #1A3A5C with white bold text
- Footer: "Prepared by The People Group" on every page
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml
from docx.enum.section import WD_ORIENT
import os

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
OUTPUT_PATH = "/home/aiciv/portal_uploads/to-portal/MadisonWall_Benefits_Advisory_July2026.docx"

NAVY = RGBColor(0x1A, 0x3A, 0x5C)
DARK_TEXT = RGBColor(0x33, 0x33, 0x33)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GRAY = RGBColor(0x99, 0x99, 0x99)


# ===================================================================
# Helper functions
# ===================================================================

def create_doc():
    """Create document with Calibri 11pt Normal style."""
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = DARK_TEXT
    style.paragraph_format.space_after = Pt(6)

    # Configure Heading 1 (not used much, but set it)
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Calibri'
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = NAVY
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(6)

    # Heading 2 -- Section headers (14pt bold navy)
    h2 = doc.styles['Heading 2']
    h2.font.name = 'Calibri'
    h2.font.size = Pt(14)
    h2.font.bold = True
    h2.font.color.rgb = NAVY
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(4)

    # Heading 3
    h3 = doc.styles['Heading 3']
    h3.font.name = 'Calibri'
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = NAVY
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(4)

    return doc


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def style_header_row(row, color_hex="1A3A5C"):
    for cell in row.cells:
        set_cell_shading(cell, color_hex)
        for pr in cell.paragraphs:
            for run in pr.runs:
                run.font.color.rgb = WHITE
                run.font.bold = True
                run.font.size = Pt(10)
                run.font.name = "Calibri"


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '</w:tblBorders>')
    tblPr.append(borders)


def make_table(doc, headers, rows_data, col_widths=None, bold_first_col=False):
    """Create a table with navy header row, borders, and optional bold first column."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)

    # Set column widths if provided
    if col_widths:
        for i, width in enumerate(col_widths):
            table.columns[i].width = width

    hdr = table.rows[0]
    for i, h in enumerate(headers):
        hdr.cells[i].text = h
    style_header_row(hdr)

    for row_data in rows_data:
        row = table.add_row()
        for i, text in enumerate(row_data):
            cell = row.cells[i]
            cell.text = ""
            pr = cell.paragraphs[0]
            run = pr.add_run(str(text))
            run.font.size = Pt(10)
            run.font.name = "Calibri"
            if bold_first_col and i == 0:
                run.font.bold = True

    return table


def add_footer(doc, text="Prepared by The People Group"):
    """Add footer to all sections."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        run.font.color.rgb = GRAY
        run.italic = True


def P(doc, text):
    """Add a normal paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    return p


def P_bold(doc, text):
    """Add a bold paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.bold = True
    return p


def P_italic(doc, text):
    """Add an italic paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.italic = True
    return p


def H2(doc, text):
    doc.add_heading(text, level=2)


def H3(doc, text):
    doc.add_heading(text, level=3)


def add_mixed_para(doc, parts):
    """Add a paragraph with mixed formatting.
    parts is a list of (text, bold, italic) tuples.
    """
    p = doc.add_paragraph()
    for text, bold, italic in parts:
        run = p.add_run(text)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.bold = bold
        run.italic = italic
    return p


# ===================================================================
# Build the document
# ===================================================================

def build():
    doc = create_doc()

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # ---------------------------------------------------------------
    # EMAIL HEADER BLOCK
    # ---------------------------------------------------------------
    header_data = [
        ("TO:", "Amy, Brian"),
        ("CC:", "Amr Galal"),
        ("FROM:", "Mike Daser, The People Group"),
        ("RE:", "Healthcare Benefits Structure - Clarification and Recommendations"),
        ("DATE:", "July 8, 2026"),
    ]

    for label, value in header_data:
        p = doc.add_paragraph()
        run_label = p.add_run(label + "  ")
        run_label.font.name = "Calibri"
        run_label.font.size = Pt(11)
        run_label.bold = True
        run_value = p.add_run(value)
        run_value.font.name = "Calibri"
        run_value.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(2)

    # Horizontal rule via paragraph border
    p_hr = doc.add_paragraph()
    p_hr.paragraph_format.space_before = Pt(6)
    p_hr.paragraph_format.space_after = Pt(6)
    pPr = p_hr._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="1A3A5C"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)

    # ---------------------------------------------------------------
    # SALUTATION
    # ---------------------------------------------------------------
    P(doc, "Amy, Brian,")
    P(doc, "Great questions and I appreciate the thorough review. You are right to want full clarity before committing. Coming from the US system, the Canadian layered approach can feel like you are stacking redundant coverage when you are actually not. Let me walk through exactly what you have, what the new quotes add, and what I recommend.")

    # ---------------------------------------------------------------
    # UNDERSTANDING THE CANADIAN SYSTEM
    # ---------------------------------------------------------------
    H2(doc, "UNDERSTANDING THE CANADIAN SYSTEM (FOR AMERICANS)")

    P(doc, "In the US, your employer plan is your healthcare. In Canada, the government covers the core medical through OHIP (Ontario Health Insurance Plan) at no cost to anyone. What employers provide on top of that fills specific gaps that OHIP does not cover. Think of OHIP as Medicare for everyone, not just 65+.")

    add_mixed_para(doc, [
        ("OHIP covers: ", True, False),
        ("doctors, hospitals, ER, specialists, surgery, diagnostics", False, False),
    ])

    add_mixed_para(doc, [
        ("OHIP does NOT cover: ", True, False),
        ("dental, prescriptions, vision, physio, massage, psychology, semi-private hospital rooms", False, False),
    ])

    P(doc, "Everything below is what fills those gaps, in tiers.")

    # ---------------------------------------------------------------
    # TIER 1: OHIP
    # ---------------------------------------------------------------
    H2(doc, "TIER 1: OHIP (Government - Free)")

    P(doc, "Every Ontario resident gets this automatically. No cost to M&W. No enrollment needed beyond having a valid Ontario health card.")

    # ---------------------------------------------------------------
    # TIER 2: GROUP BENEFITS
    # ---------------------------------------------------------------
    H2(doc, "TIER 2: GROUP BENEFITS (Manulife via Bridgewell 360)")

    P(doc, "This is your core employer plan. Here is exactly what M&W is paying:")

    # Main premium table
    P_bold(doc, "Manulife Monthly Premium Breakdown - Quote Q0288893")

    make_table(doc,
        headers=["Benefit", "Coverage Details", "Monthly Premium"],
        rows_data=[
            ["Basic Life Insurance",
             "1x salary, up to $1M, reduces 50% at age 65",
             "$339.41"],
            ["AD&D",
             "1x salary, Enhanced Plan II",
             "$56.10"],
            ["Dependent Life",
             "Spouse $10,000 / Child $5,000",
             "$8.21"],
            ["Extended Health Care\n(2 families)",
             "100% Rx (pay-direct card, mandatory generic, no max), 100% paramedical ($1,000/practitioner/yr), semi-private hospital, $400 vision/2yr, $5M out-of-country emergency",
             "$728.06"],
            ["Dental (2 families)",
             "Basic 100%, $1,000/yr cap, 6-month recall, current year fee guide",
             "$238.10"],
            ["Long-Term Disability",
             "66.7% of first $2,500 + 50% of next $3,500 + 40% of remaining, $12K/month max, 112-day wait",
             "$170.40"],
            ["TOTAL CORE PLAN",
             "2 employees, both family coverage",
             "$1,540.27/month\n($18,483/year)"],
        ],
        bold_first_col=True
    )
    doc.add_paragraph()  # spacing

    P_bold(doc, "Optional Add-Ons (recommended):")

    make_table(doc,
        headers=["Add-On", "Description", "Monthly Premium"],
        rows_data=[
            ["Critical Illness", "$10,000 lump sum on diagnosis", "$49.94/month"],
            ["Healthcare Online\n(TELUS Virtual Care)", "24/7 virtual doctor access", "$7.10/month"],
            ["TOTAL WITH OPTIONS", "", "$1,597.31/month\n($19,168/year)"],
        ],
        bold_first_col=True
    )
    doc.add_paragraph()

    P_bold(doc, "Also included at no additional premium:")
    P(doc, "EFAP (unlimited counselling), Manulife mobile app, Medical Second Opinion service, Vision Care Partnership (FYidoctors/BonLook discounts), Survivor Benefit (24 months continued coverage for dependents).")

    # ---------------------------------------------------------------
    # COST SHARING RECOMMENDATION
    # ---------------------------------------------------------------
    H3(doc, "COST SHARING RECOMMENDATION")

    P(doc, "M&W pays 100% of everything EXCEPT Long-Term Disability. LTD should be employee-paid for one important reason: if LTD premiums are employer-paid, the benefit is taxable income to the employee when they collect. If the employee pays the LTD premium, the benefit is received tax-free. This is standard practice in Canada and protects your employees in the event they actually need to use it.")

    P(doc, "Additionally, I recommend M&W cover 100% of the Critical Illness and Healthcare Online add-ons. These are low-cost ($57.04/month combined) and signal that M&W invests in its people.")

    P_bold(doc, "Revised cost split:")

    make_table(doc,
        headers=["Paid By", "Amount", "Covers"],
        rows_data=[
            ["M&W pays",
             "$1,426.87/month\n($17,122/year)",
             "Life, AD&D, Dependent Life, Extended Health, Dental, Critical Illness, Healthcare Online"],
            ["Employee pays",
             "$170.40/month\n($2,045/year)\nsplit between 2 employees",
             "LTD only ($85.20/employee/month)\nThis ensures tax-free LTD benefits if ever needed"],
        ],
        bold_first_col=True
    )
    doc.add_paragraph()

    P_italic(doc, "This is not a cost-cutting measure. This is the correct tax-advantaged structure. Every benefits advisor in Canada will confirm this.")

    # ---------------------------------------------------------------
    # APPENDIX: PER-EMPLOYEE COST BREAKDOWN
    # ---------------------------------------------------------------
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    H2(doc, "Appendix: Per-Employee Cost Breakdown")

    P(doc, "The Manulife quote covers 2 employees with family coverage. Below is the per-employee breakdown so you can see exactly what each person costs, and how this scales if M&W adds employees in the future.")

    P_bold(doc, "Per Employee Monthly Cost (Family Coverage)")

    make_table(doc,
        headers=["Benefit", "Rate Per Employee", "How It Is Calculated"],
        rows_data=[
            ["Basic Life Insurance", "~$169.71*", "$0.363 per $1,000 of annual salary"],
            ["AD&D", "~$28.05*", "$0.060 per $1,000 of annual salary"],
            ["Dependent Life", "$4.10", "Flat rate per employee"],
            ["Extended Health Care (Family)", "$364.03", "Flat rate per family"],
            ["Dental (Family)", "$119.05", "Flat rate per family"],
            ["Long-Term Disability", "~$85.20*", "$0.710 per $100 of monthly insured earnings"],
            ["Per Employee Total", "~$770.14", ""],
        ],
        bold_first_col=True
    )
    doc.add_paragraph()

    P_italic(doc, "*Life, AD&D, and LTD rates are volume-based. The per-employee amount depends on individual salary. The figures above assume an even split of the $935,000 total insured volume across 2 employees.")

    P_bold(doc, "Optional Add-Ons (Per Employee)")

    make_table(doc,
        headers=["Benefit", "Rate Per Employee"],
        rows_data=[
            ["Critical Illness ($10,000 coverage)", "$24.97"],
            ["Healthcare Online (TELUS Virtual Care)", "$3.55"],
            ["Optional Total", "$28.52"],
        ],
        bold_first_col=True
    )
    doc.add_paragraph()

    P_bold(doc, "Summary: Per Employee vs. Total Annual Cost")

    make_table(doc,
        headers=["Category", "Per Employee/Month", "Per Employee/Year",
                 "Total (2 Employees)/Month", "Total (2 Employees)/Year"],
        rows_data=[
            ["Core Plan", "$770.14", "$9,242", "$1,540.27", "$18,483"],
            ["Core + Options", "$798.66", "$9,584", "$1,597.31", "$19,168"],
            ["M&W Pays (excl. LTD)", "$713.44", "$8,561", "$1,426.87", "$17,122"],
            ["Employee Pays (LTD only)", "$85.20", "$1,022", "$170.40", "$2,045"],
        ],
        bold_first_col=True
    )
    doc.add_paragraph()

    P_bold(doc, "Scaling: What Each New Employee Would Add")

    make_table(doc,
        headers=["Coverage Type", "Monthly Cost Per New Employee"],
        rows_data=[
            ["Single (no dependents)",
             "~$435 (EHC $99.29 + Dental $40.93 + Life/AD&D/LTD varies by salary + Dep Life N/A)"],
            ["Family",
             "~$770 (based on current family rates)"],
        ],
        bold_first_col=True
    )
    doc.add_paragraph()

    P_italic(doc, "Rates are guaranteed for 16 months (EHC, Dental) and 28 months (Life, AD&D, Dependent Life, LTD, Critical Illness) from the effective date. Adding employees does not change existing rates.")

    # ---------------------------------------------------------------
    # TIER 3: EXECUTIVE MEDICAL
    # ---------------------------------------------------------------
    H2(doc, "TIER 3: EXECUTIVE MEDICAL ASSESSMENT (Medcan) - RECOMMENDED")

    P(doc, "This is a comprehensive annual physical, not a doctor visit. Full-day executive screening: bloodwork panels, cardiac stress test, imaging, body composition, nutrition counselling, mental health screening. You take the detailed findings to your GP for ongoing care.")

    add_mixed_para(doc, [
        ("Cost: ", True, False),
        ("approximately $3,600 per person per year", False, False),
    ])
    add_mixed_para(doc, [
        ("Who pays: ", True, False),
        ("M&W, no employee contribution", False, False),
    ])
    add_mixed_para(doc, [
        ("Tax treatment: ", True, False),
        ("taxable benefit to the employee, but widely offered at executive level", False, False),
    ])

    add_mixed_para(doc, [
        ("MY TAKE: ", True, False),
        ("This IS a market expectation for executive-level roles in Canadian financial services. I recommend it. It is preventive, defensible, and genuinely valuable. Through Amr's preferred pricing, you may also access Cleveland Clinic or Lume Women's Health at similar rates.", False, False),
    ])

    # ---------------------------------------------------------------
    # TIER 4: CONCIERGE
    # ---------------------------------------------------------------
    H2(doc, "TIER 4: CONCIERGE / ONGOING CARE - OPTIONAL")

    P(doc, "This is what the application Amr sent is for. This provides a private GP relationship with same-day/next-day access, no walk-in waits.")

    make_table(doc,
        headers=["Option", "What You Get", "Approximate Cost"],
        rows_data=[
            ["Medcan Year Round Care", "Sophisticated walk-in clinic access", "$2,195/family/year"],
            ["Medcan Dedicated Care", "Personal GP relationship", "$3,600/person/year"],
            ["Cleveland Clinic Membership", "1 year included with executive medical", "+$300/person for family members"],
            ["Cleveland Clinic Chairman's", "Dedicated personal GP", "$10,000/person/year"],
            ["Lume Women's Health", "No membership programs yet", "N/A"],
        ],
        bold_first_col=True
    )
    doc.add_paragraph()

    add_mixed_para(doc, [
        ("MY HONEST TAKE: ", True, False),
        ("This is above market. Most Canadian executives, even in financial services, use OHIP for ongoing GP care and it works. Wait times for a GP can be longer than in the US, but with the Medcan annual assessment (Tier 3) you are already getting the most important preventive screening. Concierge medicine is a luxury differentiator, not a competitive necessity.", False, False),
    ])

    P_bold(doc, "If cost-conscious:")
    P(doc, "Skip Tier 4 entirely. OHIP + Manulife + Medcan annual gives you excellent coverage.")

    P_bold(doc, "If you want some GP access:")
    P(doc, "Medcan Year Round Care at $2,195/family is the best value. It gives you walk-in access without the $10K/person Chairman's price tag.")

    # ---------------------------------------------------------------
    # INTERNATIONAL HEALTH INSURANCE
    # ---------------------------------------------------------------
    H2(doc, "INTERNATIONAL HEALTH INSURANCE ($5M Lifetime) - WORTH CONSIDERING")

    P(doc, "This is separate from everything above. Amr quoted coverage for treatment outside Canada, including US facilities (Cleveland Clinic, etc.).")

    add_mixed_para(doc, [
        ("Coverage: ", True, False),
        ("$5M lifetime per person", False, False),
    ])
    add_mixed_para(doc, [
        ("Deductible options: ", True, False),
        ("$5,000 / $10,000 / $20,000 annually", False, False),
    ])
    add_mixed_para(doc, [
        ("Executive Medical rider: ", True, False),
        ("adults only (Medcan, Cleveland Clinic, Lume)", False, False),
    ])
    add_mixed_para(doc, [
        ("Rapid MRI/CT access: ", True, False),
        ("including medically necessary diagnostics", False, False),
    ])

    P(doc, "For Americans living in Canada who may want US facility access for certain procedures, this is genuinely valuable. Your Manulife plan already includes $5M out-of-country emergency coverage, but this international plan covers non-emergency treatment as well.")

    add_mixed_para(doc, [
        ("RECOMMENDATION: ", True, False),
        ("If you proceed, take the $10,000 or $20,000 deductible to keep premiums reasonable. The $5,000 deductible will be significantly more expensive for marginal additional value. A medical questionnaire is required (the application Amr sent).", False, False),
    ])

    # ---------------------------------------------------------------
    # HEALTH SPENDING ACCOUNT
    # ---------------------------------------------------------------
    H2(doc, "HEALTH SPENDING ACCOUNT (HSA) - RECOMMENDED")

    P(doc, "Amr is right that a standalone HSA through getmyhsa.com gives more flexibility than Manulife's built-in Cost Plus (which charges 10% admin on every claim, $25 min/$250 max per month).")

    add_mixed_para(doc, [
        ("What an HSA does: ", True, False),
        ("gives each employee a fixed annual dollar amount to spend on any CRA-eligible medical expense not covered by the group plan. It is 100% tax-deductible to M&W and tax-free to the employee. It is the most tax-efficient benefit you can offer.", False, False),
    ])

    add_mixed_para(doc, [
        ("Why M&W needs one: ", True, False),
        ("your Manulife dental plan does NOT include major dental (crowns, bridges, implants) or orthodontics because you need minimum 3 eligible employees. An HSA fills that gap immediately.", False, False),
    ])

    P_bold(doc, "Recommended HSA amounts:")

    make_table(doc,
        headers=["Employee Class", "Annual HSA Amount", "Rationale"],
        rows_data=[
            ["Executives/Partners",
             "$5,000 - $10,000",
             "Covers major dental gap, paramedical overages beyond $1,000/practitioner, portions of Tier 3/4 costs, any other CRA-eligible medical expenses"],
            ["Future general employees",
             "$1,500 - $2,500",
             "Proportional coverage (CRA requires the highest class cannot exceed the lowest by more than 10x)"],
        ],
        bold_first_col=True
    )
    doc.add_paragraph()

    add_mixed_para(doc, [
        ("CRA guideline: ", True, False),
        ("HSA allocation should not exceed 25% of the employee's gross salary. At executive compensation levels, $10,000 is well within this threshold.", False, False),
    ])

    # ---------------------------------------------------------------
    # SUMMARY TABLE
    # ---------------------------------------------------------------
    H2(doc, "SUMMARY: WHAT I RECOMMEND FOR M&W")

    make_table(doc,
        headers=["Tier", "Decision", "Annual Cost to M&W (est.)"],
        rows_data=[
            ["OHIP", "Free, automatic", "$0"],
            ["Manulife Group Benefits\n(excl. LTD)", "Yes, M&W pays 100%", "$17,122/year"],
            ["LTD", "Yes, EMPLOYEE-PAID\n(tax-free benefit)", "$0 to M&W"],
            ["Critical Illness +\nVirtual Care", "Yes, M&W pays", "Included in above"],
            ["HSA (getmyhsa.com)", "Yes, $5,000-$10,000/person", "$10,000-$20,000/year"],
            ["Medcan Annual\nAssessment", "Yes", "$7,200/year (2 people)"],
            ["Concierge Ongoing Care", "Optional, skip if\ncost-conscious", "$0 - $4,390"],
            ["International Health\nInsurance", "Worth considering,\n$10K deductible", "Quote from Amr"],
        ],
        bold_first_col=True
    )
    doc.add_paragraph()

    P_bold(doc, "TOTAL ESTIMATED ANNUAL COST TO M&W (recommended package):")
    P(doc, "$34,322 - $44,322 depending on HSA amount and concierge decision.")

    P(doc, "This gives you a benefits package that is competitive with any financial services firm in Toronto, with the correct tax structure, and without overpaying for coverage you do not need.")

    P(doc, "Happy to walk through any of this on a call next week.")

    # ---------------------------------------------------------------
    # SIGN OFF
    # ---------------------------------------------------------------
    doc.add_paragraph()  # extra space
    p_name = doc.add_paragraph()
    run = p_name.add_run("Mike Daser")
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.bold = True

    P(doc, "The People Group")
    P(doc, "mike.daser@thepeoplegroup.ca")

    # ---------------------------------------------------------------
    # FOOTER on every page
    # ---------------------------------------------------------------
    add_footer(doc)

    # ---------------------------------------------------------------
    # Save
    # ---------------------------------------------------------------
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Saved: {OUTPUT_PATH}")
    print(f"Size: {os.path.getsize(OUTPUT_PATH):,} bytes")


if __name__ == "__main__":
    build()
